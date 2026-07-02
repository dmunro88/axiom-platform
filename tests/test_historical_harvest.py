import json
import sqlite3
import tempfile
import unittest
from pathlib import Path

from docx import Document

from comparable_contract import confirm_extraction_result
from db import (
    get_conn,
    init_db,
    search_assignments,
    search_income_snapshots,
)
from harvest_contract import (
    ASSIGNMENT_CONTRACT_ID,
    INCOME_CONTRACT_ID,
    SCHEMA_VERSION,
)
from ingest import commit_extraction_result, run_extraction


def _build_fictional_report(path):
    document = Document()
    document.add_heading("Fictional Historical Appraisal", level=1)
    document.add_paragraph("Prepared for: Fictional Archive Bank")
    document.add_paragraph("Subject Property: 999 Fictional Subject Road")
    document.add_paragraph("Effective Date: January 15, 2025")
    document.add_paragraph("Date of Report: February 1, 2025")
    document.add_paragraph("Sales Comparison Approach: $1,200,000")
    document.add_paragraph("Income Approach: $1,250,000")
    document.add_paragraph("Cost Approach: $1,180,000")
    document.add_paragraph("Reconciled Value: $1,225,000")
    document.add_paragraph("Cap rates ranged from 7.5% to 8.5%.")
    document.save(path)


def _build_fictional_income_chart(path):
    document = Document()
    document.add_heading("Fictional Income Summary", level=1)
    document.add_paragraph("Period Year: 2025")
    document.add_paragraph("Period Type: Stabilized")
    document.add_paragraph("Potential Gross Income: $240,000")
    document.add_paragraph("Vacancy and Collection Loss: 5.0%")
    document.add_paragraph("Effective Gross Income: $228,000")
    document.add_paragraph("Total Operating Expenses: $78,000")
    document.add_paragraph("Expense Ratio: 34.21%")
    document.add_paragraph("Net Operating Income: $150,000")
    document.add_paragraph("Capitalization Rate Applied: 8.0%")
    document.save(path)


class HistoricalHarvestTests(unittest.TestCase):
    def test_contract_descriptor_matches_runtime(self):
        schema_path = (
            Path(__file__).resolve().parents[1]
            / "schemas"
            / "historical_harvest.v1.json"
        )
        schema = json.loads(schema_path.read_text(encoding="utf-8"))
        self.assertEqual(SCHEMA_VERSION, schema["schema_version"])
        self.assertEqual(
            ASSIGNMENT_CONTRACT_ID,
            schema["contracts"]["assignment"],
        )
        self.assertEqual(INCOME_CONTRACT_ID, schema["contracts"]["income"])

    def test_legacy_database_migrates_harvest_columns(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            db_path = Path(temp_dir) / "legacy.db"
            connection = sqlite3.connect(db_path)
            connection.executescript(
                """
                CREATE TABLE assignments (
                    assignment_id INTEGER PRIMARY KEY,
                    file_no TEXT
                );
                CREATE TABLE income_snapshots (
                    snapshot_id INTEGER PRIMARY KEY,
                    period_year INTEGER
                );
                """
            )
            connection.commit()
            connection.close()

            init_db(db_path, quiet=True)
            connection = sqlite3.connect(db_path)
            try:
                for table in ("assignments", "income_snapshots"):
                    columns = {
                        row[1]
                        for row in connection.execute(
                            f"PRAGMA table_info({table})"
                        ).fetchall()
                    }
                    self.assertTrue({
                        "identity_key",
                        "confidence",
                        "review_status",
                        "reviewed_at",
                        "source_record_json",
                    }.issubset(columns))
            finally:
                connection.close()

    def test_fictional_report_and_income_end_to_end(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            projects = root / "historical"
            assignment = (
                projects
                / "24C777 - Office - 999 Fictional Subject Road, Demo City"
            )
            assignment.mkdir(parents=True)
            _build_fictional_report(assignment / "REPORT 24C777.docx")
            _build_fictional_income_chart(assignment / "Income Chart.docx")

            staged_paths = run_extraction(
                projects,
                staged_dir=root / "staged",
            )
            self.assertEqual(1, len(staged_paths))
            staged = json.loads(staged_paths[0].read_text(encoding="utf-8"))
            conclusion = staged["assignment_record"]
            income = staged["income_snapshot"]

            self.assertEqual(ASSIGNMENT_CONTRACT_ID, conclusion["contract_id"])
            self.assertEqual(INCOME_CONTRACT_ID, income["contract_id"])
            self.assertEqual("unreviewed", conclusion["review"]["status"])
            self.assertEqual("unreviewed", income["review"]["status"])
            self.assertEqual("24C777", conclusion["data"]["file_no"])
            self.assertEqual("2025-01-15", conclusion["data"]["effective_date"])
            self.assertEqual(1_225_000, conclusion["data"]["reconciled_value"])
            self.assertEqual("SCA,IA,CA", conclusion["data"]["approaches"])
            self.assertEqual("stabilized", income["data"]["period_type"])
            self.assertEqual(0.05, income["data"]["vacancy_pct"])
            self.assertEqual(0.08, income["data"]["cap_rate_applied"])
            self.assertEqual(150_000, income["data"]["noi"])
            self.assertEqual(
                "Income Chart.docx",
                income["provenance"]["source_filename"],
            )

            conclusion["data"]["client"] = "Reviewed Fictional Archive Bank"
            conclusion["review_edits"] = [{
                "field": "client",
                "before": "fictional archive bank",
                "after": "Reviewed Fictional Archive Bank",
            }]
            confirmed = confirm_extraction_result(
                staged,
                reviewer="fictional-qa",
                reviewed_at="2026-07-01T12:00:00+00:00",
            )
            db_path = root / "harvest.db"
            init_db(db_path, quiet=True)
            connection = get_conn(db_path)
            try:
                with connection:
                    counts = commit_extraction_result(confirmed, connection)
                self.assertEqual(1, counts["assignments"])
                self.assertEqual(1, counts["income_snapshots"])

                with connection:
                    duplicate_counts = commit_extraction_result(
                        confirmed,
                        connection,
                    )
                self.assertEqual(
                    1,
                    duplicate_counts["duplicate_assignments"],
                )
                self.assertEqual(
                    1,
                    duplicate_counts["duplicate_income_snapshots"],
                )
            finally:
                connection.close()

            assignments = search_assignments(
                db_path,
                file_no="24c777",
            )
            snapshots = search_income_snapshots(
                db_path,
                period_year=2025,
                period_type="stabilized",
            )
            self.assertEqual(1, len(assignments))
            self.assertEqual("confirmed", assignments[0]["review_status"])
            self.assertEqual("Demo City", assignments[0]["address_city"])
            self.assertEqual(
                "Reviewed Fictional Archive Bank",
                assignments[0]["client"],
            )
            self.assertEqual(1, len(snapshots))
            self.assertEqual(0.3421, snapshots[0]["expense_ratio"])
            self.assertEqual("confirmed", snapshots[0]["review_status"])

    def test_unconfirmed_harvest_record_rolls_back(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            report_path = root / "REPORT TEST.docx"
            _build_fictional_report(report_path)
            raw = {
                "folder_name": "24C778 - Office - Fictional",
                "folder_meta": {
                    "file_no": "24C778",
                    "property_type": "Office",
                    "city": "Demo City",
                },
                "narrative": {
                    "data": {
                        "file_no": "24C778",
                        "reconciled_value": 500000,
                    },
                    "confidence": {"reconciled_value": "medium"},
                },
                "income_data": {},
                "comps": [],
                "lease_comps": [],
                "sources": [str(report_path)],
                "assignment_source": str(report_path),
            }
            confirmed = confirm_extraction_result(raw, reviewer="qa")
            confirmed["assignment_record"]["review"]["status"] = "unreviewed"

            db_path = root / "rollback.db"
            init_db(db_path, quiet=True)
            connection = get_conn(db_path)
            try:
                with self.assertRaisesRegex(ValueError, "not been confirmed"):
                    with connection:
                        commit_extraction_result(confirmed, connection)
                self.assertEqual(
                    0,
                    connection.execute(
                        "SELECT count(*) FROM assignments"
                    ).fetchone()[0],
                )
                self.assertEqual(
                    0,
                    connection.execute(
                        "SELECT count(*) FROM source_documents"
                    ).fetchone()[0],
                )
            finally:
                connection.close()


if __name__ == "__main__":
    unittest.main()
