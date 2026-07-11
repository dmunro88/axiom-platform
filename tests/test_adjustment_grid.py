"""
tests/test_adjustment_grid.py — Axiom Platform

Regression coverage for adjustment_grid.py (Phase 6 Adjustment Grid,
steps 5-6), the module that reads the sca_adjustment_grid /
land_adjustment_grid / sca_qualitative / land_qualitative workbook tabs
and injects each as a Word table at its [[..._GRID_BLOCK]] marker.

Includes a regression test for a real bug found while building the
DEMO-001 fixture: land_adjustment_grid (and any grid sheet with a MEAN /
summary section below its comp rows) was having those summary rows
misread as comps, because the original row scan collected any row with
a non-blank value in a header column rather than checking that the row
was actually a comp row. Fixed by anchoring on the "Sale No. N" label
every comp row is written with, and stopping the scan at the first row
that doesn't match.
"""

import sys
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import openpyxl
from docx import Document

from adjustment_grid import (
    AdjustmentGridError,
    GRIDS,
    _format_value,
    inject_adjustment_grid_block,
    inject_all_adjustment_grids,
    read_grid_rows,
)


def _make_workbook(tmp_path, sheet_name="test_grid", headers=None):
    """Build a minimal workbook with one grid sheet, header row at 6."""
    headers = headers or {
        "A": "Comp",
        "B": "Location",
        "C": "Sale Price ($)",
        "D": "Time Adj %",
        "E": "Net Adjustment %",
        "F": "Indicated Value ($/SF)",
    }
    wb = openpyxl.Workbook()
    wb.active.title = "Intake"
    ws = wb.create_sheet(sheet_name)
    for col, text in headers.items():
        ws[f"{col}6"] = text
    workbook_path = tmp_path / "workbook.xlsx"
    wb.save(workbook_path)
    wb.close()
    return workbook_path


def _make_docx(tmp_path, markers):
    """Build a minimal docx with one paragraph per marker."""
    doc = Document()
    doc.add_paragraph("Report Heading")
    for marker in markers:
        doc.add_paragraph(f"[[{marker}]]")
    doc.add_paragraph("Closing text.")
    doc_path = tmp_path / "report.docx"
    doc.save(doc_path)
    return doc_path


class ReadGridRowsTests(unittest.TestCase):
    def test_reads_populated_comp_rows(self):
        with _tmp_dir() as tmp_path:
            workbook_path = _make_workbook(tmp_path)
            wb = openpyxl.load_workbook(workbook_path)
            ws = wb["test_grid"]
            ws["B7"] = "123 Fictional Rd"
            ws["A7"] = "Sale No. 1"
            ws["C7"] = 450000
            ws["D7"] = 0.05
            ws["E7"] = 0.02
            ws["F7"] = 9.5
            ws["A8"] = "Sale No. 2"
            ws["B8"] = "456 Sample Ave"
            ws["C8"] = 380000
            wb.save(workbook_path)
            wb.close()

            headers, rows = read_grid_rows(workbook_path, "test_grid")

            self.assertEqual(
                ["Comp", "Location", "Sale Price ($)", "Time Adj %",
                 "Net Adjustment %", "Indicated Value ($/SF)"],
                headers,
            )
            self.assertEqual(2, len(rows))
            self.assertEqual("Sale No. 1", rows[0]["Comp"])
            self.assertEqual("123 Fictional Rd", rows[0]["Location"])
            self.assertEqual("$450,000.00", rows[0]["Sale Price ($)"])
            self.assertEqual("5.0%", rows[0]["Time Adj %"])
            self.assertEqual("Sale No. 2", rows[1]["Comp"])
            self.assertEqual("$380,000.00", rows[1]["Sale Price ($)"])

    def test_skips_fully_blank_comp_rows(self):
        """A row whose Comp label is pre-filled ("Sale No. N") but has no
        other data yet (the template's unused capacity rows) must not
        surface as a populated row."""
        with _tmp_dir() as tmp_path:
            workbook_path = _make_workbook(tmp_path)
            wb = openpyxl.load_workbook(workbook_path)
            ws = wb["test_grid"]
            ws["A7"] = "Sale No. 1"
            ws["B7"] = "123 Fictional Rd"
            for r in range(8, 11):
                ws[f"A{r}"] = f"Sale No. {r - 6}"
            wb.save(workbook_path)
            wb.close()

            _, rows = read_grid_rows(workbook_path, "test_grid")
            self.assertEqual(1, len(rows))
            self.assertEqual("123 Fictional Rd", rows[0]["Location"])

    def test_stops_at_mean_and_summary_rows_below_comps(self):
        """Regression test: a MEAN row and/or a conclusion section below
        the comp rows (as land_adjustment_grid has) must never be read as
        comp rows, even though they can have non-blank values sitting in
        the same header columns by coincidence of position. This is
        exactly the bug found building the DEMO-001 fixture: the MEAN row
        and 4 "LAND VALUE CONCLUSION" rows were showing up as 5 extra
        phantom comps."""
        with _tmp_dir() as tmp_path:
            workbook_path = _make_workbook(tmp_path)
            wb = openpyxl.load_workbook(workbook_path)
            ws = wb["test_grid"]
            ws["A7"] = "Sale No. 1"
            ws["B7"] = "123 Fictional Rd"
            ws["C7"] = 450000
            ws["A8"] = "Sale No. 2"
            ws["B8"] = "456 Sample Ave"
            ws["C8"] = 380000
            # Unused template capacity rows 3-10, all comp-labeled but blank.
            for r in range(9, 17):
                ws[f"A{r}"] = f"Sale No. {r - 6}"
            # Summary section below the comp rows, mirroring
            # land_adjustment_grid's real structure.
            ws["A17"] = "MEAN"
            ws["C17"] = 9.21
            ws["A19"] = "LAND VALUE CONCLUSION"
            ws["A20"] = "Indicated Range ($/SF):"
            ws["B20"] = "9.32"
            ws["C20"] = "9.96"
            ws["A21"] = "Concluded Land Value ($/SF):"
            ws["B21"] = "9.25"

            wb.save(workbook_path)
            wb.close()

            _, rows = read_grid_rows(workbook_path, "test_grid")
            self.assertEqual(2, len(rows))
            self.assertEqual(
                ["Sale No. 1", "Sale No. 2"],
                [row["Comp"] for row in rows],
            )

    def test_missing_sheet_returns_empty(self):
        with _tmp_dir() as tmp_path:
            workbook_path = _make_workbook(tmp_path)
            headers, rows = read_grid_rows(workbook_path, "does_not_exist")
            self.assertEqual([], headers)
            self.assertEqual([], rows)

    def test_missing_anchor_header_raises_loudly(self):
        """If the sheet's header row doesn't have "Comp" in column A, this
        injector must refuse to guess at a different layout (per the
        design doc's drift-protection requirement) rather than silently
        reading garbage."""
        with _tmp_dir() as tmp_path:
            workbook_path = _make_workbook(
                tmp_path,
                headers={"A": "Sale", "B": "Location"},
            )
            with self.assertRaises(AdjustmentGridError):
                read_grid_rows(workbook_path, "test_grid")

    def test_comp_row_past_max_data_row_raises_loudly(self):
        """A sheet hand-expanded with a real comp row sitting past
        MAX_DATA_ROW must fail loudly rather than silently drop that comp
        from every delivered report (Fable adversarial review finding A5)."""
        with _tmp_dir() as tmp_path:
            from adjustment_grid import MAX_DATA_ROW

            workbook_path = _make_workbook(tmp_path)
            wb = openpyxl.load_workbook(workbook_path)
            ws = wb["test_grid"]
            ws["A7"] = "Sale No. 1"
            ws["B7"] = "123 Fictional Rd"
            overflow_row = MAX_DATA_ROW + 1
            ws[f"A{overflow_row}"] = f"Sale No. {overflow_row - 6}"
            ws[f"B{overflow_row}"] = "999 Overflow Ave"
            wb.save(workbook_path)
            wb.close()

            with self.assertRaises(AdjustmentGridError):
                read_grid_rows(workbook_path, "test_grid")


class FormatValueTests(unittest.TestCase):
    def test_percentage_header_formats_as_percent(self):
        self.assertEqual("5.0%", _format_value("Time Adj %", 0.05))
        self.assertEqual("-2.0%", _format_value("Topography Adj %", -0.02))

    def test_currency_header_formats_with_dollar_sign(self):
        self.assertEqual("$450,000.00", _format_value("Sale Price ($)", 450000))
        self.assertEqual(
            "$9.86", _format_value("Adjusted Price ($/SF)", 9.86),
        )

    def test_overall_header_formats_to_two_decimals_plain(self):
        self.assertEqual("1.00", _format_value("Overall", 1))
        self.assertEqual("0.17", _format_value("Overall", 0.166666667))

    def test_plain_integer_and_float(self):
        self.assertEqual("3", _format_value("Site Area (SF)", 3))
        self.assertEqual("48,000.00", _format_value("Site Area (SF)", 48000.0))

    def test_blank_and_none_render_empty(self):
        self.assertEqual("", _format_value("Notes", None))
        self.assertEqual("", _format_value("Notes", ""))

    def test_date_formats_month_day_year(self):
        import datetime
        self.assertEqual(
            "05/01/2025",
            _format_value("Sale Date", datetime.date(2025, 5, 1)),
        )

    def test_plain_text_is_stripped(self):
        self.assertEqual("X", _format_value("Flood Zone", "  X  "))

    def test_excel_error_token_is_flagged_not_shown_verbatim(self):
        """A formula-error cached value (e.g. a blank Comp GBA making
        Ratio (Ac/As) divide by zero) must be surfaced as an unmistakable
        error, not blend in as if it were a normal text cell like a Notes
        or Flood Zone entry (Fable adversarial review finding A3)."""
        self.assertEqual(
            "[FORMULA ERROR -- #DIV/0!]",
            _format_value("Ratio (Ac/As)", "#DIV/0!"),
        )
        self.assertEqual(
            "[FORMULA ERROR -- #REF!]",
            _format_value("Indicated Value ($/SF)", "#REF!"),
        )

    def test_rate_header_without_percent_sign_still_formats_as_percent(self):
        """"Monthly Mkt Rate" is a real header on sca_adjustment_grid /
        land_adjustment_grid that carries a fractional rate (e.g. 0.005 =
        0.5%/month) but, unlike "Time Adj %", has no literal "%" in its
        header text. The old plain-number branch rounded it straight to
        "0.01", losing virtually all of its meaning (Fable adversarial
        review finding A4)."""
        self.assertEqual("0.5%", _format_value("Monthly Mkt Rate", 0.005))

    def test_dollar_header_without_price_or_value_token_still_gets_dollar_sign(self):
        """"Size Adj $/SF" is a real header (sca_adjustment_grid) that
        contains neither "price" nor "value", so the old compound condition
        (token AND unit) never matched it and it rendered with no dollar
        sign at all (Fable adversarial review finding A4)."""
        self.assertEqual("$0.95", _format_value("Size Adj $/SF", 0.948821765913758))
        self.assertEqual("-$2.32", _format_value("Size Adj $/SF", -2.32243661190965))

    def test_int_header_gets_thousands_separator(self):
        """A plain int cell (e.g. Comp GBA (SF)) used to bypass the comma
        formatting entirely -- 24000 rendered as "24000" instead of
        "24,000" (Fable adversarial review finding A4)."""
        self.assertEqual("24,000", _format_value("Comp GBA (SF)", 24000))

    def test_text_date_is_reparsed_to_month_day_year(self):
        """A Sale Date typed as plain text rather than a real Excel date
        (openpyxl then hands back a str, not a date/datetime) used to skip
        the strftime branch entirely and render however it was typed, e.g.
        ISO "2025-03-15" instead of "03/15/2025" like every real date cell
        gets (Fable adversarial review finding A4)."""
        self.assertEqual("03/15/2025", _format_value("Sale Date", "2025-03-15"))
        self.assertEqual("03/15/2025", _format_value("Sale Date", "03/15/2025"))

    def test_unparseable_text_date_falls_back_to_original_text(self):
        """A date header with a value that doesn't match any known text
        date format must still render (as typed) rather than raise or
        disappear."""
        self.assertEqual("TBD", _format_value("Sale Date", "TBD"))


class InjectAdjustmentGridBlockTests(unittest.TestCase):
    def test_injects_table_and_removes_marker(self):
        with _tmp_dir() as tmp_path:
            workbook_path = _make_workbook(tmp_path)
            wb = openpyxl.load_workbook(workbook_path)
            ws = wb["test_grid"]
            ws["A7"] = "Sale No. 1"
            ws["B7"] = "123 Fictional Rd"
            ws["C7"] = 450000
            wb.save(workbook_path)
            wb.close()

            doc_path = _make_docx(tmp_path, ["TEST_GRID_BLOCK"])

            n = inject_adjustment_grid_block(
                doc_path, workbook_path, "TEST_GRID_BLOCK", "test_grid",
            )
            self.assertEqual(1, n)

            doc = Document(str(doc_path))
            marker_found = any(
                "[[TEST_GRID_BLOCK]]" in p.text for p in doc.paragraphs
            )
            self.assertFalse(marker_found)
            self.assertEqual(1, len(doc.tables))
            table = doc.tables[0]
            self.assertEqual(
                ["Comp", "Location", "Sale Price ($)", "Time Adj %",
                 "Net Adjustment %", "Indicated Value ($/SF)"],
                [cell.text for cell in table.rows[0].cells],
            )
            self.assertEqual("Sale No. 1", table.rows[1].cells[0].text)
            self.assertEqual("123 Fictional Rd", table.rows[1].cells[1].text)
            self.assertEqual("$450,000.00", table.rows[1].cells[2].text)

    def test_marker_not_found_returns_zero(self):
        with _tmp_dir() as tmp_path:
            workbook_path = _make_workbook(tmp_path)
            wb = openpyxl.load_workbook(workbook_path)
            ws = wb["test_grid"]
            ws["A7"] = "Sale No. 1"
            ws["B7"] = "123 Fictional Rd"
            wb.save(workbook_path)
            wb.close()

            doc_path = _make_docx(tmp_path, ["OTHER_BLOCK"])

            n = inject_adjustment_grid_block(
                doc_path, workbook_path, "TEST_GRID_BLOCK", "test_grid",
            )
            self.assertEqual(0, n)

    def test_no_populated_rows_leaves_marker_in_place(self):
        with _tmp_dir() as tmp_path:
            workbook_path = _make_workbook(tmp_path)
            doc_path = _make_docx(tmp_path, ["TEST_GRID_BLOCK"])

            n = inject_adjustment_grid_block(
                doc_path, workbook_path, "TEST_GRID_BLOCK", "test_grid",
            )
            self.assertEqual(0, n)

            doc = Document(str(doc_path))
            marker_found = any(
                "[[TEST_GRID_BLOCK]]" in p.text for p in doc.paragraphs
            )
            self.assertTrue(marker_found)


class InjectAllAdjustmentGridsTests(unittest.TestCase):
    def test_injects_each_grid_independently(self):
        """Build a document/workbook pair with all 4 real grid markers and
        sheet names -- one populated, others missing/empty -- and confirm
        each is handled independently (a missing sheet for one grid
        doesn't block the others), matching the real GRIDS wiring used by
        axiom.py's deliver stage."""
        with _tmp_dir() as tmp_path:
            wb = openpyxl.Workbook()
            wb.active.title = "Intake"
            ws = wb.create_sheet("sca_adjustment_grid")
            ws["A6"] = "Comp"
            ws["B6"] = "Location"
            ws["A7"] = "Sale No. 1"
            ws["B7"] = "100 Fictional Way"

            ws2 = wb.create_sheet("land_adjustment_grid")
            ws2["A6"] = "Comp"
            ws2["B6"] = "Location"
            # land_adjustment_grid has no populated rows in this test --
            # should inject 0 without affecting the other blocks.

            workbook_path = tmp_path / "workbook.xlsx"
            wb.save(workbook_path)
            wb.close()

            doc_path = _make_docx(tmp_path, list(GRIDS.keys()))

            results = inject_all_adjustment_grids(doc_path, workbook_path)

            self.assertEqual(1, results["SCA_ADJUSTMENT_GRID_BLOCK"])
            self.assertEqual(0, results["LAND_ADJUSTMENT_GRID_BLOCK"])
            self.assertEqual(0, results["SCA_QUALITATIVE_GRID_BLOCK"])
            self.assertEqual(0, results["LAND_QUALITATIVE_GRID_BLOCK"])

            doc = Document(str(doc_path))
            # Populated block's marker gone, empty blocks' markers remain.
            texts = [p.text for p in doc.paragraphs]
            self.assertFalse(
                any("[[SCA_ADJUSTMENT_GRID_BLOCK]]" in t for t in texts)
            )
            self.assertTrue(
                any("[[LAND_ADJUSTMENT_GRID_BLOCK]]" in t for t in texts)
            )


def _tmp_dir():
    import tempfile

    class _Ctx:
        def __enter__(self):
            self._td = tempfile.TemporaryDirectory()
            return Path(self._td.name)

        def __exit__(self, *a):
            self._td.cleanup()

    return _Ctx()


if __name__ == "__main__":
    unittest.main()
