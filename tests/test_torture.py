import base64
import json
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

import openpyxl
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

import axiom
from comp_builder import inject_comp_section
from dilmore import dilmore_adj_pct, dilmore_factor
from field_registry import inventory_templates
from fill_engine import fill_document
from media_blocks import MAX_IMAGE_BYTES, inject_media_blocks
from validation import find_docx_placeholders, validate_assignment


TINY_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk"
    "/x8AAusB9Wl2nWQAAAAASUVORK5CYII="
)


def build_assignment(root, template_text, variables):
    root = Path(root)
    assignment = root / "TEST-999_Fictional_Client"
    templates = root / "templates"
    assignment.mkdir()
    templates.mkdir()

    workbook = openpyxl.Workbook()
    workbook.active.title = "outputs"
    workbook.active.append(["Label", "Key", "Raw", "Formatted"])
    workbook.save(assignment / "workbook.xlsx")
    workbook.close()
    (assignment / "TEST-999_variables.json").write_text(
        json.dumps(variables),
        encoding="utf-8",
    )

    document = Document()
    document.add_paragraph(template_text)
    document.save(templates / "report.docx")
    return assignment, templates, {"documents": [{"template": "report.docx"}]}


class TortureTests(unittest.TestCase):
    def test_assignment_lookup_does_not_prefix_match_another_file_number(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            assignments = Path(temp_dir)
            (assignments / "TEST-010_Fictional").mkdir()
            with patch.object(axiom, "ASSIGNMENTS_DIR", assignments):
                self.assertIsNone(axiom._find_assignment("TEST-01"))
                self.assertEqual(
                    assignments / "TEST-010_Fictional",
                    axiom._find_assignment("TEST-010"),
                )

    def test_new_assignment_rejects_pathlike_file_numbers(self):
        invalid_values = (
            "../escape",
            r"C:\escape",
            "TEST/001",
            "TEST..001",
            "X" * 65,
        )
        for file_no in invalid_values:
            with self.subTest(file_no=file_no), tempfile.TemporaryDirectory() as temp_dir:
                assignments = Path(temp_dir) / "assignments"
                assignments.mkdir()
                with patch.object(axiom, "ASSIGNMENTS_DIR", assignments):
                    self.assertFalse(
                        axiom.cmd_new([file_no, "Fictional Client"])
                    )
                self.assertEqual([], list(assignments.iterdir()))

    def test_client_name_is_sanitized_inside_assignment_root(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            assignments = root / "assignments"
            assignments.mkdir()
            workbook_template = root / "workbook.xlsx"
            workbook = openpyxl.Workbook()
            workbook.save(workbook_template)
            workbook.close()

            with (
                patch.object(axiom, "ASSIGNMENTS_DIR", assignments),
                patch.object(axiom, "WORKBOOK_TPL", workbook_template),
            ):
                axiom.cmd_new(
                    [
                        "TEST-001",
                        r"Fictional\..\Client",
                        "/",
                        'QA:*?"',
                    ]
                )

            created = list(assignments.iterdir())
            self.assertEqual(1, len(created))
            self.assertEqual(assignments.resolve(), created[0].resolve().parent)
            self.assertNotIn("..", created[0].name)
            self.assertFalse(
                any(character in created[0].name for character in '<>:"/\\|?*')
            )

    def test_blank_and_null_values_do_not_satisfy_placeholders(self):
        for value in ("", "   ", None, "None"):
            with self.subTest(value=value), tempfile.TemporaryDirectory() as temp_dir:
                assignment, templates, config = build_assignment(
                    temp_dir,
                    "[[REQUIRED_VALUE]]",
                    {"REQUIRED_VALUE": value},
                )
                result = validate_assignment(assignment, templates, config)
                self.assertFalse(result["ready"])
                self.assertEqual(["REQUIRED_VALUE"], result["missing"])

    def test_zero_and_false_are_preserved_as_real_values(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            template = Path(temp_dir) / "template.docx"
            output = Path(temp_dir) / "output.docx"
            document = Document()
            document.add_paragraph("[[ZERO_VALUE]] / [[FALSE_VALUE]]")
            document.save(template)

            result = fill_document(
                template,
                output,
                {"ZERO_VALUE": 0, "FALSE_VALUE": False},
            )
            self.assertEqual([], result["missing"])
            self.assertIn("0 / False", Document(output).paragraphs[0].text)

    def test_split_run_placeholder_is_detected_in_body_and_header(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "split.docx"
            document = Document()
            paragraph = document.add_paragraph()
            paragraph.add_run("[[BODY_")
            paragraph.add_run("VALUE]]")
            header = document.sections[0].header.paragraphs[0]
            header.add_run("[[HEADER_")
            header.add_run("VALUE]]")
            document.save(path)

            self.assertEqual(
                ["BODY_VALUE", "HEADER_VALUE"],
                find_docx_placeholders(path),
            )

    def test_split_run_template_placeholder_enters_contract_inventory(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            paragraph_doc = Document()
            paragraph = paragraph_doc.add_paragraph()
            paragraph.add_run("[[SPLIT_")
            paragraph.add_run("CONTRACT_KEY]]")
            paragraph_doc.save(root / "report.docx")
            stages = {
                "deliver": {
                    "documents": [{"template": "report.docx"}],
                }
            }
            self.assertEqual(
                {"SPLIT_CONTRACT_KEY"},
                inventory_templates(root, stages)["deliver"],
            )

    def test_boolean_conditional_flag_keeps_enabled_section(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            template = Path(temp_dir) / "template.docx"
            output = Path(temp_dir) / "output.docx"
            document = Document()
            document.styles.add_style(
                "MainSectionHeading",
                WD_STYLE_TYPE.PARAGRAPH,
            )
            document.add_paragraph("Cost Approach", style="MainSectionHeading")
            document.add_paragraph("Cost content")
            document.add_paragraph(
                "Sales Comparison Approach",
                style="MainSectionHeading",
            )
            document.save(template)

            result = fill_document(template, output, {"CA_DEVELOPED": True})
            self.assertEqual([], result["removed_sections"])
            self.assertIn(
                "Cost content",
                "\n".join(p.text for p in Document(output).paragraphs),
            )

    def test_malformed_json_and_workbook_fail_closed(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            assignment, templates, config = build_assignment(
                temp_dir,
                "[[VALUE]]",
                {"VALUE": "safe"},
            )
            json_path = assignment / "TEST-999_variables.json"
            json_path.write_text("{broken", encoding="utf-8")
            json_result = validate_assignment(assignment, templates, config)
            self.assertFalse(json_result["ready"])
            self.assertTrue(
                any("validation failed" in error.lower() for error in json_result["errors"])
            )

            json_path.write_text('{"VALUE": "safe"}', encoding="utf-8")
            (assignment / "workbook.xlsx").write_bytes(b"not an xlsx")
            workbook_result = validate_assignment(assignment, templates, config)
            self.assertFalse(workbook_result["ready"])
            self.assertTrue(
                any(
                    "validation failed" in error.lower()
                    for error in workbook_result["errors"]
                )
            )

    def test_draft_with_malformed_json_fails_cleanly(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            assignment, _, _ = build_assignment(
                temp_dir,
                "[[VALUE]]",
                {"VALUE": "safe"},
            )
            (assignment / "outputs").mkdir()
            state_path = assignment / ".axiom.json"
            state_path.write_text(
                json.dumps(
                    {
                        "file_no": "TEST-999",
                        "stage": "new",
                        "delivered": None,
                    }
                ),
                encoding="utf-8",
            )
            (assignment / "TEST-999_variables.json").write_text(
                "{broken",
                encoding="utf-8",
            )

            with (
                patch.object(axiom, "_find_assignment", return_value=assignment),
                patch.object(
                    axiom,
                    "check_delivery_readiness",
                    return_value={
                        "ready": False,
                        "errors": ["Malformed JSON"],
                        "missing": [],
                        "unresolved_blocks": {},
                    },
                ),
            ):
                result = axiom.cmd_deliver(["TEST-999", "--draft"])

            self.assertFalse(result)
            final_state = json.loads(state_path.read_text(encoding="utf-8"))
            self.assertEqual("new", final_state["stage"])
            self.assertIsNone(final_state["delivered"])
            self.assertEqual("input_failed", final_state["last_delivery_status"])
            self.assertEqual([], list((assignment / "outputs").iterdir()))

    def test_missing_engagement_template_cannot_mark_assignment_engaged(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            assignment, templates, _ = build_assignment(
                temp_dir,
                "[[VALUE]]",
                {"VALUE": "safe"},
            )
            (assignment / "outputs").mkdir()
            state_path = assignment / ".axiom.json"
            state_path.write_text(
                json.dumps(
                    {
                        "file_no": "TEST-999",
                        "stage": "new",
                        "engaged": None,
                    }
                ),
                encoding="utf-8",
            )
            stage = {
                "engage": {
                    "documents": [
                        {
                            "template": "missing.docx",
                            "output": "engagement.docx",
                        }
                    ]
                }
            }
            with (
                patch.object(axiom, "_find_assignment", return_value=assignment),
                patch.object(axiom, "STAGES", stage),
                patch.object(axiom, "TEMPLATES_DIR", templates),
            ):
                result = axiom.cmd_engage(["TEST-999"])

            self.assertFalse(result)
            final_state = json.loads(state_path.read_text(encoding="utf-8"))
            self.assertEqual("new", final_state["stage"])
            self.assertIsNone(final_state["engaged"])
            self.assertEqual(
                "generation_failed",
                final_state["last_engagement_status"],
            )

    def test_engagement_failure_preserves_prior_document(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            assignment, templates, _ = build_assignment(
                temp_dir,
                "[[VALUE]]",
                {"VALUE": "safe"},
            )
            outputs = assignment / "outputs"
            outputs.mkdir()
            output_path = outputs / "TEST-999_engagement.docx"
            output_path.write_bytes(b"prior-engagement")
            state_path = assignment / ".axiom.json"
            state_path.write_text(
                json.dumps(
                    {
                        "file_no": "TEST-999",
                        "stage": "new",
                        "engaged": None,
                    }
                ),
                encoding="utf-8",
            )

            def fail_after_writing(_template, output, _variables, **_kwargs):
                Path(output).write_bytes(b"partial-engagement")
                raise RuntimeError("synthetic engagement failure")

            stage = {
                "engage": {
                    "documents": [
                        {
                            "template": "report.docx",
                            "output": "engagement.docx",
                        }
                    ]
                }
            }
            with (
                patch.object(axiom, "_find_assignment", return_value=assignment),
                patch.object(axiom, "STAGES", stage),
                patch.object(axiom, "TEMPLATES_DIR", templates),
                patch.object(axiom, "fill_document", side_effect=fail_after_writing),
            ):
                result = axiom.cmd_engage(["TEST-999"])

            self.assertFalse(result)
            self.assertEqual(b"prior-engagement", output_path.read_bytes())
            final_state = json.loads(state_path.read_text(encoding="utf-8"))
            self.assertEqual("new", final_state["stage"])
            self.assertEqual(
                "generation_failed",
                final_state["last_engagement_status"],
            )
            self.assertEqual([], list(outputs.glob(".*.tmp")))

    def test_corrupt_image_is_rejected_before_injection(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            assignment, templates, config = build_assignment(
                temp_dir,
                "[[REGIONAL_MAP_IMAGE]]",
                {},
            )
            maps = assignment / "assets" / "maps"
            maps.mkdir(parents=True)
            (maps / "regional.png").write_bytes(b"this is not a png")

            result = validate_assignment(assignment, templates, config)
            self.assertFalse(result["ready"])
            reason = result["unresolved_blocks"]["REGIONAL_MAP_IMAGE"]
            self.assertIn("unreadable", reason)

    def test_oversized_image_is_rejected_before_injection(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            assignment, templates, config = build_assignment(
                temp_dir,
                "[[REGIONAL_MAP_IMAGE]]",
                {},
            )
            maps = assignment / "assets" / "maps"
            maps.mkdir(parents=True)
            image_path = maps / "regional.png"
            with open(image_path, "wb") as image_file:
                image_file.write(TINY_PNG)
                image_file.seek(MAX_IMAGE_BYTES)
                image_file.write(b"\0")

            result = validate_assignment(assignment, templates, config)
            self.assertFalse(result["ready"])
            reason = result["unresolved_blocks"]["REGIONAL_MAP_IMAGE"]
            self.assertIn("25 MB", reason)

    def test_duplicate_or_incomplete_comps_fail_quality_checks(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            assignment, templates, config = build_assignment(
                temp_dir,
                "[[COMP_SHEETS_BLOCK]]",
                {},
            )
            workbook_path = assignment / "workbook.xlsx"
            workbook = openpyxl.load_workbook(workbook_path)
            sheet = workbook.create_sheet("comp_data")
            sheet.append(["COMP_NO", "", "ADDRESS", "", "", "PRICE"])
            sheet.append(["Sale No. 1", "", "1 Fictional Way", "", "", ""])
            sheet.append(["Sale No. 1", "", "", "", "", "$1,000"])
            workbook.save(workbook_path)
            workbook.close()

            result = validate_assignment(assignment, templates, config)
            self.assertFalse(result["ready"])
            reason = result["unresolved_blocks"]["COMP_SHEETS_BLOCK"]
            self.assertIn("duplicate comparable number", reason)
            self.assertIn("missing sale price", reason)
            self.assertIn("missing address", reason)

    def test_fifty_photos_inject_without_losing_or_reordering_markers(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            assignment = Path(temp_dir) / "assignment"
            photos = assignment / "assets" / "photos" / "subject"
            photos.mkdir(parents=True)
            for index in range(50):
                (photos / f"{index:02d}.png").write_bytes(TINY_PNG)

            output = Path(temp_dir) / "photos.docx"
            document = Document()
            document.add_paragraph("Before [[SUBJECT_PHOTOS_BLOCK]] After")
            document.save(output)

            self.assertEqual(
                {"SUBJECT_PHOTOS_BLOCK": 50},
                inject_media_blocks(output, assignment),
            )
            rendered = Document(output)
            self.assertEqual(50, len(rendered.inline_shapes))
            self.assertEqual([], find_docx_placeholders(output))
            self.assertIn("Before", rendered.paragraphs[0].text)
            self.assertIn("After", rendered.paragraphs[0].text)

    def test_fifty_comps_inject_with_page_breaks_and_no_placeholders(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            workbook_path = root / "comps.xlsx"
            workbook = openpyxl.Workbook()
            sheet = workbook.active
            sheet.title = "comp_data"
            sheet.append(["COMP_NO", "", "ADDRESS", "", "", "PRICE"])
            for index in range(1, 51):
                sheet.append(
                    [
                        f"Sale No. {index}",
                        "",
                        f"{index} Fictional Stress Test Way",
                        "",
                        "",
                        f"${index * 100000:,}",
                    ]
                )
            workbook.save(workbook_path)
            workbook.close()

            template = root / "comp-template.docx"
            template_doc = Document()
            template_doc.add_paragraph(
                "[[COMP_NO]] | [[COMP_ADDRESS_LINE1]] | [[COMP_SALE_PRICE]]"
            )
            template_doc.save(template)

            report = root / "report.docx"
            report_doc = Document()
            report_doc.add_paragraph("[[COMP_SHEETS_BLOCK]]")
            report_doc.save(report)

            self.assertEqual(
                50,
                inject_comp_section(report, template, workbook_path),
            )
            self.assertEqual([], find_docx_placeholders(report))
            generated = Document(report)
            page_breaks = [
                node
                for node in generated.element.body.iter(qn("w:br"))
                if node.get(qn("w:type")) == "page"
            ]
            self.assertEqual(49, len(page_breaks))

    def test_split_run_comp_template_placeholder_is_filled(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            workbook_path = root / "comps.xlsx"
            workbook = openpyxl.Workbook()
            sheet = workbook.active
            sheet.title = "comp_data"
            sheet.append(["COMP_NO", "", "ADDRESS", "", "", "PRICE"])
            sheet.append(["Sale No. 1", "", "1 Fictional Way", "", "", "$1"])
            workbook.save(workbook_path)
            workbook.close()

            template = root / "comp-template.docx"
            template_doc = Document()
            paragraph = template_doc.add_paragraph()
            paragraph.add_run("[[COMP_")
            paragraph.add_run("NO]]")
            template_doc.save(template)

            report = root / "report.docx"
            report_doc = Document()
            report_doc.add_paragraph("[[COMP_SHEETS_BLOCK]]")
            report_doc.save(report)

            inject_comp_section(report, template, workbook_path)
            self.assertEqual([], find_docx_placeholders(report))
            self.assertIn("Sale No. 1", Document(report).paragraphs[0].text)

    def test_unicode_and_large_text_round_trip(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            template = Path(temp_dir) / "template.docx"
            output = Path(temp_dir) / "output.docx"
            document = Document()
            document.add_paragraph("[[LONG_VALUE]]")
            document.save(template)
            value = "Åxiom — 東京 — 🏢 " + ("Fictional data. " * 4000)

            result = fill_document(template, output, {"LONG_VALUE": value})
            self.assertEqual([], result["missing"])
            self.assertEqual(value, Document(output).paragraphs[0].text)

    def test_generation_failure_preserves_existing_output_and_state(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            assignment, templates, _ = build_assignment(
                root,
                "[[VALUE]]",
                {"VALUE": "safe"},
            )
            outputs = assignment / "outputs"
            outputs.mkdir()
            output_path = outputs / "TEST-999_Appraisal.docx"
            output_path.write_bytes(b"known-good-output")
            state_path = assignment / ".axiom.json"
            state_path.write_text(
                json.dumps(
                    {
                        "file_no": "TEST-999",
                        "stage": "new",
                        "delivered": None,
                    }
                ),
                encoding="utf-8",
            )

            def fail_after_writing(_template, output, _variables, **_kwargs):
                Path(output).write_bytes(b"partial-output")
                raise RuntimeError("synthetic generation failure")

            stage = {
                "deliver": {
                    "documents": [
                        {
                            "template": "report.docx",
                            "output": "Appraisal.docx",
                        }
                    ]
                }
            }
            with (
                patch.object(axiom, "_find_assignment", return_value=assignment),
                patch.object(
                    axiom,
                    "check_delivery_readiness",
                    return_value={
                        "ready": True,
                        "errors": [],
                        "missing": [],
                        "unresolved_blocks": {},
                    },
                ),
                patch.object(axiom, "STAGES", stage),
                patch.object(axiom, "TEMPLATES_DIR", templates),
                patch.object(axiom, "fill_document", side_effect=fail_after_writing),
            ):
                result = axiom.cmd_deliver(["TEST-999"])

            self.assertFalse(result)
            self.assertEqual(b"known-good-output", output_path.read_bytes())
            final_state = json.loads(state_path.read_text(encoding="utf-8"))
            self.assertEqual("new", final_state["stage"])
            self.assertIsNone(final_state["delivered"])
            self.assertEqual(
                "generation_failed",
                final_state["last_delivery_status"],
            )
            self.assertEqual([], list(outputs.glob(".*.tmp")))

    def test_locked_output_preserves_prior_report_and_records_failure(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            assignment, templates, _ = build_assignment(
                root,
                "[[VALUE]]",
                {"VALUE": "safe"},
            )
            outputs = assignment / "outputs"
            outputs.mkdir()
            output_path = outputs / "TEST-999_Appraisal.docx"
            output_path.write_bytes(b"prior-reviewed-report")
            state_path = assignment / ".axiom.json"
            state_path.write_text(
                json.dumps(
                    {
                        "file_no": "TEST-999",
                        "stage": "new",
                        "delivered": None,
                    }
                ),
                encoding="utf-8",
            )
            stage = {
                "deliver": {
                    "documents": [
                        {
                            "template": "report.docx",
                            "output": "Appraisal.docx",
                        }
                    ]
                }
            }
            with (
                patch.object(axiom, "_find_assignment", return_value=assignment),
                patch.object(
                    axiom,
                    "check_delivery_readiness",
                    return_value={
                        "ready": True,
                        "errors": [],
                        "missing": [],
                        "unresolved_blocks": {},
                    },
                ),
                patch.object(axiom, "STAGES", stage),
                patch.object(axiom, "TEMPLATES_DIR", templates),
                patch.object(
                    Path,
                    "replace",
                    side_effect=PermissionError("synthetic locked output"),
                ),
            ):
                result = axiom.cmd_deliver(["TEST-999"])

            self.assertFalse(result)
            self.assertEqual(b"prior-reviewed-report", output_path.read_bytes())
            final_state = json.loads(state_path.read_text(encoding="utf-8"))
            self.assertEqual("generation_failed", final_state["last_delivery_status"])
            self.assertIn("locked output", final_state["last_delivery_error"])
            self.assertEqual([], list(outputs.glob(".*.tmp")))

    def test_dilmore_uses_correct_ratio_direction_and_signature(self):
        """cmd_dilmore used to call dilmore_factor/dilmore_adj_pct as
        (subject_gba, comp_gba, curve) -- 3 positional args against their
        real 2-arg (ratio, curve) signature -- which raised TypeError on
        every real run, and separately had the ratio backwards (subject/comp
        instead of comp/subject). This exercises a real subject/comp pair
        where the two directions produce materially different, verifiable
        numbers: a comp twice the subject's size should get a positive size
        adjustment, and a comp half the subject's size should get a negative
        one, each matching dilmore_factor/dilmore_adj_pct computed directly
        with the correct ratio."""
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            assignments = root / "assignments"
            assignments.mkdir()
            assignment = assignments / "TEST-500_Fictional_Client"
            assignment.mkdir()

            workbook = openpyxl.Workbook()
            intake = workbook.active
            intake.title = "Intake"
            intake.append(["Field", "Value"])
            intake.append(["GBA", 10000])

            # Mirrors the real templates/workbook.xlsx size_adj header row
            # exactly (A=Comp, B=Comp GBA, C=Ratio formula, D=Size Factor,
            # E=Adj %, F=Adj $/SF, G=Notes) -- including a real Ratio formula
            # in column C, so this test actually catches a regression that
            # writes Size Factor/Adj % into the wrong columns and clobbers
            # it (as a prior version of this fix did: it wrote to columns
            # 3/4 -- C/D -- instead of the real 4/5 -- D/E).
            size_adj = workbook.create_sheet("size_adj")
            size_adj.cell(row=6, column=1).value = "Comp"
            size_adj.cell(row=6, column=2).value = "Comp GBA (SF)"
            size_adj.cell(row=6, column=3).value = "Ratio (Ac/As)"
            size_adj.cell(row=6, column=4).value = "Size Factor"
            size_adj.cell(row=6, column=5).value = "Adj %"
            size_adj.cell(row=6, column=6).value = "Adj $ / SF"
            size_adj.cell(row=6, column=7).value = "Notes"
            size_adj["B3"] = 85
            size_adj.cell(row=7, column=2).value = 20000  # Comp 1: 2x subject
            size_adj.cell(row=7, column=3).value = '=IF(B7="","",IFERROR(B7/$B$4,""))'
            size_adj.cell(row=8, column=2).value = 5000   # Comp 2: half subject
            size_adj.cell(row=8, column=3).value = '=IF(B8="","",IFERROR(B8/$B$4,""))'

            workbook_path = assignment / "workbook.xlsx"
            workbook.save(workbook_path)
            workbook.close()

            with patch.object(axiom, "ASSIGNMENTS_DIR", assignments):
                axiom.cmd_dilmore(["TEST-500"])

            result_wb = openpyxl.load_workbook(workbook_path)
            result_sa = result_wb["size_adj"]

            expected_factor_1 = round(dilmore_factor(20000 / 10000, 85), 4)
            # size_adj!E6 is a true Excel percentage format (expects a
            # fraction, e.g. 0.18 -> "+18.0%"), not dilmore_adj_pct()'s
            # percentage-point number (18.0, which would display as
            # "1800.0%") -- so the written value must be factor - 1, not
            # dilmore_adj_pct() directly. See axiom.py's cmd_dilmore.
            expected_adj_1 = round(expected_factor_1 - 1, 4)
            expected_factor_2 = round(dilmore_factor(5000 / 10000, 85), 4)
            expected_adj_2 = round(expected_factor_2 - 1, 4)

            # Size Factor -> column D (4), Adj % -> column E (5).
            self.assertEqual(expected_factor_1, result_sa.cell(row=7, column=4).value)
            self.assertEqual(expected_adj_1, result_sa.cell(row=7, column=5).value)
            self.assertEqual(expected_factor_2, result_sa.cell(row=8, column=4).value)
            self.assertEqual(expected_adj_2, result_sa.cell(row=8, column=5).value)

            # Column C's pre-existing Ratio formula must survive untouched --
            # this is the exact cell the earlier buggy column mapping
            # clobbered with the Size Factor value instead.
            self.assertEqual(
                '=IF(B7="","",IFERROR(B7/$B$4,""))',
                result_sa.cell(row=7, column=3).value,
            )
            self.assertEqual(
                '=IF(B8="","",IFERROR(B8/$B$4,""))',
                result_sa.cell(row=8, column=3).value,
            )

            # Comp 1 (larger than subject) gets a positive adjustment; comp 2
            # (smaller than subject) gets a negative one. With the old
            # backwards ratio these signs would flip.
            self.assertGreater(result_sa.cell(row=7, column=5).value, 0)
            self.assertLess(result_sa.cell(row=8, column=5).value, 0)
            result_wb.close()

    def test_dilmore_invalid_curve_fails_loudly_without_writing(self):
        """An out-of-table curve value in size_adj!B3 must produce a clear
        error and leave the workbook byte-for-byte untouched, not a raw
        traceback or a partial write."""
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            assignments = root / "assignments"
            assignments.mkdir()
            assignment = assignments / "TEST-501_Fictional_Client"
            assignment.mkdir()

            workbook = openpyxl.Workbook()
            intake = workbook.active
            intake.title = "Intake"
            intake.append(["Field", "Value"])
            intake.append(["GBA", 10000])

            size_adj = workbook.create_sheet("size_adj")
            size_adj["B3"] = 83  # not one of 80, 82.5, 85, 87.5, 90
            size_adj.cell(row=7, column=2).value = 20000

            workbook_path = assignment / "workbook.xlsx"
            workbook.save(workbook_path)
            workbook.close()
            original_bytes = workbook_path.read_bytes()

            with patch.object(axiom, "ASSIGNMENTS_DIR", assignments):
                axiom.cmd_dilmore(["TEST-501"])

            self.assertEqual(original_bytes, workbook_path.read_bytes())

    def test_dilmore_prefers_sca_adjustment_grid_over_size_adj(self):
        """Phase 6's sca_adjustment_grid tab uses different column
        positions than the older size_adj tab (Comp GBA in C not B,
        Size Factor/Adj % in K/L not D/E). _run_dilmore_calc must detect
        which tab a given workbook actually has and use the matching
        layout -- not assume the old size_adj positions unconditionally."""
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            assignments = root / "assignments"
            assignments.mkdir()
            assignment = assignments / "TEST-502_Fictional_Client"
            assignment.mkdir()

            workbook = openpyxl.Workbook()
            intake = workbook.active
            intake.title = "Intake"
            intake.append(["Field", "Value"])
            intake.append(["GBA", 10000])

            grid = workbook.create_sheet("sca_adjustment_grid")
            grid["B3"] = 85
            grid.cell(row=7, column=1).value = "Sale No. 1"
            # New layout: Comp GBA in column C (3), not B (2).
            grid.cell(row=7, column=3).value = 20000  # 2x subject -> positive

            workbook_path = assignment / "workbook.xlsx"
            workbook.save(workbook_path)
            workbook.close()

            with patch.object(axiom, "ASSIGNMENTS_DIR", assignments):
                axiom.cmd_dilmore(["TEST-502"])

            result_wb = openpyxl.load_workbook(workbook_path)
            grid_result = result_wb["sca_adjustment_grid"]

            expected_factor = round(dilmore_factor(20000 / 10000, 85), 4)
            expected_adj = round(expected_factor - 1, 4)
            # New layout: Size Factor in column K (11), Adj % in L (12).
            self.assertEqual(expected_factor, grid_result.cell(row=7, column=11).value)
            self.assertEqual(expected_adj, grid_result.cell(row=7, column=12).value)
            # Old size_adj columns D/E must be untouched (there is no
            # size_adj tab at all in this fixture -- if the old column
            # positions were used against this new layout by mistake, this
            # would have written into what's actually the Sale Date/Months
            # columns instead).
            self.assertIsNone(grid_result.cell(row=7, column=4).value)
            self.assertIsNone(grid_result.cell(row=7, column=5).value)
            result_wb.close()

    def test_dilmore_writes_comps_hand_expanded_past_row_16(self):
        """sca_adjustment_grid's fixed capacity is rows 7-16 ("Sale No.
        1".."Sale No. 10"), but adjustment_grid.py's own MAX_DATA_ROW
        design lets Derek hand-expand a sheet with more comps below that,
        anchored the same way ("Sale No. 11", etc.). _run_dilmore_calc
        used to scan only the fixed rows 7-16, so a hand-expanded comp's
        GBA would silently never get a Size Factor/Adj % written at all
        (Fable adversarial review finding N1) -- this exercises comp 11
        at row 17."""
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            assignments = root / "assignments"
            assignments.mkdir()
            assignment = assignments / "TEST-503_Fictional_Client"
            assignment.mkdir()

            workbook = openpyxl.Workbook()
            intake = workbook.active
            intake.title = "Intake"
            intake.append(["Field", "Value"])
            intake.append(["GBA", 10000])

            grid = workbook.create_sheet("sca_adjustment_grid")
            grid["B3"] = 85
            # The real template pre-labels every one of its built-in rows
            # 7-16 with "Sale No. 1".."Sale No. 10" regardless of whether
            # a comp's data is filled in yet -- _run_dilmore_calc's scan
            # (like read_grid_rows') stops at the first row whose anchor
            # doesn't match, so a gap in the labels would hide every comp
            # past the gap, including the intentionally-tested row 17.
            for offset in range(10):
                grid.cell(row=7 + offset, column=1).value = f"Sale No. {offset + 1}"
            grid.cell(row=7, column=3).value = 20000
            # Hand-expanded 11th comp, anchored the same way real comp
            # rows are.
            grid.cell(row=17, column=1).value = "Sale No. 11"
            grid.cell(row=17, column=3).value = 15000

            workbook_path = assignment / "workbook.xlsx"
            workbook.save(workbook_path)
            workbook.close()

            with patch.object(axiom, "ASSIGNMENTS_DIR", assignments):
                axiom.cmd_dilmore(["TEST-503"])

            result_wb = openpyxl.load_workbook(workbook_path)
            grid_result = result_wb["sca_adjustment_grid"]

            expected_factor_11 = round(dilmore_factor(15000 / 10000, 85), 4)
            self.assertEqual(
                expected_factor_11, grid_result.cell(row=17, column=11).value
            )
            self.assertIsNotNone(grid_result.cell(row=17, column=12).value)
            result_wb.close()

    def test_dilmore_staleness_warning_is_read_only(self):
        """validate's staleness check must never write to the workbook --
        only cmd_dilmore/deliver's auto-run may write. A comp GBA with no
        Size Factor yet should produce exactly one warning naming the tab
        and comp number, and the workbook must be byte-for-byte unchanged
        afterward."""
        with tempfile.TemporaryDirectory() as temp_dir:
            assignment = Path(temp_dir)
            workbook = openpyxl.Workbook()
            intake = workbook.active
            intake.title = "Intake"
            intake.append(["Field", "Value"])
            intake.append(["GBA", 10000])

            grid = workbook.create_sheet("sca_adjustment_grid")
            grid.cell(row=7, column=1).value = "Sale No. 1"
            grid.cell(row=7, column=3).value = 20000   # GBA entered
            # Size Factor (column 11) intentionally left blank -- stale.
            grid.cell(row=8, column=1).value = "Sale No. 2"
            grid.cell(row=8, column=3).value = 5000    # GBA entered
            grid.cell(row=8, column=11).value = 0.85   # already computed

            workbook_path = assignment / "workbook.xlsx"
            workbook.save(workbook_path)
            workbook.close()
            original_bytes = workbook_path.read_bytes()

            warnings = axiom._dilmore_staleness_warnings(workbook_path)

            self.assertEqual(original_bytes, workbook_path.read_bytes())
            self.assertEqual(1, len(warnings))
            self.assertIn("sca_adjustment_grid", warnings[0])
            self.assertIn("Comp 1", warnings[0])
            self.assertNotIn("Comp 2", warnings[0])

    def test_deliver_auto_runs_dilmore_but_hard_stops_when_it_writes(self):
        """Per Derek's explicit choice (2026-07-10), Dilmore stays a tested
        Python calculation rather than a live Excel formula, and it must
        run automatically as part of deliver so Derek never has to
        remember a separate `dilmore` step after editing a comp's GBA.

        But openpyxl has no formula engine: any save() that actually writes
        a changed Size Factor/Adj % also silently wipes every OTHER cached
        formula result in the workbook (Net Adjustment, Indicated Value,
        the outputs tab -- see _run_dilmore_calc's docstring, and the
        Fable adversarial review that caught this as finding A1). So a
        real write must compute the right factor AND hard-stop delivery
        with a clear message, rather than silently proceeding to build a
        report from a workbook whose cache it just invalidated."""
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            assignment, templates, _ = build_assignment(
                root,
                "[[VALUE]]",
                {"VALUE": "safe"},
            )

            workbook_path = assignment / "workbook.xlsx"
            workbook = openpyxl.load_workbook(workbook_path)
            intake = workbook.create_sheet("Intake")
            intake.append(["Field", "Value"])
            intake.append(["GBA", 10000])
            grid = workbook.create_sheet("sca_adjustment_grid")
            grid["B3"] = 85
            grid.cell(row=7, column=1).value = "Sale No. 1"
            grid.cell(row=7, column=3).value = 20000
            workbook.save(workbook_path)
            workbook.close()

            stage = {
                "deliver": {
                    "documents": [
                        {"template": "report.docx", "output": "Appraisal.docx"}
                    ]
                }
            }

            def fake_fill_document(_template, output, _variables, **_kwargs):
                Document().save(output)
                return {"missing": [], "blocks": [], "removed_sections": []}

            with (
                patch.object(axiom, "_find_assignment", return_value=assignment),
                patch.object(
                    axiom,
                    "check_delivery_readiness",
                    return_value={
                        "ready": True,
                        "errors": [],
                        "missing": [],
                        "unresolved_blocks": {},
                    },
                ),
                patch.object(axiom, "STAGES", stage),
                patch.object(axiom, "TEMPLATES_DIR", templates),
                patch.object(axiom, "fill_document", side_effect=fake_fill_document),
            ):
                # NOT --draft: draft mode now skips the Dilmore auto-run
                # entirely (Fable adversarial review finding P4), so the
                # hard-stop-on-real-write behavior this test exercises only
                # applies to a normal (non-draft) deliver.
                result = axiom.cmd_deliver(["TEST-999"])

            # The write happened and computed the right factor ...
            result_wb = openpyxl.load_workbook(workbook_path)
            grid_result = result_wb["sca_adjustment_grid"]
            expected_factor = round(dilmore_factor(20000 / 10000, 85), 4)
            self.assertEqual(expected_factor, grid_result.cell(row=7, column=11).value)
            result_wb.close()

            # ... but delivery must NOT have completed, since the workbook's
            # other cached formula values are now stale/blank.
            self.assertFalse(result)
            state = axiom._load_state(assignment)
            self.assertEqual("input_failed", state["last_delivery_status"])
            self.assertIn("recalculate", state["last_delivery_error"].lower())
            self.assertNotEqual("delivered", state.get("stage"))

    def test_deliver_proceeds_when_dilmore_values_already_current(self):
        """If the adjustment grid's Size Factor/Adj % are already correct
        (the common case -- nothing about the comps changed since the last
        deliver), _run_dilmore_calc must skip the write entirely and
        deliver must proceed normally, per the "changed": False path added
        alongside the hard-stop above."""
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            assignment, templates, _ = build_assignment(
                root,
                "[[VALUE]]",
                {"VALUE": "safe"},
            )

            workbook_path = assignment / "workbook.xlsx"
            workbook = openpyxl.load_workbook(workbook_path)
            intake = workbook.create_sheet("Intake")
            intake.append(["Field", "Value"])
            intake.append(["GBA", 10000])
            grid = workbook.create_sheet("sca_adjustment_grid")
            grid["B3"] = 85
            grid.cell(row=7, column=1).value = "Sale No. 1"
            grid.cell(row=7, column=3).value = 20000
            expected_factor = round(dilmore_factor(20000 / 10000, 85), 4)
            expected_adj_pct = round(expected_factor - 1, 4)
            grid.cell(row=7, column=11).value = expected_factor
            grid.cell(row=7, column=12).value = expected_adj_pct
            workbook.save(workbook_path)
            workbook.close()

            stage = {
                "deliver": {
                    "documents": [
                        {"template": "report.docx", "output": "Appraisal.docx"}
                    ]
                }
            }

            def fake_fill_document(_template, output, _variables, **_kwargs):
                Document().save(output)
                return {"missing": [], "blocks": [], "removed_sections": []}

            with (
                patch.object(axiom, "_find_assignment", return_value=assignment),
                patch.object(
                    axiom,
                    "check_delivery_readiness",
                    return_value={
                        "ready": True,
                        "errors": [],
                        "missing": [],
                        "unresolved_blocks": {},
                    },
                ),
                patch.object(axiom, "STAGES", stage),
                patch.object(axiom, "TEMPLATES_DIR", templates),
                patch.object(axiom, "fill_document", side_effect=fake_fill_document),
            ):
                # NOT --draft: exercises _run_dilmore_calc's "changed": False
                # skip path for real, rather than draft mode's unconditional
                # Dilmore skip (Fable adversarial review finding P4).
                result = axiom.cmd_deliver(["TEST-999"])

            self.assertTrue(result)

    def test_draft_mode_never_runs_dilmore_or_mutates_anything(self):
        """Fable adversarial review finding P4: draft mode's banner promises
        "Assignment delivery state will not be changed", but the Dilmore
        auto-run used to execute unconditionally even in draft mode -- a
        real write there both mutates workbook.xlsx (wipes every other
        cached formula result) and sets .axiom.json's formula_cache_stale
        flag, either of which breaks that promise. Draft mode must now skip
        the Dilmore auto-run entirely, leaving the workbook and state
        byte-for-byte/key-for-key untouched, even when Dilmore would
        otherwise have real, uncomputed size adjustments to write."""
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            assignment, templates, _ = build_assignment(
                root,
                "[[VALUE]]",
                {"VALUE": "safe"},
            )

            workbook_path = assignment / "workbook.xlsx"
            workbook = openpyxl.load_workbook(workbook_path)
            intake = workbook.create_sheet("Intake")
            intake.append(["Field", "Value"])
            intake.append(["GBA", 10000])
            grid = workbook.create_sheet("sca_adjustment_grid")
            grid["B3"] = 85
            grid.cell(row=7, column=1).value = "Sale No. 1"
            # A real, uncomputed size adjustment: Dilmore would write here
            # if this were a normal (non-draft) deliver.
            grid.cell(row=7, column=3).value = 20000
            workbook.save(workbook_path)
            workbook.close()
            original_bytes = workbook_path.read_bytes()

            stage = {
                "deliver": {
                    "documents": [
                        {"template": "report.docx", "output": "Appraisal.docx"}
                    ]
                }
            }

            def fake_fill_document(_template, output, _variables, **_kwargs):
                Document().save(output)
                return {"missing": [], "blocks": [], "removed_sections": []}

            with (
                patch.object(axiom, "_find_assignment", return_value=assignment),
                patch.object(
                    axiom,
                    "check_delivery_readiness",
                    return_value={
                        "ready": True,
                        "errors": [],
                        "missing": [],
                        "unresolved_blocks": {},
                    },
                ),
                patch.object(axiom, "STAGES", stage),
                patch.object(axiom, "TEMPLATES_DIR", templates),
                patch.object(axiom, "fill_document", side_effect=fake_fill_document),
            ):
                result = axiom.cmd_deliver(["TEST-999", "--draft"])

            self.assertTrue(result)
            self.assertEqual(original_bytes, workbook_path.read_bytes())
            state = axiom._load_state(assignment)
            self.assertNotIn("formula_cache_stale", state)
            self.assertNotIn("formula_cache_stale_mtime", state)
            self.assertEqual("draft_generated", state["last_delivery_status"])


if __name__ == "__main__":
    unittest.main()
