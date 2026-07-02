import json
import tempfile
import unittest
from pathlib import Path

from docx import Document

from comparable_contract import confirm_extraction_result
from db import get_conn, init_db, search_market_observations
from harvest_contract import OBSERVATION_CONTRACT_ID
from ingest import commit_extraction_result, run_extraction


def _build_market_report(path):
    document = Document()
    document.add_paragraph("Subject Property: 456 Fictional Market Road")
    document.add_paragraph("Effective Date: May 31, 2025")
    document.add_paragraph("Reconciled Value: $3,000,000")
    document.add_heading("Market Area Analysis", level=1)
    document.add_paragraph(
        "The fictional market area contains a diverse employment base and "
        "stable transportation access. This paragraph is synthetic test "
        "content and is not evidence about any real place."
    )
    document.add_paragraph(
        "Fictional population and employment indicators were generally stable "
        "during the test period, with no real-world inference intended."
    )
    document.add_heading("Supply and Demand", level=1)
    document.add_paragraph(
        "The fictional competitive inventory remained balanced during the "
        "synthetic observation period. Leasing velocity and vacancy references "
        "are invented solely to exercise bounded section harvesting."
    )
    document.add_heading("Neighborhood Analysis", level=1)
    document.add_paragraph("Too short to retain.")
    document.add_heading("Reconciliation", level=1)
    document.add_paragraph(
        "This valuation discussion must not leak into the preceding market "
        "observation."
    )
    document.save(path)


class ObservationHarvestTests(unittest.TestCase):
    def test_fictional_observations_extract_review_commit_and_search(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            assignment = (
                root
                / "historical"
                / "25C950 - Office - 456 Fictional Market Road, Demo City"
            )
            assignment.mkdir(parents=True)
            _build_market_report(assignment / "REPORT 25C950.docx")

            staged_path = run_extraction(
                root / "historical",
                staged_dir=root / "staged",
            )[0]
            staged = json.loads(staged_path.read_text(encoding="utf-8"))
            observations = staged["market_observations"]
            self.assertEqual(2, len(observations))
            self.assertEqual(
                ["market_area", "supply_demand"],
                [record["data"]["category"] for record in observations],
            )
            first = observations[0]
            self.assertEqual(OBSERVATION_CONTRACT_ID, first["contract_id"])
            self.assertEqual("2025-05-31", first["data"]["effective_date"])
            self.assertEqual("Demo City", first["data"]["geography"])
            self.assertEqual("Office", first["data"]["property_type"])
            self.assertEqual(
                "paragraphs:4-6",
                first["provenance"]["source_locator"],
            )
            self.assertNotIn("valuation discussion", first["data"]["text"])
            self.assertEqual("unreviewed", first["review"]["status"])

            first["data"]["text"] += "\n\nReviewed fictional clarification."
            first["review_edits"] = [{
                "field": "text",
                "before": "(fictional extracted text)",
                "after": "(fictional reviewed text)",
            }]
            confirmed = confirm_extraction_result(
                staged,
                reviewer="fictional-observation-qa",
                reviewed_at="2026-07-01T17:00:00+00:00",
            )

            db_path = root / "observations.db"
            init_db(db_path, quiet=True)
            connection = get_conn(db_path)
            try:
                with connection:
                    counts = commit_extraction_result(confirmed, connection)
                self.assertEqual(2, counts["market_observations"])
                with connection:
                    duplicates = commit_extraction_result(confirmed, connection)
                self.assertEqual(
                    2,
                    duplicates["duplicate_market_observations"],
                )
            finally:
                connection.close()

            records = search_market_observations(
                db_path,
                category="market_area",
                geography="Demo",
                property_type="Office",
                text_contains="clarification",
                effective_date_from="2025-01-01",
                effective_date_to="2025-12-31",
            )
            self.assertEqual(1, len(records))
            self.assertEqual("25C950", records[0]["file_no"])
            self.assertEqual("confirmed", records[0]["review_status"])
            self.assertIn(
                "Reviewed fictional clarification",
                records[0]["observation_text"],
            )

    def test_unconfirmed_observation_rolls_back_assignment_and_source(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            assignment = (
                root
                / "historical"
                / "25C951 - Office - 654 Fictional Market Road, Demo City"
            )
            assignment.mkdir(parents=True)
            _build_market_report(assignment / "REPORT 25C951.docx")
            staged_path = run_extraction(
                root / "historical",
                staged_dir=root / "staged",
            )[0]
            staged = json.loads(staged_path.read_text(encoding="utf-8"))
            confirmed = confirm_extraction_result(staged, reviewer="qa")
            confirmed["market_observations"][0]["review"]["status"] = "unreviewed"

            db_path = root / "rollback.db"
            init_db(db_path, quiet=True)
            connection = get_conn(db_path)
            try:
                with self.assertRaisesRegex(ValueError, "not been confirmed"):
                    with connection:
                        commit_extraction_result(confirmed, connection)
                for table in (
                    "assignments",
                    "market_observations",
                    "source_documents",
                ):
                    self.assertEqual(
                        0,
                        connection.execute(
                            f"SELECT count(*) FROM {table}"
                        ).fetchone()[0],
                    )
            finally:
                connection.close()

    def test_contract_descriptor_and_database_table(self):
        schema_path = (
            Path(__file__).resolve().parents[1]
            / "schemas"
            / "historical_harvest.v1.json"
        )
        schema = json.loads(schema_path.read_text(encoding="utf-8"))
        self.assertEqual(
            OBSERVATION_CONTRACT_ID,
            schema["contracts"]["observation"],
        )
        with tempfile.TemporaryDirectory() as temp_dir:
            db_path = Path(temp_dir) / "fresh.db"
            init_db(db_path, quiet=True)
            connection = get_conn(db_path)
            try:
                columns = {
                    row[1]
                    for row in connection.execute(
                        "PRAGMA table_info(market_observations)"
                    )
                }
                self.assertIn("observation_text", columns)
                self.assertIn("source_record_json", columns)
            finally:
                connection.close()


if __name__ == "__main__":
    unittest.main()
