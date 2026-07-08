import json
import tempfile
import unittest
from datetime import date
from pathlib import Path

import openpyxl
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.pagesizes import landscape, letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from comparable_contract import confirm_extraction_result
from db import (
    get_conn,
    init_db,
    search_operating_expenses,
    search_rent_roll_entries,
)
from harvest_contract import (
    EXPENSE_CONTRACT_ID,
    RENT_ROLL_CONTRACT_ID,
)
from ingest import commit_extraction_result, run_extraction
from financial_extractor import (
    RENT_ROLL_SYNONYMS,
    _find_header,
    extract_financial_workbook,
)
from pdf_financial_extractor import _ocr_available, extract_financial_pdf

# PyMuPDF/numpy/Pillow are only needed to *build* the synthetic scanned-PDF
# fixtures below; unlike every other test dependency in this file, they're
# new as of the OCR lane and not guaranteed to be installed yet (there's no
# repo-tracked requirements.txt). Guard them so a missing package skips only
# the OCR tests instead of crashing this entire module's collection.
try:
    import fitz  # PyMuPDF
    import numpy as np
    from PIL import Image
    _OCR_TEST_DEPS_AVAILABLE = True
except ImportError:
    _OCR_TEST_DEPS_AVAILABLE = False

# Whether the Tesseract OCR *binary* (not just the pytesseract pip wrapper)
# is actually installed and reachable. `_ocr_available()` is safe to call
# even when PyMuPDF/pytesseract failed to import -- it's guarded internally
# in pdf_financial_extractor.py and just returns False.
_TESSERACT_AVAILABLE = _ocr_available()


def _build_report(path):
    document = Document()
    document.add_paragraph("Subject Property: 777 Fictional Finance Road")
    document.add_paragraph("Effective Date: June 30, 2025")
    document.add_paragraph("Reconciled Value: $2,500,000")
    document.save(path)


def _build_rent_roll(path):
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Rent Roll"
    sheet["A1"] = "Fictional Rent Roll As of June 30, 2025"
    sheet.append([])
    sheet.append([
        "Suite",
        "Tenant",
        "Use",
        "Rentable SF",
        "Lease Start",
        "Expiration",
        "Monthly Rent",
        "Annual Rent",
        "Rent/SF",
        "Rent Structure",
        "Status",
    ])
    first = [
        "101",
        "Fictional Tenant Alpha",
        "Office",
        1_000,
        date(2024, 1, 1),
        date(2028, 12, 31),
        2_000,
        24_000,
        24,
        "Modified Gross",
        "Occupied",
    ]
    sheet.append(first)
    sheet.append([
        "102",
        "Fictional Tenant Beta",
        "Office",
        1_500,
        date(2025, 3, 1),
        date(2030, 2, 28),
        3_250,
        39_000,
        26,
        "NNN",
        "Occupied",
    ])
    sheet.append(first)  # exact duplicate should collapse before review
    sheet.append(["Total", None, None, 2_500, None, None, 5_250])
    workbook.save(path)
    workbook.close()


def _build_expenses(path):
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Operating Expenses"
    sheet.append([
        "Expense Category",
        "Period Year",
        "Period Type",
        "Amount",
        "Per SF",
        "Notes",
    ])
    first = [
        "Real Estate Taxes",
        2025,
        "Actual",
        25_000,
        10,
        "Fictional QA",
    ]
    sheet.append(first)
    sheet.append([
        "Insurance",
        2025,
        "Actual",
        12_500,
        5,
        "Fictional QA",
    ])
    sheet.append(first)  # exact duplicate should collapse
    sheet.append(["Total Operating Expenses", 2025, "Actual", 37_500, 15])
    workbook.save(path)
    workbook.close()


def _build_wide_expenses(path):
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Operating Statement"
    sheet.append([None, "2023", None, "2024", None, "2025 Budget", "Notes"])
    sheet.append([
        "Expense Category",
        "Actual",
        "Actual $/SF",
        "Actual",
        "Actual $/SF",
        "Budget",
        "Notes",
    ])
    sheet.append([
        "Real Estate Taxes",
        20_000,
        8,
        25_000,
        10,
        30_000,
        "Fictional two-row header",
    ])
    sheet.append([
        "Insurance",
        10_000,
        4,
        12_000,
        4.8,
        15_000,
        "Fictional two-row header",
    ])
    sheet.append([
        "Total Operating Expenses",
        30_000,
        12,
        37_000,
        14.8,
        45_000,
        None,
    ])
    workbook.save(path)
    workbook.close()


def _build_specialty_rent_rolls(path):
    workbook = openpyxl.Workbook()
    mini = workbook.active
    mini.title = "master list"
    mini.append([
        "Unit #",
        "Tenant",
        "Move-In",
        "Monthly Rent",
        "Unit Type",
        "Status",
        "Notes",
    ])
    mini.append([
        "A-101",
        "Fictional Storage Tenant Alpha",
        date(2022, 7, 26),
        125,
        "10x10 Outside",
        "Occupied",
        "Fictional mini-storage master row",
    ])
    mini.append([
        "A-102",
        "Fictional Storage Tenant Beta",
        date(2022, 7, 26),
        95,
        "5x10 Inside",
        "Occupied",
        "Fictional mini-storage master row",
    ])
    mini.append(["Total", None, None, 220])

    category = workbook.create_sheet("10X10 - OUTSIDE")
    category.append([
        "Unit #",
        "Tenant",
        "Move-In",
        "Monthly Rent",
        "Unit Type",
        "Status",
        "Notes",
    ])
    category.append([
        "A-101",
        "Fictional Storage Tenant Alpha",
        date(2022, 7, 26),
        125,
        "10x10 Outside",
        "Occupied",
        "Duplicate category sheet row",
    ])

    mobile_home = workbook.create_sheet("JANUARY 2026 MH")
    mobile_home.append([
        "Lot #",
        "Resident",
        "Move In",
        "Rent",
        "Discounts",
        "Status",
        "Notes",
    ])
    mobile_home.append([
        "L-12",
        "Fictional Mobile Resident",
        date(2025, 1, 15),
        450,
        25,
        "Occupied",
        "Fictional mobile-home lot",
    ])

    apartment = workbook.create_sheet("RUGBY")
    apartment.append([
        "Apt",
        "Name",
        "Move-In Date",
        "Lease Exp",
        "Rent",
        "Status",
    ])
    apartment.append([
        "2B",
        "Fictional Apartment Tenant",
        date(2024, 5, 1),
        date(2025, 4, 30),
        850,
        "Occupied",
    ])

    rv = workbook.create_sheet("Proforma Rent Roll")
    rv.append(["Site", "Guest Name", "Arrival", "Monthly Rent", "Site Type"])
    rv.append([
        "RV-7",
        "Fictional RV Guest",
        date(2024, 11, 14),
        600,
        "Pull-through RV site",
    ])
    workbook.save(path)
    workbook.close()


def _build_pdf_rent_roll(path):
    document = SimpleDocTemplate(
        str(path),
        pagesize=landscape(letter),
        title="Fictional Native PDF Rent Roll",
    )
    styles = getSampleStyleSheet()
    rows = [
        ["Fictional Commercial Rent Roll As of 06/30/2025"],
        [
            "Suite",
            "Tenant",
            "Use",
            "Rentable SF",
            "Lease Start",
            "Expiration",
            "Monthly Rent",
            "Annual Rent",
            "Rent/SF",
            "Status",
        ],
        [
            "201",
            "Fictional PDF Tenant Alpha",
            "Retail",
            "1,200",
            "01/01/2024",
            "12/31/2028",
            "$2,400",
            "$28,800",
            "$24.00",
            "Occupied",
        ],
        [
            "202",
            "Fictional PDF Tenant Beta",
            "Office",
            "900",
            "03/01/2025",
            "02/28/2030",
            "$1,950",
            "$23,400",
            "$26.00",
            "Occupied",
        ],
        ["Total", "", "", "2,100", "", "", "$4,350", "$52,200", "", ""],
    ]
    table = Table(rows, repeatRows=2)
    table.setStyle(TableStyle([
        ("SPAN", (0, 0), (-1, 0)),
        ("BACKGROUND", (0, 0), (-1, 1), colors.lightgrey),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("FONTNAME", (0, 0), (-1, 1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("ALIGN", (3, 2), (-1, -1), "RIGHT"),
    ]))
    document.build([
        Paragraph("Fictional Native PDF Rent Roll", styles["Title"]),
        Spacer(1, 8),
        table,
    ])


def _build_pdf_rent_roll_with_bad_annual_rent(path):
    """A single row whose annual rent doesn't reconcile with monthly rent x
    12 (every individual field still looks plausible) -- simulates an OCR
    digit misread that the arithmetic cross-check is meant to catch."""
    document = SimpleDocTemplate(str(path), pagesize=landscape(letter))
    styles = getSampleStyleSheet()
    rows = [
        [
            "Suite", "Tenant", "Use", "Rentable SF", "Lease Start",
            "Expiration", "Monthly Rent", "Annual Rent", "Rent/SF", "Status",
        ],
        [
            "201", "Fictional PDF Tenant Alpha", "Retail", "1,200",
            "01/01/2024", "12/31/2028", "$2,400", "$38,800", "$24.00",
            "Occupied",
        ],
    ]
    table = Table(rows, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
    ]))
    document.build([
        Paragraph("Fictional Mismatched Rent Roll", styles["Title"]),
        Spacer(1, 8),
        table,
    ])


def _build_pdf_rent_roll_with_unheaded_continuation_page(path):
    """Page 1: a normal headered rent-roll table. Page 2: a continuation
    table with a data row only and no repeated header -- a common
    real-world multi-page rent-roll layout that the original OCR lane
    silently dropped without any warning."""
    document = SimpleDocTemplate(str(path), pagesize=landscape(letter))
    styles = getSampleStyleSheet()
    page1_rows = [
        [
            "Suite", "Tenant", "Use", "Rentable SF", "Lease Start",
            "Expiration", "Monthly Rent", "Annual Rent", "Rent/SF", "Status",
        ],
        [
            "201", "Fictional PDF Tenant Alpha", "Retail", "1,200",
            "01/01/2024", "12/31/2028", "$2,400", "$28,800", "$24.00",
            "Occupied",
        ],
    ]
    page1_table = Table(page1_rows, repeatRows=1)
    page1_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
    ]))
    page2_rows = [
        [
            "202", "Fictional PDF Tenant Beta", "Office", "900",
            "03/01/2025", "02/28/2030", "$1,950", "$23,400", "$26.00",
            "Occupied",
        ],
    ]
    page2_table = Table(page2_rows)
    page2_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
    ]))
    document.build([
        Paragraph("Fictional Two-Page Rent Roll", styles["Title"]),
        Spacer(1, 8),
        page1_table,
        PageBreak(),
        page2_table,
    ])


def _rasterize_all_pages_to_scanned_pdf(native_pdf_path, dest_path, dpi=300):
    """Like _rasterize_to_scanned_pdf but renders every page of the source
    PDF into the resulting image-only PDF, instead of just page 1 -- used
    to simulate a multi-page scanned document."""
    zoom = dpi / 72.0
    matrix = fitz.Matrix(zoom, zoom)
    src = fitz.open(str(native_pdf_path))
    doc = fitz.open()
    img_paths = []
    try:
        for page in src:
            pix = page.get_pixmap(matrix=matrix)
            image = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            img_path = dest_path.with_suffix(f".p{page.number}.png")
            image.save(img_path, format="PNG")
            img_paths.append(img_path)
            rect = fitz.Rect(0, 0, image.width * 72 / dpi, image.height * 72 / dpi)
            new_page = doc.new_page(width=rect.width, height=rect.height)
            new_page.insert_image(rect, filename=str(img_path))
        doc.save(str(dest_path))
    finally:
        doc.close()
        src.close()
        for img_path in img_paths:
            img_path.unlink()


def _build_pdf_profit_and_loss(path):
    pdf = canvas.Canvas(str(path), pagesize=letter)
    pdf.setTitle("Fictional Native Text P&L")
    width, height = letter
    y = height - 72
    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(72, y, "Fictional Profit and Loss Statement 2024 Actual")
    y -= 30
    pdf.setFont("Helvetica-Bold", 11)
    pdf.drawString(72, y, "Income")
    y -= 18
    pdf.setFont("Helvetica", 10)
    pdf.drawString(90, y, "Rental Income")
    pdf.drawRightString(width - 72, y, "$120,000")
    y -= 18
    pdf.setFont("Helvetica-Bold", 10)
    pdf.drawString(72, y, "Total Income")
    pdf.drawRightString(width - 72, y, "$120,000")
    y -= 28
    pdf.setFont("Helvetica-Bold", 11)
    pdf.drawString(72, y, "Expenses")
    y -= 18
    pdf.setFont("Helvetica", 10)
    for label, amount in (
        ("Repairs and Maintenance", "$12,500"),
        ("Insurance", "$4,200"),
        ("Utilities", "$7,800"),
    ):
        pdf.drawString(90, y, label)
        pdf.drawRightString(width - 72, y, amount)
        y -= 18
    pdf.setFont("Helvetica-Bold", 10)
    pdf.drawString(72, y, "Total Expenses")
    pdf.drawRightString(width - 72, y, "$24,500")
    y -= 18
    pdf.drawString(72, y, "Net Income")
    pdf.drawRightString(width - 72, y, "$95,500")
    pdf.save()


class NoDimensionSheet:
    title = "No Dimension"
    max_row = None

    def iter_rows(self, min_row=1, max_row=None, values_only=True):
        rows = [
            ["Lot #", "Resident", "Move In", "Rent"],
            ["L-1", "Fictional Resident", date(2025, 1, 1), 450],
        ]
        return iter(rows[min_row - 1:max_row])



def _rasterize_to_scanned_pdf(native_pdf_path, dest_path, rotate_degrees=0, dpi=300):
    """Render page 1 of a native PDF to an image and re-embed it as a new
    image-only PDF (no text layer), optionally rotated, to simulate a
    scanned/photographed document for OCR-lane testing."""
    zoom = dpi / 72.0
    matrix = fitz.Matrix(zoom, zoom)
    src = fitz.open(str(native_pdf_path))
    pix = src[0].get_pixmap(matrix=matrix)
    image = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
    src.close()
    if rotate_degrees:
        image = image.rotate(-rotate_degrees, expand=True)
    img_path = dest_path.with_suffix(".png")
    image.save(img_path, format="PNG")
    doc = fitz.open()
    rect = fitz.Rect(0, 0, image.width * 72 / dpi, image.height * 72 / dpi)
    page = doc.new_page(width=rect.width, height=rect.height)
    page.insert_image(rect, filename=str(img_path))
    doc.save(str(dest_path))
    doc.close()
    img_path.unlink()


def _build_illegible_scan_pdf(dest_path, dpi=300):
    """Build an image-only PDF containing pure random noise -- no legible
    text at all -- to exercise the OCR lane's low-confidence bail-out path."""
    rng = np.random.default_rng(42)
    # A small noisy source image upscaled to a full page is still
    # illegible and keeps fixture generation fast.
    small = rng.integers(0, 256, size=(200, 160, 3), dtype=np.uint8)
    image = Image.fromarray(small, mode="RGB").resize(
        (int(8.5 * dpi), int(11 * dpi)), Image.NEAREST
    )
    img_path = dest_path.with_suffix(".png")
    image.save(img_path, format="PNG")
    doc = fitz.open()
    rect = fitz.Rect(0, 0, image.width * 72 / dpi, image.height * 72 / dpi)
    page = doc.new_page(width=rect.width, height=rect.height)
    page.insert_image(rect, filename=str(img_path))
    doc.save(str(dest_path))
    doc.close()
    img_path.unlink()


class FinancialHarvestTests(unittest.TestCase):
    def test_contract_descriptor_matches_financial_runtime(self):
        schema_path = (
            Path(__file__).resolve().parents[1]
            / "schemas"
            / "historical_harvest.v1.json"
        )
        schema = json.loads(schema_path.read_text(encoding="utf-8"))
        self.assertEqual(
            RENT_ROLL_CONTRACT_ID,
            schema["contracts"]["rent_roll"],
        )
        self.assertEqual(
            EXPENSE_CONTRACT_ID,
            schema["contracts"]["expense"],
        )

    def test_fictional_rent_roll_and_expenses_end_to_end(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            assignment = (
                root
                / "historical"
                / "25C900 - Office - 777 Fictional Finance Road, Demo City"
            )
            assignment.mkdir(parents=True)
            _build_report(assignment / "REPORT 25C900.docx")
            _build_rent_roll(assignment / "Fictional Rent Roll.xlsx")
            _build_expenses(assignment / "Operating Expenses 2025.xlsx")

            staged_paths = run_extraction(
                root / "historical",
                staged_dir=root / "staged",
            )
            self.assertEqual(1, len(staged_paths))
            staged = json.loads(staged_paths[0].read_text(encoding="utf-8"))
            self.assertEqual(2, len(staged["rent_roll_entries"]))
            self.assertEqual(2, len(staged["expense_records"]))

            rent = staged["rent_roll_entries"][0]
            expense = staged["expense_records"][0]
            self.assertEqual(RENT_ROLL_CONTRACT_ID, rent["contract_id"])
            self.assertEqual(EXPENSE_CONTRACT_ID, expense["contract_id"])
            self.assertEqual("2025-06-30", rent["data"]["as_of_date"])
            self.assertEqual("2024-01-01", rent["data"]["lease_start"])
            self.assertEqual(
                "worksheet:Rent Roll:row:4",
                rent["provenance"]["source_locator"],
            )
            self.assertEqual(
                "worksheet:Operating Expenses:row:2",
                expense["provenance"]["source_locator"],
            )
            self.assertEqual("unreviewed", rent["review"]["status"])

            rent["data"]["monthly_rent"] = 2_100
            rent["review_edits"] = [{
                "field": "monthly_rent",
                "before": 2_000,
                "after": 2_100,
            }]
            confirmed = confirm_extraction_result(
                staged,
                reviewer="fictional-financial-qa",
                reviewed_at="2026-07-01T15:00:00+00:00",
            )

            db_path = root / "financial.db"
            init_db(db_path, quiet=True)
            connection = get_conn(db_path)
            try:
                with connection:
                    counts = commit_extraction_result(confirmed, connection)
                self.assertEqual(2, counts["rent_roll_entries"])
                self.assertEqual(2, counts["operating_expenses"])
                with connection:
                    duplicates = commit_extraction_result(confirmed, connection)
                self.assertEqual(
                    2,
                    duplicates["duplicate_rent_roll_entries"],
                )
                self.assertEqual(
                    2,
                    duplicates["duplicate_operating_expenses"],
                )
            finally:
                connection.close()

            rent_rows = search_rent_roll_entries(
                db_path,
                tenant_contains="Alpha",
                as_of_date="2025-06-30",
            )
            expense_rows = search_operating_expenses(
                db_path,
                period_year=2025,
                category_contains="tax",
            )
            self.assertEqual(1, len(rent_rows))
            self.assertEqual(2_100, rent_rows[0]["monthly_rent"])
            self.assertEqual("25C900", rent_rows[0]["file_no"])
            self.assertEqual("confirmed", rent_rows[0]["review_status"])
            self.assertEqual(1, len(expense_rows))
            self.assertEqual(25_000, expense_rows[0]["amount"])
            self.assertEqual("confirmed", expense_rows[0]["review_status"])

    def test_wide_operating_statement_explodes_to_expense_lines(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            assignment = (
                root
                / "historical"
                / "25C902 - Office - 999 Fictional Finance Road, Demo City"
            )
            assignment.mkdir(parents=True)
            _build_report(assignment / "REPORT 25C902.docx")
            _build_wide_expenses(
                assignment / "Wide Operating Statement 2023-2025.xlsx"
            )

            staged_path = run_extraction(
                root / "historical",
                staged_dir=root / "staged",
            )[0]
            staged = json.loads(staged_path.read_text(encoding="utf-8"))
            self.assertEqual(6, len(staged["expense_records"]))

            taxes_2023 = next(
                record
                for record in staged["expense_records"]
                if record["data"]["category"] == "Real Estate Taxes"
                and record["data"]["period_year"] == 2023
            )
            taxes_2025 = next(
                record
                for record in staged["expense_records"]
                if record["data"]["category"] == "Real Estate Taxes"
                and record["data"]["period_year"] == 2025
            )
            self.assertEqual(20_000, taxes_2023["data"]["amount"])
            self.assertEqual(8, taxes_2023["data"]["amount_per_sf"])
            self.assertEqual("actual", taxes_2023["data"]["period_type"])
            self.assertEqual("budget", taxes_2025["data"]["period_type"])
            self.assertEqual(
                "worksheet:Operating Statement:row:3:cols:2-3",
                taxes_2023["provenance"]["source_locator"],
            )
            self.assertEqual(
                "wide_operating_statement",
                taxes_2023["provenance"]["layout"],
            )

            confirmed = confirm_extraction_result(
                staged,
                reviewer="fictional-wide-financial-qa",
                reviewed_at="2026-07-02T15:00:00+00:00",
            )
            db_path = root / "wide-financial.db"
            init_db(db_path, quiet=True)
            connection = get_conn(db_path)
            try:
                with connection:
                    counts = commit_extraction_result(confirmed, connection)
                self.assertEqual(6, counts["operating_expenses"])
                rows = connection.execute(
                    """
                    SELECT amount, amount_per_sf, period_type, source_record_json
                    FROM operating_expenses
                    WHERE category = ? AND period_year = ?
                    """,
                    ("Real Estate Taxes", 2025),
                ).fetchall()
                self.assertEqual(1, len(rows))
                self.assertEqual(30_000, rows[0]["amount"])
                self.assertEqual("budget", rows[0]["period_type"])
                self.assertIn(
                    "wide_operating_statement",
                    rows[0]["source_record_json"],
                )
            finally:
                connection.close()

            searched = search_operating_expenses(
                db_path,
                period_year=2024,
                category_contains="insurance",
            )
            self.assertEqual(1, len(searched))
            self.assertEqual(12_000, searched[0]["amount"])
            self.assertEqual(4.8, searched[0]["amount_per_sf"])

    def test_specialty_rent_roll_layouts_are_normalized_and_deduped(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            assignment = (
                root
                / "historical"
                / "25C903 - RV Park - 111 Fictional Specialty Road, Demo City"
            )
            assignment.mkdir(parents=True)
            _build_report(assignment / "REPORT 25C903.docx")
            _build_specialty_rent_rolls(
                assignment / "Fictional Specialty Rent Roll.xlsx"
            )

            staged_path = run_extraction(
                root / "historical",
                staged_dir=root / "staged",
            )[0]
            staged = json.loads(staged_path.read_text(encoding="utf-8"))
            rents = staged["rent_roll_entries"]
            self.assertEqual(5, len(rents))

            storage = next(
                record
                for record in rents
                if record["data"].get("unit_id") == "A-101"
            )
            mobile_home = next(
                record
                for record in rents
                if record["data"].get("unit_id") == "L-12"
            )
            apartment = next(
                record
                for record in rents
                if record["data"].get("unit_id") == "2B"
            )
            rv = next(
                record
                for record in rents
                if record["data"].get("unit_id") == "RV-7"
            )
            self.assertEqual(125, storage["data"]["monthly_rent"])
            self.assertEqual("10x10 Outside", storage["data"]["tenant_use"])
            self.assertEqual(1, len(storage.get("alternate_provenance", [])))
            self.assertEqual("2026-01-01", mobile_home["data"]["as_of_date"])
            self.assertEqual(
                "discounts: 25.0",
                mobile_home["data"]["reimbursement_structure"],
            )
            self.assertEqual(
                "2025-04-30",
                apartment["data"]["lease_end"],
            )
            self.assertEqual("Pull-through RV site", rv["data"]["tenant_use"])

            confirmed = confirm_extraction_result(
                staged,
                reviewer="fictional-specialty-rent-qa",
                reviewed_at="2026-07-03T15:00:00+00:00",
            )
            db_path = root / "specialty-rent.db"
            init_db(db_path, quiet=True)
            connection = get_conn(db_path)
            try:
                with connection:
                    counts = commit_extraction_result(confirmed, connection)
                self.assertEqual(5, counts["rent_roll_entries"])
            finally:
                connection.close()

            rows = search_rent_roll_entries(
                db_path,
                tenant_contains="Mobile",
                as_of_date="2026-01-01",
            )
            self.assertEqual(1, len(rows))
            self.assertEqual("L-12", rows[0]["unit_id"])
            self.assertEqual("confirmed", rows[0]["review_status"])

    def test_native_pdf_rent_roll_table_end_to_end(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            assignment = (
                root
                / "historical"
                / "25C904 - Retail - 222 Fictional PDF Road, Demo City"
            )
            assignment.mkdir(parents=True)
            _build_report(assignment / "REPORT 25C904.docx")
            pdf_path = assignment / "Fictional Rent Roll.pdf"
            _build_pdf_rent_roll(pdf_path)

            direct = extract_financial_pdf(pdf_path)
            self.assertEqual(2, len(direct["rent_roll_entries"]))
            self.assertEqual([], direct["warnings"])

            staged_path = run_extraction(
                root / "historical",
                staged_dir=root / "staged",
            )[0]
            staged = json.loads(staged_path.read_text(encoding="utf-8"))
            self.assertEqual(2, len(staged["rent_roll_entries"]))
            first = staged["rent_roll_entries"][0]
            self.assertEqual(RENT_ROLL_CONTRACT_ID, first["contract_id"])
            self.assertEqual("201", first["data"]["suite"])
            self.assertEqual("Fictional PDF Tenant Alpha", first["data"]["tenant_name"])
            self.assertEqual(1_200, first["data"]["sf_leased"])
            self.assertEqual(2_400, first["data"]["monthly_rent"])
            self.assertEqual("2025-06-30", first["data"]["as_of_date"])
            self.assertEqual(
                "native_pdf_table_extractor",
                first["provenance"]["extraction_method"],
            )
            self.assertTrue(
                first["provenance"]["source_locator"].startswith(
                    "pdf:page:1:table:"
                )
            )

            confirmed = confirm_extraction_result(
                staged,
                reviewer="fictional-pdf-rent-qa",
                reviewed_at="2026-07-04T15:00:00+00:00",
            )
            db_path = root / "pdf-rent.db"
            init_db(db_path, quiet=True)
            connection = get_conn(db_path)
            try:
                with connection:
                    counts = commit_extraction_result(confirmed, connection)
                self.assertEqual(2, counts["rent_roll_entries"])
            finally:
                connection.close()

            rows = search_rent_roll_entries(
                db_path,
                tenant_contains="PDF Tenant Alpha",
                as_of_date="2025-06-30",
            )
            self.assertEqual(1, len(rows))
            self.assertEqual("25C904", rows[0]["file_no"])
            self.assertEqual("confirmed", rows[0]["review_status"])

    def test_native_text_position_pdf_expenses_end_to_end(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            assignment = (
                root
                / "historical"
                / "25C905 - Retail - 333 Fictional Text PDF Road, Demo City"
            )
            assignment.mkdir(parents=True)
            _build_report(assignment / "REPORT 25C905.docx")
            pdf_path = assignment / "Fictional Profit and Loss 2024.pdf"
            _build_pdf_profit_and_loss(pdf_path)

            direct = extract_financial_pdf(pdf_path)
            self.assertEqual(3, len(direct["expense_records"]))
            self.assertEqual([], direct["warnings"])

            staged_path = run_extraction(
                root / "historical",
                staged_dir=root / "staged",
            )[0]
            staged = json.loads(staged_path.read_text(encoding="utf-8"))
            self.assertEqual(3, len(staged["expense_records"]))
            categories = {
                record["data"]["category"]: record
                for record in staged["expense_records"]
            }
            self.assertEqual(
                12_500,
                categories["Repairs and Maintenance"]["data"]["amount"],
            )
            self.assertEqual(
                "actual",
                categories["Repairs and Maintenance"]["data"]["period_type"],
            )
            self.assertEqual(
                2024,
                categories["Repairs and Maintenance"]["data"]["period_year"],
            )
            self.assertEqual(
                "native_pdf_text_position_extractor",
                categories["Repairs and Maintenance"]["provenance"][
                    "extraction_method"
                ],
            )
            self.assertTrue(
                categories["Repairs and Maintenance"]["provenance"][
                    "source_locator"
                ].startswith("pdf:page:1:line:")
            )

            confirmed = confirm_extraction_result(
                staged,
                reviewer="fictional-pdf-expense-qa",
                reviewed_at="2026-07-05T15:00:00+00:00",
            )
            db_path = root / "pdf-expense.db"
            init_db(db_path, quiet=True)
            connection = get_conn(db_path)
            try:
                with connection:
                    counts = commit_extraction_result(confirmed, connection)
                self.assertEqual(3, counts["operating_expenses"])
            finally:
                connection.close()

            rows = search_operating_expenses(
                db_path,
                period_year=2024,
                category_contains="insurance",
            )
            self.assertEqual(1, len(rows))
            self.assertEqual(4_200, rows[0]["amount"])
            self.assertEqual("confirmed", rows[0]["review_status"])

    def test_header_detection_survives_missing_worksheet_dimensions(self):
        header_row, mapping = _find_header(
            NoDimensionSheet(),
            RENT_ROLL_SYNONYMS,
            3,
        )
        self.assertEqual(1, header_row)
        self.assertIn("unit_id", mapping.values())
        self.assertIn("tenant_name", mapping.values())
        self.assertIn("monthly_rent", mapping.values())

    def test_unconfirmed_financial_row_rolls_back_everything(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            assignment = (
                root
                / "historical"
                / "25C901 - Office - 888 Fictional Finance Road, Demo City"
            )
            assignment.mkdir(parents=True)
            _build_report(assignment / "REPORT 25C901.docx")
            _build_rent_roll(assignment / "Fictional Rent Roll.xlsx")
            staged_path = run_extraction(
                root / "historical",
                staged_dir=root / "staged",
            )[0]
            staged = json.loads(staged_path.read_text(encoding="utf-8"))
            confirmed = confirm_extraction_result(staged, reviewer="qa")
            confirmed["rent_roll_entries"][0]["review"]["status"] = "unreviewed"

            db_path = root / "rollback.db"
            init_db(db_path, quiet=True)
            connection = get_conn(db_path)
            try:
                with self.assertRaisesRegex(ValueError, "not been confirmed"):
                    with connection:
                        commit_extraction_result(confirmed, connection)
                for table in (
                    "assignments",
                    "rent_roll_entries",
                    "source_documents",
                ):
                    self.assertEqual(
                        0,
                        connection.execute(
                            f"SELECT count(*) FROM {table}"
                        ).fetchone()[0],
                    )
            finally:
                connection.close()

    def test_database_initialization_adds_financial_tables(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            db_path = Path(temp_dir) / "fresh.db"
            init_db(db_path, quiet=True)
            connection = get_conn(db_path)
            try:
                tables = {
                    row[0]
                    for row in connection.execute(
                        "SELECT name FROM sqlite_master WHERE type = 'table'"
                    )
                }
                self.assertIn("rent_roll_entries", tables)
                self.assertIn("operating_expenses", tables)
            finally:
                connection.close()


    @unittest.skipUnless(
        _OCR_TEST_DEPS_AVAILABLE and _TESSERACT_AVAILABLE,
        "requires PyMuPDF/Pillow and a working local Tesseract install",
    )
    def test_ocr_scanned_pdf_rent_roll_end_to_end(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            native_pdf = root / "native.pdf"
            _build_pdf_rent_roll(native_pdf)
            scanned_pdf = root / "Fictional Scanned Rent Roll.pdf"
            _rasterize_to_scanned_pdf(native_pdf, scanned_pdf)

            result = extract_financial_pdf(scanned_pdf)
            self.assertEqual([], result["warnings"])
            self.assertEqual(2, len(result["rent_roll_entries"]))

            first = result["rent_roll_entries"][0]
            self.assertEqual(
                "ocr_pdf_table_extractor",
                first["provenance"]["extraction_method"],
            )
            self.assertEqual("tesseract", first["provenance"]["ocr_engine"])
            self.assertGreater(
                first["provenance"]["ocr_avg_word_confidence"], 0
            )
            self.assertTrue(
                all(value == "low" for value in first["confidence"].values()),
                "every OCR-derived field must be tagged confidence low",
            )
            rendered = first["provenance"].get("rendered_page_image")
            self.assertIsNotNone(rendered)
            image_path = Path(__file__).resolve().parents[1] / rendered
            self.assertTrue(image_path.is_file())

            # Numeric/date columns should still resolve correctly even though
            # the source has no text layer at all.
            self.assertEqual("201", first["data"]["suite"])
            self.assertEqual(1_200, first["data"]["sf_leased"])
            self.assertEqual(2_400, first["data"]["monthly_rent"])
            self.assertEqual(28_800, first["data"]["annual_rent"])
            self.assertEqual(24.0, first["data"]["rent_psf"])
            self.assertEqual("2024-01-01", first["data"]["lease_start"])
            self.assertEqual("2028-12-31", first["data"]["lease_end"])
            self.assertEqual("2025-06-30", first["data"]["as_of_date"])
            self.assertEqual("Occupied", first["data"]["occupancy_status"])

            second = result["rent_roll_entries"][1]
            self.assertEqual("202", second["data"]["suite"])
            self.assertEqual(1_950, second["data"]["monthly_rent"])

            # Confirm this flows through the normal stage -> review -> commit
            # gate exactly like every other harvest record -- no special
            # OCR-only commit path exists.
            assignment = root / "historical" / "25C906 - Retail - 444 Fictional Scan Road, Demo City"
            assignment.mkdir(parents=True)
            _build_report(assignment / "REPORT 25C906.docx")
            import shutil
            shutil.copy(scanned_pdf, assignment / "Fictional Scanned Rent Roll.pdf")

            staged_path = run_extraction(
                root / "historical",
                staged_dir=root / "staged",
            )[0]
            staged = json.loads(staged_path.read_text(encoding="utf-8"))
            self.assertEqual(2, len(staged["rent_roll_entries"]))
            self.assertEqual("unreviewed", staged["rent_roll_entries"][0]["review"]["status"])

            confirmed = confirm_extraction_result(
                staged,
                reviewer="fictional-ocr-qa",
                reviewed_at="2026-07-08T15:00:00+00:00",
            )
            db_path = root / "ocr-rent.db"
            init_db(db_path, quiet=True)
            connection = get_conn(db_path)
            try:
                with connection:
                    counts = commit_extraction_result(confirmed, connection)
                self.assertEqual(2, counts["rent_roll_entries"])
            finally:
                connection.close()

            rows = search_rent_roll_entries(
                db_path,
                tenant_contains="",
                as_of_date="2025-06-30",
            )
            self.assertEqual(2, len(rows))
            self.assertEqual("confirmed", rows[0]["review_status"])

    @unittest.skipUnless(
        _OCR_TEST_DEPS_AVAILABLE and _TESSERACT_AVAILABLE,
        "requires PyMuPDF/Pillow and a working local Tesseract install",
    )
    def test_ocr_rotated_scan_recovers_orientation(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            native_pdf = root / "native.pdf"
            _build_pdf_rent_roll(native_pdf)
            scanned_pdf = root / "Fictional Rotated Rent Roll.pdf"
            _rasterize_to_scanned_pdf(native_pdf, scanned_pdf, rotate_degrees=180)

            result = extract_financial_pdf(scanned_pdf)
            self.assertEqual(2, len(result["rent_roll_entries"]))
            first = result["rent_roll_entries"][0]
            self.assertEqual(180, first["provenance"]["rotation_degrees_applied"])
            self.assertEqual("201", first["data"]["suite"])
            self.assertEqual(2_400, first["data"]["monthly_rent"])

    @unittest.skipUnless(
        _OCR_TEST_DEPS_AVAILABLE,
        "requires PyMuPDF/numpy/Pillow to build the synthetic scan fixture",
    )
    def test_ocr_illegible_scan_bails_out_with_warning(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            garbage_pdf = root / "Fictional Illegible Scan.pdf"
            _build_illegible_scan_pdf(garbage_pdf)

            result = extract_financial_pdf(garbage_pdf)
            self.assertEqual([], result["rent_roll_entries"])
            self.assertEqual([], result["expense_records"])
            self.assertTrue(result["warnings"], "expected a bail-out warning")
            joined = " ".join(result["warnings"]).lower()
            self.assertTrue(
                "confidence too low" in joined
                or "no recognizable" in joined
                or "ocr is required" in joined,
                result["warnings"],
            )

    @unittest.skipUnless(
        _OCR_TEST_DEPS_AVAILABLE,
        "requires PyMuPDF/Pillow to build the synthetic scan fixture",
    )
    def test_ocr_lane_degrades_gracefully_without_tesseract(self):
        import pdf_financial_extractor as pf

        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            native_pdf = root / "native.pdf"
            _build_pdf_rent_roll(native_pdf)
            scanned_pdf = root / "Fictional Scanned Rent Roll No Engine.pdf"
            _rasterize_to_scanned_pdf(native_pdf, scanned_pdf)

            original = pf._tesseract_binary_ok
            pf._tesseract_binary_ok = False
            try:
                result = extract_financial_pdf(scanned_pdf)
            finally:
                pf._tesseract_binary_ok = original

            self.assertEqual([], result["rent_roll_entries"])
            self.assertEqual([], result["expense_records"])
            self.assertEqual(1, len(result["warnings"]))
            self.assertIn("Tesseract is not installed", result["warnings"][0])

    @unittest.skipUnless(
        _OCR_TEST_DEPS_AVAILABLE and _TESSERACT_AVAILABLE,
        "requires PyMuPDF/Pillow and a working local Tesseract install",
    )
    def test_ocr_arithmetic_mismatch_warns_without_dropping_row(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            native_pdf = root / "native_mismatch.pdf"
            _build_pdf_rent_roll_with_bad_annual_rent(native_pdf)
            scanned_pdf = root / "Fictional Mismatched Scan.pdf"
            _rasterize_to_scanned_pdf(native_pdf, scanned_pdf)

            result = extract_financial_pdf(scanned_pdf)

            # The row is still staged for review -- confidence="low" plus a
            # warning, not silently dropped -- because a checksum failure
            # doesn't mean the row is worthless, just that it needs a human
            # to check it against the source scan.
            self.assertEqual(1, len(result["rent_roll_entries"]))
            self.assertEqual(
                38_800, result["rent_roll_entries"][0]["data"]["annual_rent"]
            )
            joined = " ".join(result["warnings"]).lower()
            self.assertIn("does not reconcile", joined)

    @unittest.skipUnless(
        _OCR_TEST_DEPS_AVAILABLE and _TESSERACT_AVAILABLE,
        "requires PyMuPDF/Pillow and a working local Tesseract install",
    )
    def test_ocr_continuation_page_without_header_warns_instead_of_silent_loss(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            native_pdf = root / "native_two_page.pdf"
            _build_pdf_rent_roll_with_unheaded_continuation_page(native_pdf)
            scanned_pdf = root / "Fictional Two Page Scan.pdf"
            _rasterize_all_pages_to_scanned_pdf(native_pdf, scanned_pdf)

            result = extract_financial_pdf(scanned_pdf)

            # Page 1's headered row still extracts normally.
            self.assertEqual(1, len(result["rent_roll_entries"]))
            self.assertEqual(
                "201", result["rent_roll_entries"][0]["data"]["suite"]
            )
            # Page 2's headerless continuation row can't be column-matched
            # and is correctly not extracted -- but that gap must be
            # surfaced as a warning naming the page, not swallowed silently.
            joined = " ".join(result["warnings"]).lower()
            self.assertIn("no rent-roll header recognized", joined)
            self.assertIn("page 2", joined)


if __name__ == "__main__":
    unittest.main()
