import base64
import os
import shutil
import tempfile
import unittest
import json
from pathlib import Path
from unittest.mock import patch

import openpyxl
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

import axiom
from comp_builder import inject_comp_section, load_comp_data
from field_registry import (
    audit_assignment_contract,
    build_registry,
    load_registry,
)
from fill_engine import load_variables
from media_blocks import inject_media_blocks
from narrative_generator import _get_model
from presentation_variants import derive_presentation_variants
from structured_blocks import inject_ownership_history
from validation import find_docx_placeholders, validate_assignment


def _build_assignment(root, template_text, variables=None):
    root = Path(root)
    assignment = root / "TEST-900_Fictional_Client"
    templates = root / "templates"
    assignment.mkdir()
    templates.mkdir()

    workbook = openpyxl.Workbook()
    outputs = workbook.active
    outputs.title = "outputs"
    outputs.append(["Label", "Key", "Raw", "Formatted"])
    workbook.save(assignment / "workbook.xlsx")

    variables = variables or {}
    with open(assignment / "TEST-900_variables.json", "w", encoding="utf-8") as f:
        json.dump(variables, f)

    document = Document()
    document.add_paragraph(template_text)
    document.save(templates / "report.docx")

    config = {"documents": [{"template": "report.docx"}]}
    return assignment, templates, config


def _write_registry(root, fields):
    registry_path = Path(root) / "field_registry.json"
    registry_path.write_text(
        json.dumps(
            {
                "contract_id": "axiom.appraisal.fields",
                "schema_version": "test",
                "fields": fields,
                "blocks": {},
            }
        ),
        encoding="utf-8",
    )
    return registry_path


class ValidateAssignmentTests(unittest.TestCase):
    TINY_PNG = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk"
        "/x8AAusB9Wl2nWQAAAAASUVORK5CYII="
    )

    def test_ready_when_required_values_are_present(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            assignment, templates, config = _build_assignment(
                temp_dir,
                "Value: [[VALUE_CONCLUSION]]",
                {"VALUE_CONCLUSION": "$1,000,000"},
            )

            result = validate_assignment(assignment, templates, config)

            self.assertTrue(result["checked"])
            self.assertTrue(result["ready"])
            self.assertEqual([], result["missing"])
            self.assertEqual({}, result["unresolved_blocks"])

    def test_presentation_variants_derive_from_canonical_fields(self):
        variables = derive_presentation_variants(
            {
                "PROPERTY_CLASS": "Class B",
                "PROPERTY_CLASS_LOWER": "stale",
                "PROPERTY_SUBTYPE_FULL": "Multi-Tenant Office Building",
                "VALUE_INTEREST": "As-Is Market Value in Leased Fee Estate",
                "VALUE_WORDS": (
                    "ONE MILLION THREE HUNDRED SEVENTY-FIVE THOUSAND DOLLARS"
                ),
                "ZONING_CLASS": "General Business District",
                "ZONING_CODE": "C-5",
            }
        )

        self.assertEqual("class b", variables["PROPERTY_CLASS_LOWER"])
        self.assertEqual(
            "multi-tenant office building",
            variables["PROPERTY_SUBTYPE_LOWER"],
        )
        self.assertEqual(
            "as-is market value in leased fee estate",
            variables["VALUE_INTEREST_LOWER"],
        )
        self.assertEqual(
            "One Million Three Hundred Seventy-Five Thousand Dollars",
            variables["VALUE_WORDS_FORMAL"],
        )
        self.assertEqual(
            "General Business District",
            variables["ZONING_CLASS_TABLE"],
        )
        self.assertEqual("C-5", variables["ZONING_CODE_TABLE"])

    def test_legacy_variant_survives_when_canonical_source_is_absent(self):
        variables = derive_presentation_variants(
            {"PROPERTY_CLASS_LOWER": "legacy class"}
        )
        self.assertEqual("legacy class", variables["PROPERTY_CLASS_LOWER"])

    def test_variable_loading_refreshes_stale_presentation_variants(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            json_path = Path(temp_dir) / "variables.json"
            json_path.write_text(
                json.dumps(
                    {
                        "PROPERTY_CLASS": "Class C",
                        "PROPERTY_CLASS_LOWER": "stale class",
                    }
                ),
                encoding="utf-8",
            )
            variables = load_variables(json_path=json_path)
            self.assertEqual("class c", variables["PROPERTY_CLASS_LOWER"])

    def test_intake_json_drift_identifies_changed_fields(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            assignment, templates, config = _build_assignment(
                temp_dir,
                "Client: [[CLIENT_NAME]]",
                {"CLIENT_NAME": "Old Example Client"},
            )
            workbook_path = assignment / "workbook.xlsx"
            workbook = openpyxl.load_workbook(workbook_path)
            intake = workbook.create_sheet("Intake")
            intake.append(["CLIENT_NAME", "New Example Client"])
            workbook.save(workbook_path)
            workbook.close()
            registry_path = _write_registry(
                temp_dir,
                {
                    "CLIENT_NAME": {
                        "value_kind": "text",
                        "source_of_truth": "intake",
                        "producers": ["intake"],
                        "used_in": ["deliver"],
                    }
                },
            )

            result = validate_assignment(
                assignment,
                templates,
                config,
                registry_path=registry_path,
            )

            self.assertFalse(result["ready"])
            self.assertEqual(["CLIENT_NAME"], result["stale_intake_fields"])
            self.assertTrue(
                any(
                    "Re-export JSON" in error
                    for error in result["errors"]
                )
            )

    def test_engagement_stops_when_intake_json_is_stale(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            assignment = Path(temp_dir) / "TEST-902_Fictional_Client"
            assignment.mkdir()
            (assignment / "outputs").mkdir()
            (assignment / "TEST-902_variables.json").write_text(
                json.dumps({"CLIENT_NAME": "Old Example Client"}),
                encoding="utf-8",
            )
            workbook = openpyxl.Workbook()
            intake = workbook.active
            intake.title = "Intake"
            intake.append(["CLIENT_NAME", "New Example Client"])
            workbook.save(assignment / "workbook.xlsx")
            workbook.close()

            with patch.object(
                axiom,
                "_find_assignment",
                return_value=assignment,
            ):
                result = axiom.cmd_engage(["TEST-902"])

            self.assertFalse(result)
            self.assertEqual([], list((assignment / "outputs").iterdir()))

    def test_older_json_with_matching_intake_is_not_stale(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            assignment, templates, config = _build_assignment(
                temp_dir,
                "Fee: [[FEE_AMOUNT]]",
                {"FEE_AMOUNT": "$3,250.00"},
            )
            workbook_path = assignment / "workbook.xlsx"
            workbook = openpyxl.load_workbook(workbook_path)
            intake = workbook.create_sheet("Intake")
            intake.append(["FEE_AMOUNT", 3250])
            workbook.save(workbook_path)
            workbook.close()
            registry_path = _write_registry(
                temp_dir,
                {
                    "FEE_AMOUNT": {
                        "value_kind": "currency_text",
                        "source_of_truth": "intake",
                        "producers": ["intake"],
                        "used_in": ["deliver"],
                    }
                },
            )
            json_path = assignment / "TEST-900_variables.json"
            os.utime(json_path, (1, 1))

            result = validate_assignment(
                assignment,
                templates,
                config,
                registry_path=registry_path,
            )

            self.assertTrue(result["ready"])
            self.assertEqual([], result["stale_intake_fields"])
            self.assertFalse(
                any("older than workbook" in warning for warning in result["warnings"])
            )

    def test_disabled_approach_ignores_its_formula_errors(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            assignment, templates, config = _build_assignment(
                temp_dir,
                "temporary",
                {
                    "CA_DEVELOPED": "No",
                    "VALUE_CONCLUSION": "$1,000,000",
                },
            )
            workbook_path = assignment / "workbook.xlsx"
            workbook = openpyxl.load_workbook(workbook_path)
            outputs = workbook["outputs"]
            outputs.append(["", "CA_VALUE", "#DIV/0!", "#DIV/0!"])
            workbook.save(workbook_path)
            workbook.close()

            document = Document()
            document.styles.add_style(
                "MainSectionHeading",
                WD_STYLE_TYPE.PARAGRAPH,
            )
            document.add_paragraph("[[VALUE_CONCLUSION]]")
            document.add_paragraph(
                "Cost Approach",
                style="MainSectionHeading",
            )
            document.add_paragraph("[[CA_VALUE]]")
            document.add_paragraph(
                "Sales Comparison Approach",
                style="MainSectionHeading",
            )
            document.save(templates / "report.docx")

            disabled = validate_assignment(assignment, templates, config)
            self.assertTrue(disabled["ready"])

            json_path = assignment / "TEST-900_variables.json"
            json_path.write_text(
                json.dumps(
                    {
                        "CA_DEVELOPED": "Yes",
                        "VALUE_CONCLUSION": "$1,000,000",
                    }
                ),
                encoding="utf-8",
            )
            enabled = validate_assignment(assignment, templates, config)
            self.assertFalse(enabled["ready"])
            self.assertTrue(
                any("CA_VALUE" in error for error in enabled["errors"])
            )

    def test_missing_value_and_unsupported_block_prevent_readiness(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            assignment, templates, config = _build_assignment(
                temp_dir,
                "[[VALUE_CONCLUSION]] [[SUBJECT_PHOTOS_BLOCK]]",
            )

            result = validate_assignment(assignment, templates, config)

            self.assertFalse(result["ready"])
            self.assertEqual(["VALUE_CONCLUSION"], result["missing"])
            self.assertIn(
                "SUBJECT_PHOTOS_BLOCK",
                result["unresolved_blocks"],
            )

    def test_narrative_requires_local_api_capability(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            assignment, templates, config = _build_assignment(
                temp_dir,
                "[[INSPECTION_NARRATIVE]]",
            )

            with patch.dict(os.environ, {}, clear=True):
                result = validate_assignment(assignment, templates, config)

            self.assertFalse(result["ready"])
            self.assertIn(
                "INSPECTION_NARRATIVE",
                result["unresolved_blocks"],
            )

    def test_failed_validation_cannot_mark_assignment_delivered(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            assignment = Path(temp_dir) / "TEST-901_Fictional_Client"
            assignment.mkdir()
            initial_state = {
                "file_no": "TEST-901",
                "stage": "new",
                "delivered": None,
            }
            with open(assignment / ".axiom.json", "w", encoding="utf-8") as f:
                json.dump(initial_state, f)

            failed_result = {
                "checked": True,
                "ready": False,
                "errors": [],
                "missing": ["VALUE_CONCLUSION"],
                "unresolved_blocks": {"SUBJECT_PHOTOS_BLOCK": "No handler."},
            }

            with (
                patch.object(axiom, "_find_assignment", return_value=assignment),
                patch.object(
                    axiom,
                    "check_delivery_readiness",
                    return_value=failed_result,
                ),
            ):
                command_result = axiom.cmd_deliver(["TEST-901"])

            with open(assignment / ".axiom.json", encoding="utf-8") as f:
                final_state = json.load(f)

            self.assertFalse(command_result)
            self.assertEqual("new", final_state["stage"])
            self.assertIsNone(final_state["delivered"])
            self.assertEqual("blocked", final_state["last_delivery_status"])
            self.assertEqual(2, final_state["last_delivery_blocker_count"])

    def test_sanitized_fixture_is_readable(self):
        project_root = Path(__file__).resolve().parents[1]
        with open(project_root / "config.json", encoding="utf-8") as config_file:
            config = json.load(config_file)

        result = validate_assignment(
            project_root / "tests" / "fixtures" / "DEMO-001",
            project_root / config["templates_dir"],
            config["stages"]["deliver"],
            registry_path=project_root / config["field_registry"],
        )

        self.assertTrue(result["checked"])
        self.assertEqual([], result["missing"])
        self.assertEqual([], result["stale_intake_fields"])
        self.assertEqual("DEMO-001", result["assignment"])

    def test_media_asset_is_validated_and_injected(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            assignment, templates, config = _build_assignment(
                temp_dir,
                "[[REGIONAL_MAP_IMAGE]]",
            )
            maps_dir = assignment / "assets" / "maps"
            maps_dir.mkdir(parents=True)
            (maps_dir / "regional.png").write_bytes(self.TINY_PNG)

            result = validate_assignment(assignment, templates, config)
            self.assertTrue(result["ready"])
            self.assertEqual(["REGIONAL_MAP_IMAGE"], result["handled_blocks"])

            output_path = Path(temp_dir) / "media-output.docx"
            shutil.copy(templates / "report.docx", output_path)
            injected = inject_media_blocks(output_path, assignment)
            self.assertEqual({"REGIONAL_MAP_IMAGE": 1}, injected)
            self.assertNotIn(
                "REGIONAL_MAP_IMAGE",
                find_docx_placeholders(output_path),
            )

    def test_demo_fixture_injects_all_synthetic_media(self):
        project_root = Path(__file__).resolve().parents[1]
        assignment = project_root / "tests" / "fixtures" / "DEMO-001"
        expected = {
            "REGIONAL_MAP_IMAGE": 1,
            "AERIAL_MAP_IMAGE": 1,
            "PARCEL_MAP_IMAGE": 1,
            "SCA_SALE_LOCATION_MAP": 1,
            "LAND_SALE_LOCATION_MAP": 1,
            "LEASE_COMP_LOCATION_MAP": 1,
            "BUILDING_SKETCH_BLOCK": 1,
            "SUBJECT_PHOTOS_BLOCK": 2,
            "LEASE_COMP_PHOTOS_BLOCK": 2,
        }

        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "all-media.docx"
            document = Document()
            for block in expected:
                document.add_paragraph(f"[[{block}]]")
            document.save(output_path)

            self.assertEqual(
                expected,
                inject_media_blocks(output_path, assignment),
            )
            self.assertEqual([], find_docx_placeholders(output_path))
            self.assertEqual(11, len(Document(output_path).inline_shapes))

    def test_demo_fixture_injects_three_comparable_sale_pages(self):
        project_root = Path(__file__).resolve().parents[1]
        assignment = project_root / "tests" / "fixtures" / "DEMO-001"
        workbook_path = assignment / "workbook.xlsx"
        comps = load_comp_data(workbook_path)
        self.assertEqual(
            ["Sale No. 1", "Sale No. 2", "Sale No. 3"],
            [comp["COMP_NO"] for comp in comps],
        )

        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "comp-pages.docx"
            document = Document()
            document.add_paragraph("Comparable Sales")
            document.add_paragraph("[[COMP_SHEETS_BLOCK]]")
            document.save(output_path)

            count = inject_comp_section(
                output_path,
                project_root / "templates" / "comp_block_template.docx",
                workbook_path,
            )
            self.assertEqual(3, count)
            placeholders = find_docx_placeholders(output_path)
            self.assertNotIn("COMP_SHEETS_BLOCK", placeholders)
            self.assertFalse(
                [name for name in placeholders if name.startswith("COMP_")]
            )
            output_text = "\n".join(
                node.text or ""
                for node in Document(output_path).element.body.iter(qn("w:t"))
            )
            for sale_no in ("Sale No. 1", "Sale No. 2", "Sale No. 3"):
                self.assertIn(sale_no, output_text)

    def test_ownership_table_uses_existing_assignment_fields(self):
        variables = {
            "OWNER_NAME": "Example Owner, LLC",
            "PRIOR_SALE_DATE": "No arm's-length transfers in three years.",
            "PRIOR_SALE_PRICE": "N/A",
        }
        with tempfile.TemporaryDirectory() as temp_dir:
            assignment, templates, config = _build_assignment(
                temp_dir,
                "[[OWNERSHIP_HISTORY_TABLE]]",
                variables,
            )

            result = validate_assignment(assignment, templates, config)
            self.assertTrue(result["ready"])
            self.assertEqual(
                ["OWNERSHIP_HISTORY_TABLE"],
                result["handled_blocks"],
            )

            output_path = Path(temp_dir) / "ownership-output.docx"
            shutil.copy(templates / "report.docx", output_path)
            self.assertTrue(inject_ownership_history(output_path, variables))
            self.assertNotIn(
                "OWNERSHIP_HISTORY_TABLE",
                find_docx_placeholders(output_path),
            )
            output_doc = Document(output_path)
            self.assertEqual("Owner of Record", output_doc.tables[0].cell(0, 0).text)
            self.assertEqual("Example Owner, LLC", output_doc.tables[0].cell(0, 1).text)

    def test_narrative_model_uses_nested_command_routing(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            config_path = Path(temp_dir) / "config.json"
            config_path.write_text(
                json.dumps(
                    {
                        "models": {
                            "default": "default-model",
                            "per_command": {
                                "draft": "draft-model",
                                "reconcile": "reconcile-model",
                            },
                        }
                    }
                ),
                encoding="utf-8",
            )

            self.assertEqual(
                "draft-model",
                _get_model(config_path, command="draft"),
            )
            self.assertEqual(
                "reconcile-model",
                _get_model(config_path, command="reconcile"),
            )
            self.assertEqual(
                "default-model",
                _get_model(config_path, command="unconfigured"),
            )

    def test_committed_field_registry_matches_baseline_sources(self):
        project_root = Path(__file__).resolve().parents[1]
        with open(project_root / "config.json", encoding="utf-8") as config_file:
            config = json.load(config_file)

        built = build_registry(
            workbook_path=project_root / config["workbook_template"],
            templates_dir=project_root / config["templates_dir"],
            stages=config["stages"],
            fixture_json_path=(
                project_root
                / "tests"
                / "fixtures"
                / "DEMO-001"
                / "DEMO-001_variables.json"
            ),
        )
        committed = load_registry(project_root / config["field_registry"])
        self.assertTrue(set(built["fields"]).issubset(committed["fields"]))
        self.assertTrue(set(built["blocks"]).issubset(committed["blocks"]))

        template_paths = [
            project_root / config["templates_dir"] / document["template"]
            for stage in config["stages"].values()
            for document in stage["documents"]
        ]
        with open(
            project_root
            / "tests"
            / "fixtures"
            / "DEMO-001"
            / "DEMO-001_variables.json",
            encoding="utf-8",
        ) as fixture_file:
            fixture_variables = json.load(fixture_file)
        audit = audit_assignment_contract(
            project_root / config["field_registry"],
            project_root / config["workbook_template"],
            template_paths,
            fixture_variables,
        )
        self.assertEqual([], audit["errors"])
        self.assertEqual([], audit["warnings"])


if __name__ == "__main__":
    unittest.main()
