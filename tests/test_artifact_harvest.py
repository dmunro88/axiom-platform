import base64
import json
import shutil
import sqlite3
import tempfile
import unittest
from pathlib import Path

from docx import Document
import openpyxl
from openpyxl.chart import LineChart, Reference

from comparable_contract import confirm_extraction_result
from db import (
    comp_photo_artifacts,
    get_conn,
    init_db,
    insert_comp,
    insert_manual_comp_photo,
    insert_property,
    insert_source_document,
    search_source_artifacts,
)
from harvest_contract import ARTIFACT_CONTRACT_ID
from ingest import commit_extraction_result, run_extraction


TINY_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk"
    "/x8AAusB9Wl2nWQAAAAASUVORK5CYII="
)


def _build_artifact_assignment(root, file_no="25C980"):
    assignment = (
        root
        / "historical"
        / f"{file_no} - Office - 321 Fictional Artifact Road, Demo City"
    )
    assignment.mkdir(parents=True)
    embedded_source = root / "embedded-map.png"
    embedded_source.write_bytes(TINY_PNG)

    document = Document()
    document.add_paragraph("Subject Property: 321 Fictional Artifact Road")
    document.add_paragraph("Effective Date: April 30, 2025")
    document.add_paragraph("Reconciled Value: $1,500,000")
    drawing = document.add_paragraph().add_run().add_picture(
        str(embedded_source)
    )
    drawing._inline.docPr.set("descr", "Fictional Location Map")
    document.save(assignment / f"REPORT {file_no}.docx")

    subject_photos = assignment / "assets" / "photos" / "subject"
    subject_photos.mkdir(parents=True)
    primary_photo = subject_photos / "Fictional Subject Photo.png"
    primary_photo.write_bytes(TINY_PNG)
    duplicate_dir = assignment / "pic"
    duplicate_dir.mkdir()
    shutil.copyfile(primary_photo, duplicate_dir / "Duplicate Photo.png")

    exhibits = assignment / "exhibits"
    exhibits.mkdir()
    (exhibits / "Fictional Zoning Map.pdf").write_bytes(
        b"%PDF-1.4\n% fictional QA artifact\n%%EOF\n"
    )
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Data"
    sheet.append(["Period", "Vacancy"])
    sheet.append(["2024", 0.08])
    sheet.append(["2025", 0.07])
    chart = LineChart()
    chart.title = "Fictional Vacancy Trend"
    chart.add_data(
        Reference(sheet, min_col=2, min_row=1, max_row=3),
        titles_from_data=True,
    )
    chart.set_categories(Reference(sheet, min_col=1, min_row=2, max_row=3))
    sheet.add_chart(chart, "D2")
    workbook.save(assignment / "Fictional Market Chart.xlsx")
    workbook.close()
    return assignment, primary_photo


class ArtifactHarvestTests(unittest.TestCase):
    def test_external_and_embedded_artifacts_end_to_end(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            assignment, _ = _build_artifact_assignment(root)
            staged_path = run_extraction(
                root / "historical",
                staged_dir=root / "staged",
            )[0]
            staged = json.loads(staged_path.read_text(encoding="utf-8"))
            artifacts = staged["artifacts"]

            self.assertEqual(4, len(artifacts))
            self.assertEqual(
                ["chart", "map", "map", "photo"],
                sorted(record["data"]["artifact_kind"] for record in artifacts),
            )
            photo = next(
                record
                for record in artifacts
                if record["data"]["artifact_kind"] == "photo"
            )
            embedded = next(
                record
                for record in artifacts
                if record["data"]["title"] == "Fictional Location Map"
            )
            self.assertEqual(ARTIFACT_CONTRACT_ID, photo["contract_id"])
            self.assertEqual(1, len(photo["alternate_provenance"]))
            self.assertEqual(64, len(photo["data"]["artifact_sha256"]))
            self.assertEqual(1, photo["data"]["width_px"])
            self.assertEqual(1, photo["data"]["height_px"])
            self.assertIn(
                "word/document.xml:drawing:1:word/media/",
                embedded["provenance"]["source_locator"],
            )
            self.assertEqual(
                f"REPORT {assignment.name.split(' - ')[0]}.docx",
                embedded["data"]["container_filename"],
            )
            self.assertEqual("2025-04-30", embedded["data"]["effective_date"])
            self.assertEqual("Demo City", embedded["data"]["geography"])
            native_chart = next(
                record
                for record in artifacts
                if record["data"]["artifact_kind"] == "chart"
            )
            self.assertIn(
                "xl/charts/chart1.xml",
                native_chart["provenance"]["source_locator"],
            )
            self.assertIn("Fictional Vacancy Trend", native_chart["data"]["title"])

            embedded["data"]["description"] = (
                "Reviewed fictional location-map exhibit."
            )
            embedded["review_edits"] = [{
                "field": "description",
                "before": None,
                "after": "Reviewed fictional location-map exhibit.",
            }]
            confirmed = confirm_extraction_result(
                staged,
                reviewer="fictional-artifact-qa",
                reviewed_at="2026-07-01T19:00:00+00:00",
            )

            db_path = root / "artifacts.db"
            init_db(db_path, quiet=True)
            connection = get_conn(db_path)
            try:
                with connection:
                    counts = commit_extraction_result(confirmed, connection)
                self.assertEqual(4, counts["source_artifacts"])
                with connection:
                    duplicates = commit_extraction_result(confirmed, connection)
                self.assertEqual(
                    4,
                    duplicates["duplicate_source_artifacts"],
                )
            finally:
                connection.close()

            records = search_source_artifacts(
                db_path,
                artifact_kind="map",
                title_contains="Location",
                geography="Demo",
                property_type="Office",
            )
            self.assertEqual(1, len(records))
            self.assertEqual("confirmed", records[0]["review_status"])
            self.assertEqual(
                "Reviewed fictional location-map exhibit.",
                records[0]["description"],
            )
            self.assertTrue(records[0]["source_path"].endswith(".docx"))

    def test_changed_external_artifact_requires_reextract(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            _, primary_photo = _build_artifact_assignment(
                root,
                file_no="25C981",
            )
            staged_path = run_extraction(
                root / "historical",
                staged_dir=root / "staged",
            )[0]
            staged = json.loads(staged_path.read_text(encoding="utf-8"))
            confirmed = confirm_extraction_result(staged, reviewer="qa")
            primary_photo.write_bytes(TINY_PNG + b"changed")

            db_path = root / "changed.db"
            init_db(db_path, quiet=True)
            connection = get_conn(db_path)
            try:
                with self.assertRaisesRegex(
                    ValueError,
                    "Source changed after extraction",
                ):
                    with connection:
                        commit_extraction_result(confirmed, connection)
                self.assertEqual(
                    0,
                    connection.execute(
                        "SELECT count(*) FROM source_artifacts"
                    ).fetchone()[0],
                )
            finally:
                connection.close()

    def test_unconfirmed_artifact_rolls_back_batch(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            _build_artifact_assignment(root, file_no="25C982")
            staged_path = run_extraction(
                root / "historical",
                staged_dir=root / "staged",
            )[0]
            staged = json.loads(staged_path.read_text(encoding="utf-8"))
            confirmed = confirm_extraction_result(staged, reviewer="qa")
            confirmed["artifacts"][0]["review"]["status"] = "unreviewed"

            db_path = root / "rollback.db"
            init_db(db_path, quiet=True)
            connection = get_conn(db_path)
            try:
                with self.assertRaisesRegex(ValueError, "not been confirmed"):
                    with connection:
                        commit_extraction_result(confirmed, connection)
                for table in (
                    "assignments",
                    "source_artifacts",
                    "source_documents",
                ):
                    self.assertEqual(
                        0,
                        connection.execute(
                            f"SELECT count(*) FROM {table}"
                        ).fetchone()[0],
                    )
            finally:
                connection.close()

    def test_contract_descriptor_and_artifact_table(self):
        schema_path = (
            Path(__file__).resolve().parents[1]
            / "schemas"
            / "historical_harvest.v1.json"
        )
        schema = json.loads(schema_path.read_text(encoding="utf-8"))
        self.assertEqual(
            ARTIFACT_CONTRACT_ID,
            schema["contracts"]["artifact"],
        )
        with tempfile.TemporaryDirectory() as temp_dir:
            db_path = Path(temp_dir) / "fresh.db"
            init_db(db_path, quiet=True)
            connection = get_conn(db_path)
            try:
                columns = {
                    row[1]
                    for row in connection.execute(
                        "PRAGMA table_info(source_artifacts)"
                    )
                }
                self.assertIn("artifact_sha256", columns)
                self.assertIn("source_record_json", columns)
                self.assertIn("comp_id", columns)
                self.assertIn("lease_comp_id", columns)
            finally:
                connection.close()

    def test_legacy_artifact_table_adds_comp_link_columns(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            db_path = Path(temp_dir) / "legacy-artifacts.db"
            connection = sqlite3.connect(db_path)
            connection.executescript(
                """
                CREATE TABLE source_artifacts (
                    artifact_id INTEGER PRIMARY KEY,
                    artifact_kind TEXT,
                    artifact_sha256 TEXT,
                    identity_key TEXT,
                    review_status TEXT
                );
                """
            )
            connection.commit()
            connection.close()

            init_db(db_path, quiet=True)
            connection = sqlite3.connect(db_path)
            try:
                columns = {
                    row[1]
                    for row in connection.execute(
                        "PRAGMA table_info(source_artifacts)"
                    )
                }
                self.assertIn("comp_id", columns)
                self.assertIn("lease_comp_id", columns)
            finally:
                connection.close()

    def test_manual_comp_photo_attachment_links_artifact_to_sale_comp(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            db_path = root / "manual-photo.db"
            photo_path = root / "fictional-comp-photo.png"
            photo_path.write_bytes(TINY_PNG)

            init_db(db_path, quiet=True)
            connection = get_conn(db_path)
            try:
                property_id = insert_property(
                    {
                        "address_street": "700 Fictional Photo Way",
                        "address_city": "Demo City",
                        "property_type": "Office",
                    },
                    connection,
                )
                source_doc_id = insert_source_document(
                    root / "Fictional Source.docx",
                    "report",
                    conn=connection,
                )
                comp_id = insert_comp(
                    {
                        "sale_price": 1_250_000,
                        "sale_date": "2025-05-01",
                    },
                    property_id,
                    source_doc_id,
                    {"sale_price": "high"},
                    connection,
                    identity_key="fictional-sale-photo-test",
                    review={"status": "confirmed"},
                )
                first_artifact_id = insert_manual_comp_photo(
                    "sale",
                    comp_id,
                    photo_path,
                    connection,
                    title="Front elevation",
                    original_filename="upload.png",
                )
                second_artifact_id = insert_manual_comp_photo(
                    "sale",
                    comp_id,
                    photo_path,
                    connection,
                    title="Duplicate upload title should not matter",
                    original_filename="upload-again.png",
                )
                connection.commit()
                self.assertEqual(first_artifact_id, second_artifact_id)
            finally:
                connection.close()

            photos = comp_photo_artifacts(
                db_path,
                record_kind="sale",
                record_id=comp_id,
            )
            self.assertEqual(1, len(photos))
            self.assertEqual(comp_id, photos[0]["comp_id"])
            self.assertEqual("confirmed", photos[0]["review_status"])
            self.assertEqual("Front elevation", photos[0]["title"])
            self.assertEqual(str(photo_path), photos[0]["artifact_filename"])
            self.assertEqual(1, photos[0]["width_px"])
            self.assertEqual(1, photos[0]["height_px"])

            search_records = search_source_artifacts(db_path, comp_id=comp_id)
            self.assertEqual(1, len(search_records))
            self.assertEqual(first_artifact_id, search_records[0]["artifact_id"])


if __name__ == "__main__":
    unittest.main()
