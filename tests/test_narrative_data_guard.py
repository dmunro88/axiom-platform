"""
tests/test_narrative_data_guard.py — Axiom Platform

Regression coverage for narrative_generator.py's pre-flight data sanity
checks, added 2026-07-10 after live-testing Phase 7 against DEMO-001 showed
that broken workbook data (unresolved Excel formula errors, or a concluded
value of $0/negative) caused the AI model to correctly refuse to fabricate
numbers -- but its refusal/meta-commentary text was getting injected into
the document verbatim instead of a clean placeholder. These tests lock in
that the bad-data cases are caught *before* any API call is made.
"""

import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import narrative_generator as ng


class HasErrorTokenTests(unittest.TestCase):
    def test_detects_each_known_error_token(self):
        for token in ["#DIV/0!", "#NUM!", "#VALUE!", "#REF!", "#NAME?", "#NULL!", "#N/A"]:
            with self.subTest(token=token):
                self.assertTrue(ng._has_error_token(token))
                self.assertTrue(ng._has_error_token(f"$ {token} per SF"))

    def test_clean_values_are_not_flagged(self):
        self.assertFalse(ng._has_error_token("$91.26"))
        self.assertFalse(ng._has_error_token(""))
        self.assertFalse(ng._has_error_token(None))


class ParseMoneyTests(unittest.TestCase):
    def test_parses_plain_and_formatted_currency(self):
        self.assertEqual(ng._parse_money("1234.5"), 1234.5)
        self.assertEqual(ng._parse_money("$1,234,567"), 1234567.0)
        self.assertEqual(ng._parse_money("$0"), 0.0)

    def test_parses_negative_forms(self):
        self.assertEqual(ng._parse_money("-1600000"), -1600000.0)
        self.assertEqual(ng._parse_money("($1,600,000)"), -1600000.0)

    def test_non_numeric_returns_none(self):
        self.assertIsNone(ng._parse_money(""))
        self.assertIsNone(ng._parse_money(None))
        self.assertIsNone(ng._parse_money("Demo North submarket"))


class FieldsDataIssueTests(unittest.TestCase):
    def test_flags_error_token_field(self):
        v = {"SCA_VALUE": "#VALUE!"}
        issue = ng._fields_data_issue(v, ["SCA_VALUE"])
        self.assertIsNotNone(issue)
        self.assertIn("SCA_VALUE", issue)

    def test_flags_zero_value_field(self):
        v = {"SCA_VALUE": "$0"}
        issue = ng._fields_data_issue(v, ["SCA_VALUE"])
        self.assertIsNotNone(issue)

    def test_flags_negative_value_field(self):
        v = {"IA_VALUE": "-$1,600,000"}
        issue = ng._fields_data_issue(v, ["IA_VALUE"])
        self.assertIsNotNone(issue)

    def test_clean_fields_pass(self):
        v = {
            "SCA_ADJ_NARROW_LOW": "$71.20", "SCA_ADJ_NARROW_HIGH": "$92.40",
            "SCA_ADJ_UNIT_MEAN": "$81.50", "SCA_ADJ_UNIT_MEDIAN": "$80.00",
            "SCA_VALUE": "$1,850,000",
        }
        issue = ng._fields_data_issue(v, list(v.keys()))
        self.assertIsNone(issue)


class ReconciliationDataIssueTests(unittest.TestCase):
    def test_only_checks_developed_approaches(self):
        # Income approach not developed -> its garbage value should be ignored.
        v = {
            "SCA_DEVELOPED": "Yes", "IA_DEVELOPED": "No", "CA_DEVELOPED": "No",
            "SCA_VALUE": "$1,850,000",
            "IA_VALUE": "-$1,600,000",  # not developed, must not be checked
            "VALUE_CONCLUSION": "$1,850,000",
        }
        self.assertIsNone(ng._reconciliation_data_issue(v))

    def test_flags_developed_approach_with_bad_value(self):
        v = {
            "SCA_DEVELOPED": "Yes", "IA_DEVELOPED": "Yes", "CA_DEVELOPED": "No",
            "SCA_VALUE": "$0",
            "IA_VALUE": "-$1,600,000",
            "VALUE_CONCLUSION": "$0",
        }
        issue = ng._reconciliation_data_issue(v)
        self.assertIsNotNone(issue)

    def test_flags_bad_final_value_conclusion_even_if_approaches_look_fine(self):
        v = {
            "SCA_DEVELOPED": "Yes", "IA_DEVELOPED": "No", "CA_DEVELOPED": "No",
            "SCA_VALUE": "$1,850,000",
            "VALUE_CONCLUSION": "#VALUE!",
        }
        issue = ng._reconciliation_data_issue(v)
        self.assertIsNotNone(issue)


class DataIssueForDispatchTests(unittest.TestCase):
    def test_dispatches_reconciliation_key(self):
        v = {"SCA_DEVELOPED": "No", "IA_DEVELOPED": "No", "CA_DEVELOPED": "No",
             "VALUE_CONCLUSION": "$0"}
        issue = ng._data_issue_for("RECONCILIATION_NARRATIVE", v)
        self.assertIsNotNone(issue)

    def test_dispatches_critical_fields_key(self):
        v = {"SCA_VALUE": "#VALUE!"}
        issue = ng._data_issue_for("SCA_CONCLUSION_NARRATIVE", v)
        self.assertIsNotNone(issue)

    def test_keys_without_a_dedicated_check_return_none(self):
        # MARKET_AREA_OVERVIEW has no numeric-value guard -- always proceeds
        # to the normal API call path.
        self.assertIsNone(ng._data_issue_for("MARKET_AREA_OVERVIEW", {}))


class InjectAllNarrativesSkipsApiOnBadDataTests(unittest.TestCase):
    """End-to-end: a document containing only a broken-data narrative
    placeholder must get a clean placeholder note injected, and must never
    call the Claude API at all."""

    def test_broken_sca_conclusion_data_skips_api_call(self):
        from docx import Document

        doc = Document()
        doc.add_paragraph("[[SCA_CONCLUSION_NARRATIVE]]")
        tmp_doc = Path(tempfile.gettempdir()) / "_test_narrative_guard.docx"
        doc.save(str(tmp_doc))

        variables = {
            "SCA_ADJ_NARROW_LOW": "#VALUE!",
            "SCA_ADJ_NARROW_HIGH": "#VALUE!",
            "SCA_ADJ_UNIT_MEAN": "#VALUE!",
            "SCA_ADJ_UNIT_MEDIAN": "#VALUE!",
            "SCA_VALUE": "$0",
        }

        with patch.object(ng, "_call_claude") as mock_call:
            results = ng.inject_all_narratives(
                tmp_doc, tmp_doc, variables, config_path=None
            )

        mock_call.assert_not_called()
        self.assertIn("SCA_CONCLUSION_NARRATIVE", results)
        self.assertTrue(results["SCA_CONCLUSION_NARRATIVE"].startswith("skipped (data issue:"))

        result_doc = Document(str(tmp_doc))
        full_text = "\n".join(p.text for p in result_doc.paragraphs)
        self.assertNotIn("[[SCA_CONCLUSION_NARRATIVE]]", full_text)
        self.assertIn("Pending", full_text)
        tmp_doc.unlink()


if __name__ == "__main__":
    unittest.main()
