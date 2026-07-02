"""Discover, validate, and inject assignment media into Word report blocks."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.image.image import Image
from docx.shared import Inches


IMAGE_SUFFIXES = frozenset({".jpg", ".jpeg", ".png"})

SINGLE_MEDIA_BLOCKS = {
    "REGIONAL_MAP_IMAGE": "assets/maps/regional",
    "AERIAL_MAP_IMAGE": "assets/maps/aerial",
    "PARCEL_MAP_IMAGE": "assets/maps/parcel",
    "SCA_SALE_LOCATION_MAP": "assets/maps/sca-sale-location",
    "LEASE_COMP_LOCATION_MAP": "assets/maps/lease-comp-location",
    "BUILDING_SKETCH_BLOCK": "assets/building-sketch",
}

MULTI_MEDIA_BLOCKS = {
    "SUBJECT_PHOTOS_BLOCK": "assets/photos/subject",
    "LEASE_COMP_PHOTOS_BLOCK": "assets/photos/lease-comps",
}

MEDIA_BLOCKS = frozenset(SINGLE_MEDIA_BLOCKS) | frozenset(MULTI_MEDIA_BLOCKS)


def create_media_directories(assignment_dir):
    """Create the standard assignment media folders."""
    assignment_dir = Path(assignment_dir)
    for relative in (
        "assets/maps",
        "assets/photos/subject",
        "assets/photos/lease-comps",
    ):
        (assignment_dir / relative).mkdir(parents=True, exist_ok=True)


def _image_files(directory):
    directory = Path(directory)
    if not directory.is_dir():
        return []
    return sorted(
        path
        for path in directory.iterdir()
        if path.is_file() and path.suffix.lower() in IMAGE_SUFFIXES
    )


def media_files_for_block(block, assignment_dir):
    """Return the conventionally named image files available for *block*."""
    assignment_dir = Path(assignment_dir)

    if block in SINGLE_MEDIA_BLOCKS:
        base_path = assignment_dir / SINGLE_MEDIA_BLOCKS[block]
        return [
            candidate
            for candidate in sorted(base_path.parent.glob(f"{base_path.name}.*"))
            if candidate.is_file() and candidate.suffix.lower() in IMAGE_SUFFIXES
        ][:1]

    if block in MULTI_MEDIA_BLOCKS:
        return _image_files(assignment_dir / MULTI_MEDIA_BLOCKS[block])

    return []


def missing_media_reason(block):
    """Return an actionable description of where media should be placed."""
    if block in SINGLE_MEDIA_BLOCKS:
        relative = SINGLE_MEDIA_BLOCKS[block]
        return f"Add {relative}.jpg or .png to the assignment before delivery."
    if block in MULTI_MEDIA_BLOCKS:
        relative = MULTI_MEDIA_BLOCKS[block]
        return (
            f"Add one or more .jpg or .png files under {relative}/ "
            "before delivery."
        )
    return "No media convention is registered for this block."


def _iter_document_paragraphs(doc):
    def walk_container(container):
        yield from container.paragraphs
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from walk_container(cell)

    yield from walk_container(doc)
    for section in doc.sections:
        yield from walk_container(section.header)
        yield from walk_container(section.footer)


def _clear_runs(paragraph):
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)


def _fitted_dimensions(image_path, max_width, max_height):
    image = Image.from_file(str(image_path))
    scale = min(max_width / image.width, max_height / image.height)
    return int(image.width * scale), int(image.height * scale)


def _inject_files(paragraph, marker, files, multi_image):
    full_text = "".join(run.text for run in paragraph.runs)
    before, after = full_text.split(marker, 1)
    _clear_runs(paragraph)

    if before:
        paragraph.add_run(before.rstrip())
        paragraph.add_run().add_break()

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    max_width = Inches(3.0 if multi_image else 6.5)
    max_height = Inches(3.5 if multi_image else 7.0)

    for index, image_path in enumerate(files):
        width, height = _fitted_dimensions(image_path, max_width, max_height)
        paragraph.add_run().add_picture(
            str(image_path),
            width=width,
            height=height,
        )
        if index < len(files) - 1:
            paragraph.add_run().add_break()

    if after:
        paragraph.add_run().add_break()
        paragraph.add_run(after.lstrip())


def inject_media_blocks(doc_path, assignment_dir):
    """
    Replace supported media placeholders in *doc_path*.

    Returns ``{block: count}`` for injected blocks. Missing assets are left
    visible so the post-generation placeholder scan can block final delivery.
    """
    doc_path = Path(doc_path)
    doc = Document(str(doc_path))
    injected = {}

    for paragraph in _iter_document_paragraphs(doc):
        text = "".join(run.text for run in paragraph.runs)
        if "[[" not in text:
            continue

        for block in MEDIA_BLOCKS:
            marker = f"[[{block}]]"
            if marker not in text:
                continue
            files = media_files_for_block(block, assignment_dir)
            if not files:
                continue
            _inject_files(
                paragraph,
                marker,
                files,
                multi_image=block in MULTI_MEDIA_BLOCKS,
            )
            injected[block] = len(files)
            break

    if injected:
        doc.save(str(doc_path))
    return injected
