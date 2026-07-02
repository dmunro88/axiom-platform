"""
narrative_generator.py — Axiom Commercial Appraisal Platform
=============================================================
AI-powered narrative generator. Handles every [[X_NARRATIVE]] and
[[X_OVERVIEW]] placeholder in the appraisal report template.

Public interface
----------------
    inject_all_narratives(doc_path, workbook_path, variables, config_path=None)
        Scan a filled Word document for known narrative placeholders and replace
        each one with AI-generated USPAP-compliant prose.  Call this after
        fill_engine and comp_builder have already run.

    generate_adjustment_narrative(workbook_path, variables, config_path=None)
        Legacy entry point kept for backward compatibility.
        Generates only the LAND_ADJUSTMENT_NARRATIVE from the workbook land tab.

Supported placeholders
----------------------
    [[INSPECTION_NARRATIVE]]     — inspection date, duration, observations
    [[MARKET_AREA_OVERVIEW]]     — regional/submarket context
    [[SCA_APPROACH_NARRATIVE]]   — methodology, comp search, unit of comparison
    [[SCA_ADJUSTMENT_NARRATIVE]] — adjustment categories, post-adj range
    [[SCA_CONCLUSION_NARRATIVE]] — reconciliation of adjusted indications
    [[CAP_RATE_NARRATIVE]]       — cap rate derivation and support
    [[ENCUMBRANCES_NARRATIVE]]   — property rights, encumbrances, sale history
    [[RECONCILIATION_NARRATIVE]] — final approach weighting and value conclusion
    [[LAND_ADJUSTMENT_NARRATIVE]]— land comp adjustment summary (workbook-driven)
"""

import copy
import json
import os
from pathlib import Path


# ── Model selection ────────────────────────────────────────────────────────────

DEFAULT_MODEL = "claude-sonnet-4-6"


def _get_model(config_path=None, command="draft"):
    """Resolve a command model from the top-level ``models`` configuration."""
    if config_path is None:
        config_path = Path(__file__).parent / "config.json"
    try:
        with open(config_path, encoding="utf-8") as f:
            cfg = json.load(f)
        models = cfg.get("models", cfg)
        return models.get("per_command", {}).get(
            command,
            models.get("default", DEFAULT_MODEL),
        )
    except Exception:
        return DEFAULT_MODEL


# ── Claude API caller ─────────────────────────────────────────────────────────

def _call_claude(prompt, model, max_tokens=900):
    try:
        import anthropic
    except ImportError:
        raise RuntimeError("pip install anthropic  |  set ANTHROPIC_API_KEY")

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        raise RuntimeError("ANTHROPIC_API_KEY environment variable not set.")

    client = anthropic.Anthropic(api_key=api_key)
    msg = client.messages.create(
        model=model,
        max_tokens=max_tokens,
        messages=[{"role": "user", "content": prompt}]
    )
    return msg.content[0].text.strip()


# ── Paragraph injector ────────────────────────────────────────────────────────

def _inject_narrative_into_doc(doc, placeholder, narrative_text):
    """
    Find the paragraph in *doc* that contains *placeholder*, replace it with
    *narrative_text*.  Multi-paragraph text (separated by blank lines) is
    expanded into multiple Word paragraphs, each inheriting the original
    paragraph's style.

    Returns True if the placeholder was found and replaced, False otherwise.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tag = f"[[{placeholder}]]"

    for para in doc.paragraphs:
        if tag not in para.text:
            continue

        pPr = para._p.find(qn("w:pPr"))
        chunks = [c.strip() for c in narrative_text.split("\n\n") if c.strip()]

        # Replace text in the first paragraph
        for run in para.runs:
            run.text = run.text.replace(tag, "")
        if para.runs:
            para.runs[0].text = chunks[0]
        else:
            para.add_run(chunks[0])

        # Insert remaining paragraphs after the anchor
        insert_after = para._p
        for chunk in chunks[1:]:
            new_p = OxmlElement("w:p")
            if pPr is not None:
                new_p.append(copy.deepcopy(pPr))
            new_r = OxmlElement("w:r")
            new_t = OxmlElement("w:t")
            new_t.text = chunk
            new_r.append(new_t)
            new_p.append(new_r)
            insert_after.addnext(new_p)
            insert_after = new_p

        return True

    return False


# ─────────────────────────────────────────────────────────────────────────────
# Individual prompt builders
# Each returns a (prompt_str, max_tokens) tuple.
# ─────────────────────────────────────────────────────────────────────────────

def _prompt_inspection(v):
    return (f"""You are writing the Inspection narrative section of a commercial real estate appraisal report. Write in formal, USPAP-compliant appraisal language. Active voice, no bullet points.

Subject property details:
  Address:           {v.get('PROPERTY_ADDRESS', '')} {v.get('PROPERTY_CITY', '')} {v.get('PROPERTY_STATE', '')} {v.get('PROPERTY_ZIP', '')}
  Property type:     {v.get('PROPERTY_SUBTYPE_FULL', 'commercial property')}
  GBA (sq ft):       {v.get('GBA', '')}
  Number of units:   {v.get('UNIT_COUNT', '')}
  Year built:        {v.get('YEAR_BUILT', '')}
  Inspection date:   {v.get('INSPECTION_DATE', '')}
  Inspection duration: {v.get('INSPECTION_DURATION', '')}
  General condition: {v.get('SUBJECT_CONDITION', '')}
  Appraiser name:    {v.get('APPRAISER_NAME', 'the appraiser')}

Write exactly TWO paragraphs:

1. INSPECTION STATEMENT: State the date the property was personally inspected, the duration, and that an interior and exterior inspection was performed. Note that the inspection was conducted for the purpose of developing the appraisal. Do not use "I" — use "the appraiser" or passive voice.

2. PROPERTY OBSERVATIONS: Briefly describe what was observed during the inspection — general condition, general state of improvements, and any notable features or limitations observed. Keep it factual and objective. Two to three sentences.

Output only the two paragraphs, no headers, no commentary.""", 600)


def _prompt_market_overview(v):
    return (f"""You are writing the Market Area Overview narrative for a commercial real estate appraisal report. Write in formal, USPAP-compliant appraisal language. Active voice, no bullet points.

Subject property details:
  City/State:        {v.get('PROPERTY_CITY', '')} {v.get('PROPERTY_STATE', '')}
  County:            {v.get('PROPERTY_COUNTY', '')} County
  Submarket:         {v.get('SUBMARKET', '')}
  Property type:     {v.get('PROPERTY_SUBTYPE_FULL', 'commercial property')}
  Adjacent uses:     {v.get('ADJACENT_USES', '')}
  Zoning:            {v.get('ZONING_CLASS', '')} ({v.get('ZONING_CODE', '')})
  Permitted uses:    {v.get('PERMITTED_USES', '')}
  Subject address:   {v.get('PROPERTY_ADDRESS', '')}
  Frontage street:   {v.get('FRONTAGE_STREET', '')}
  Subject access:    {v.get('SUBJECT_ACCESS', '')}
  Subject visibility:{v.get('SUBJECT_VISIBILITY', '')}

Write exactly THREE paragraphs:

1. REGIONAL CONTEXT: Describe the city and county in general terms — location within the state, general economic character, regional significance. Two to three sentences.

2. SUBMARKET/NEIGHBORHOOD: Describe the immediate market area or submarket around the subject property — character of the area, nearby land uses, access and visibility of the location. Reference the submarket name. Three to four sentences.

3. MARKET CONDITIONS OVERVIEW: Describe the general commercial market conditions for this property type in this submarket — demand drivers, typical occupancy, general investor appetite. Two to three sentences. Do not state specific vacancy or cap rate numbers (those appear elsewhere in the report).

Output only the three paragraphs, no headers, no commentary.""", 700)


def _prompt_sca_approach(v):
    return (f"""You are writing the introductory methodology paragraph(s) for the Sales Comparison Approach section of a commercial real estate appraisal report. Write in formal, USPAP-compliant language. Active voice, no bullet points.

Subject property details:
  Property type:     {v.get('PROPERTY_SUBTYPE_FULL', 'commercial property')}
  GBA (sq ft):       {v.get('GBA', '')}
  Location:          {v.get('PROPERTY_CITY', '')} {v.get('PROPERTY_STATE', '')} / {v.get('SUBMARKET', '')} submarket
  Unit of comparison:{v.get('UNIT_OF_COMPARISON', 'price per square foot of gross building area')}
  Comp size range:   {v.get('SCA_COMP_SIZE_LOW', '')} – {v.get('SCA_COMP_SIZE_HIGH', '')} SF (mean {v.get('SCA_COMP_SIZE_MEAN', '')}, median {v.get('SCA_COMP_SIZE_MEDIAN', '')} SF)
  Comp unadj range:  {v.get('SCA_COMP_UNIT_LOW', '')} – {v.get('SCA_COMP_UNIT_HIGH', '')} per SF (mean {v.get('SCA_COMP_UNIT_MEAN', '')}, median {v.get('SCA_COMP_UNIT_MEDIAN', '')} per SF)
  Effective date:    {v.get('EFFECTIVE_DATE', '')}

Write exactly TWO paragraphs:

1. METHODOLOGY: Describe the Sales Comparison Approach methodology — that it estimates value by comparing the subject to recently sold properties with similar characteristics, using the stated unit of comparison, and that differences between the subject and each comparable are addressed through adjustments. Note the geographic and time scope of the search (recent sales in the market area).

2. COMPARABLE SUMMARY: Describe the unadjusted range and central tendency of the comparables selected, referencing the size range and price range. Note that the comparables selected are considered the most relevant to the subject based on location, physical characteristics, and market conditions.

Output only the two paragraphs, no headers, no commentary.""", 500)


def _prompt_sca_adjustment(v):
    return (f"""You are writing the Adjustment Summary narrative for the Sales Comparison Approach section of a commercial real estate appraisal report. Write in formal, USPAP-compliant appraisal language. Active voice, no bullet points.

Subject and comparable data:
  Property type:       {v.get('PROPERTY_SUBTYPE_FULL', 'commercial property')}
  Submarket:           {v.get('SUBMARKET', '')}
  Unit of comparison:  {v.get('UNIT_OF_COMPARISON', 'price per square foot')}
  Unadjusted range:    {v.get('SCA_COMP_UNIT_LOW', '')} – {v.get('SCA_COMP_UNIT_HIGH', '')} per SF
  Adjusted unit range: {v.get('SCA_ADJ_UNIT_LOW', '')} – {v.get('SCA_ADJ_UNIT_HIGH', '')} per SF
  Adjusted unit mean:  {v.get('SCA_ADJ_UNIT_MEAN', '')} per SF
  Adjusted unit median:{v.get('SCA_ADJ_UNIT_MEDIAN', '')} per SF
  Narrowed adj range:  {v.get('SCA_ADJ_NARROW_LOW', '')} – {v.get('SCA_ADJ_NARROW_HIGH', '')} per SF
  Narrowed mean:       {v.get('SCA_ADJ_NARROW_MEAN', '')} per SF
  Narrowed median:     {v.get('SCA_ADJ_NARROW_MEDIAN', '')} per SF
  Qualitative summary: {v.get('SCA_QUAL_CONCLUSION', '')}
  Effective date:      {v.get('EFFECTIVE_DATE', '')}

Write exactly THREE paragraphs:

1. ADJUSTMENT OVERVIEW: Describe the adjustment categories applied — market conditions (time), location, size/physical characteristics, and condition. Explain that adjustments were applied to each comparable based on its characteristics relative to the subject. Do not list specific dollar amounts or percentages.

2. POST-ADJUSTMENT ANALYSIS: Describe the narrowing effect of adjustments — the range tightened from the unadjusted spread to the adjusted range. Reference the adjusted mean and median as central tendency indicators. Note which comparables fall near the central tendency.

3. QUALITATIVE SUMMARY: Describe the qualitative position of the subject relative to the comparables in general terms (similar, inferior, superior). State that after considering all adjustments and the qualitative factors, the indications point toward a defined range per square foot that will be reconciled in the value conclusion. Do not state the final concluded value here.

Output only the three paragraphs, no headers, no commentary.""", 700)


def _prompt_sca_conclusion(v):
    return (f"""You are writing the Value Conclusion paragraph(s) for the Sales Comparison Approach section of a commercial real estate appraisal report. Write in formal, USPAP-compliant language. Active voice, no bullet points.

Adjustment data:
  Narrowed adjusted range: {v.get('SCA_ADJ_NARROW_LOW', '')} – {v.get('SCA_ADJ_NARROW_HIGH', '')} per SF
  Adjusted mean:           {v.get('SCA_ADJ_UNIT_MEAN', '')} per SF
  Adjusted median:         {v.get('SCA_ADJ_UNIT_MEDIAN', '')} per SF
  Qualitative summary:     {v.get('SCA_QUAL_CONCLUSION', '')}
  SCA concluded value:     {v.get('SCA_VALUE', '')}
  Subject GBA:             {v.get('GBA', '')} SF
  Unit of comparison:      {v.get('UNIT_OF_COMPARISON', 'price per square foot of gross building area')}
  Value interest:          {v.get('VALUE_INTEREST', 'as-is market value')}
  Effective date:          {v.get('EFFECTIVE_DATE', '')}

Write exactly TWO paragraphs:

1. RECONCILIATION OF INDICATIONS: Describe how the concluded per-unit value was selected from within the adjusted range — which comparables carry more weight and why (e.g., fewest adjustments, most similar size, most recent). Reference the median and mean as support. Do not use bullet points.

2. VALUE CONCLUSION: State the concluded value for the Sales Comparison Approach. Format: "Based on the foregoing analysis, the indicated value by the Sales Comparison Approach is [SCA_VALUE], as of [EFFECTIVE_DATE]." Keep this paragraph to two sentences maximum.

Output only the two paragraphs, no headers, no commentary.""", 500)


def _prompt_cap_rate(v):
    return (f"""You are writing the Cap Rate Derivation narrative for the Income Approach section of a commercial real estate appraisal report. Write in formal, USPAP-compliant appraisal language. Active voice, no bullet points.

Income Approach data:
  Property type:       {v.get('PROPERTY_SUBTYPE_FULL', 'commercial property')}
  Submarket:           {v.get('SUBMARKET', '')}
  Concluded cap rate:  {v.get('CAP_RATE', '')}
  Mortgage loan rate:  {v.get('LOAN_RATE', '')}
  Market rent (per SF):{v.get('MARKET_RENT', '')} annually
  Market rent range:   {v.get('MARKET_RENT_RANGE', '')} per SF annually
  Vacancy rate used:   {v.get('VACANCY_RATE', '')}%
  Property class:      {v.get('PROPERTY_CLASS', '')}
  City/State:          {v.get('PROPERTY_CITY', '')} {v.get('PROPERTY_STATE', '')}

Write exactly THREE paragraphs:

1. DERIVATION METHODOLOGY: Describe the cap rate derivation process — that the capitalization rate was derived from market data including sales of similar income-producing properties, investor surveys, and the band-of-investment technique. Explain that the rate reflects market participant return expectations for this property type and location.

2. MARKET SUPPORT: Describe the relationship between the concluded cap rate and the prevailing mortgage rate (note whether the concluded rate is above or below the loan rate and what this implies about equity returns). Reference the property class and submarket as factors influencing the rate. Two to three sentences.

3. CONCLUDED RATE: State the concluded overall capitalization rate and that it reflects the risk profile, income stability, and market conditions applicable to the subject property. One to two sentences.

Output only the three paragraphs, no headers, no commentary.""", 600)


def _prompt_encumbrances(v):
    return (f"""You are writing the Encumbrances and Property History narrative for a commercial real estate appraisal report. Write in formal, USPAP-compliant appraisal language. Active voice, no bullet points.

Property details:
  Property rights appraised: {v.get('PROPERTY_RIGHTS', '')}
  Current contract/listing:  {v.get('CURRENT_CONTRACT_PRICE', '')}
  Prior sale date:           {v.get('PRIOR_SALE_DATE', '')}
  Prior sale price:          {v.get('PRIOR_SALE_PRICE', 'N/A')}
  Owner name:                {v.get('OWNER_NAME', '')}
  Extraordinary assumption:  {v.get('EXTRAORDINARY_ASSUMPTION', 'None')}
  Hypothetical condition:    {v.get('HYPOTHETICAL_CONDITION', 'None')}
  Value interest:            {v.get('VALUE_INTEREST', '')}

Write exactly TWO paragraphs:

1. PROPERTY RIGHTS AND ENCUMBRANCES: State the property rights being appraised. Note that based on the appraiser's review of available information, no adverse easements, restrictions, or encumbrances that would materially affect value were identified beyond those typical for this property type in this jurisdiction. If there are extraordinary assumptions or hypothetical conditions, note them briefly.

2. OWNERSHIP AND TRANSACTION HISTORY: State the current owner and describe the prior sale history (or lack thereof) in terms of arm's-length transactions within the prior three years. State whether the property is currently listed for sale or under contract, and if so, at what price — or state that it is not. Keep this to two to three sentences.

Output only the two paragraphs, no headers, no commentary.""", 550)


def _prompt_reconciliation(v):
    sca_dev  = v.get('SCA_DEVELOPED', 'No').strip().lower() == 'yes'
    ia_dev   = v.get('IA_DEVELOPED',  'No').strip().lower() == 'yes'
    ca_dev   = v.get('CA_DEVELOPED',  'No').strip().lower() == 'yes'

    approaches = []
    if sca_dev:
        approaches.append(f"Sales Comparison Approach ({v.get('SCA_VALUE','')}, weighted {v.get('SCA_WEIGHT','')})")
    if ia_dev:
        approaches.append(f"Income Approach ({v.get('IA_VALUE','')}, weighted {v.get('IA_WEIGHT','')})")
    if ca_dev:
        approaches.append(f"Cost Approach ({v.get('COST_APPROACH_VALUE_ROUNDED','')}, weighted {v.get('CA_WEIGHT','')})")

    not_developed = []
    if not sca_dev:
        not_developed.append("Sales Comparison Approach")
    if not ia_dev:
        not_developed.append("Income Approach")
    if not ca_dev:
        not_developed.append(f"Cost Approach — {v.get('CA_DEVELOPED_FULL','not developed')}")

    return (f"""You are writing the final Reconciliation narrative for a commercial real estate appraisal report. Write in formal, USPAP-compliant appraisal language. Active voice, no bullet points.

Valuation summary:
  Property type:        {v.get('PROPERTY_SUBTYPE_FULL', '')}
  Approaches developed: {'; '.join(approaches) if approaches else 'None listed'}
  Not developed:        {'; '.join(not_developed) if not_developed else 'All developed'}
  SCA value:            {v.get('SCA_VALUE', 'N/A')}
  SCA weight:           {v.get('SCA_WEIGHT', 'N/A')}
  Income value:         {v.get('IA_VALUE', 'N/A')}
  Income weight:        {v.get('IA_WEIGHT', 'N/A')}
  Cost value:           {v.get('COST_APPROACH_VALUE_ROUNDED', 'N/A')}
  Cost weight:          {v.get('CA_WEIGHT', 'N/A')}
  Final value:          {v.get('VALUE_CONCLUSION', '')}
  Value type:           {v.get('VALUE_TYPE', 'Market Value')}
  Value interest:       {v.get('VALUE_INTEREST', '')}
  Effective date:       {v.get('EFFECTIVE_DATE', '')}

Write exactly THREE paragraphs:

1. APPROACHES AND RELIABILITY: List the approaches that were developed and briefly note which was not developed and why (if any). Describe the relative reliability of each developed approach for this property type — which approach(es) are typically most meaningful to market participants for this asset class, and which are secondary.

2. WEIGHT RATIONALE: Explain the weighting rationale. Describe why the primary approach receives greater weight (e.g., active sales market, most direct evidence of market behavior) and why the secondary approach receives lesser weight (e.g., corroborative, supports the primary). Do not simply repeat the percentages — explain the reasoning.

3. VALUE CONCLUSION: State the final value conclusion. Begin with: "Based on the foregoing analysis, the {v.get('VALUE_TYPE_SHORT','As-Is')} {v.get('VALUE_TYPE','Market Value')} of the {v.get('VALUE_INTEREST','fee simple interest')} in the subject property, as of {v.get('EFFECTIVE_DATE','')}, is concluded to be:" followed by the value in words (e.g., ONE MILLION THREE HUNDRED SEVENTY-FIVE THOUSAND DOLLARS — {v.get('VALUE_CONCLUSION','')}).

Output only the three paragraphs, no headers, no commentary.""", 800)


# ── Land adjustment narrative (workbook-driven, legacy) ───────────────────────

def _read_land_adj(workbook_path):
    try:
        import openpyxl
    except ImportError:
        raise RuntimeError("pip install openpyxl")

    wb = openpyxl.load_workbook(str(workbook_path), data_only=True)
    if "land" not in wb.sheetnames:
        return []
    ws = wb["land"]
    comps = []

    for r in range(5, 15):
        comp_no  = ws.cell(r, 2).value
        location = ws.cell(r, 3).value
        if not comp_no:
            continue

        def _pct(col):
            v = ws.cell(r, col).value
            if v is None: return 0.0
            try:
                f = float(v)
                return f if abs(f) <= 1.0 else f / 100.0
            except (ValueError, TypeError):
                return 0.0

        def _note(col):
            v = ws.cell(r, col).value
            return str(v).strip() if v and str(v).strip() not in ("", "None") else ""

        mkt_cond_pct  = _pct(4);  mkt_cond_note = _note(5)
        location_pct  = _pct(6);  location_note = _note(7)
        other_pct     = _pct(8);  other_note    = _note(9)
        dilmore_pct   = _pct(10)

        is_placeholder = not location or "address" in str(location).lower()
        all_zero = all(p == 0.0 for p in [mkt_cond_pct, location_pct, other_pct, dilmore_pct])
        if is_placeholder and all_zero:
            continue

        comps.append({
            "comp_no": int(comp_no), "location": str(location) if location else f"Comp {comp_no}",
            "mkt_cond_pct": mkt_cond_pct, "mkt_cond_note": mkt_cond_note,
            "location_pct": location_pct, "location_note": location_note,
            "other_pct": other_pct, "other_note": other_note,
            "dilmore_pct": dilmore_pct,
        })

    wb.close()
    return comps


def _prompt_land_adjustment(comps, variables):
    subject = variables.get("SUBJECT_ADDRESS", variables.get("PROPERTY_ADDRESS", "the subject"))
    lines = []
    for c in comps:
        parts = [f"  Comp {c['comp_no']} ({c['location']})"]
        if c["mkt_cond_pct"] != 0:
            d = "upward" if c["mkt_cond_pct"] > 0 else "downward"
            n = f" — {c['mkt_cond_note']}" if c["mkt_cond_note"] else ""
            parts.append(f"    Market Conditions: {d} {abs(c['mkt_cond_pct'])*100:.0f}%{n}")
        else:
            parts.append("    Market Conditions: no adjustment")
        if c["location_pct"] != 0:
            d = "upward" if c["location_pct"] > 0 else "downward"
            n = f" — {c['location_note']}" if c["location_note"] else ""
            parts.append(f"    Location: {d} {abs(c['location_pct'])*100:.0f}%{n}")
        else:
            parts.append("    Location: no adjustment")
        if c["dilmore_pct"] != 0:
            d = "upward" if c["dilmore_pct"] > 0 else "downward"
            parts.append(f"    Size (Dilmore): {d} {abs(c['dilmore_pct'])*100:.1f}%")
        if c["other_pct"] != 0:
            d = "upward" if c["other_pct"] > 0 else "downward"
            n = f" — {c['other_note']}" if c["other_note"] else ""
            parts.append(f"    Other: {d} {abs(c['other_pct'])*100:.0f}%{n}")
        lines.append("\n".join(parts))

    return (f"""You are writing the Adjustment Summary narrative for the Cost Approach — Land Valuation section of a commercial real estate appraisal report. Write in formal, USPAP-compliant appraisal language.

Subject property: {subject}

Per-comp adjustment data:
{chr(10).join(lines)}

Write exactly THREE paragraphs:

1. MARKET CONDITIONS paragraph: Describe which sales received market condition adjustments, in which direction, and the general reason. Sales with no market condition adjustment should be noted as requiring no adjustment. Do not list percentages — describe directionally only.

2. SIZE AND LOCATION paragraph: First sentence: size adjustments using the Dilmore Size Adjustment Scale along the 85% curve. Then describe location adjustments by direction. Do not list percentages.

3. ADDITIONAL FACTORS paragraph: One sentence noting that additional elements of comparison were considered but no further quantitative adjustments were applied, as these factors are addressed qualitatively in the reconciliation.

Output only the three paragraphs, no headers, no commentary. Each 2–4 sentences. Active voice. No bullets.""", 700)


# ── Dispatch map ──────────────────────────────────────────────────────────────

# Maps placeholder key (without [[ ]]) → prompt-builder function that takes (variables,)
# Land adjustment is handled separately (needs workbook)
_PROMPT_BUILDERS = {
    "INSPECTION_NARRATIVE":     _prompt_inspection,
    "MARKET_AREA_OVERVIEW":     _prompt_market_overview,
    "SCA_APPROACH_NARRATIVE":   _prompt_sca_approach,
    "SCA_ADJUSTMENT_NARRATIVE": _prompt_sca_adjustment,
    "SCA_CONCLUSION_NARRATIVE": _prompt_sca_conclusion,
    "CAP_RATE_NARRATIVE":       _prompt_cap_rate,
    "ENCUMBRANCES_NARRATIVE":   _prompt_encumbrances,
    "RECONCILIATION_NARRATIVE": _prompt_reconciliation,
}

_MODEL_COMMANDS = {
    "SCA_ADJUSTMENT_NARRATIVE": "adj-justify",
    "LAND_ADJUSTMENT_NARRATIVE": "adj-justify",
    "SCA_CONCLUSION_NARRATIVE": "reconcile",
    "RECONCILIATION_NARRATIVE": "reconcile",
}


# ── Public interface ──────────────────────────────────────────────────────────

def inject_all_narratives(doc_path, workbook_path, variables, config_path=None):
    """
    Scan *doc_path* (a filled Word document) for all known narrative placeholders
    and replace each one with AI-generated prose via the Claude API.

    Parameters
    ----------
    doc_path : str | Path
    workbook_path : str | Path   — used only for LAND_ADJUSTMENT_NARRATIVE
    variables : dict             — from fill_engine.load_variables()
    config_path : str | Path, optional

    Returns
    -------
    dict  — {placeholder: 'injected' | 'skipped' | 'error: <msg>'}
    """
    from docx import Document as _Document

    doc_path      = Path(doc_path)
    workbook_path = Path(workbook_path)
    doc           = _Document(str(doc_path))

    # Quick scan: which placeholders actually appear in this document?
    all_text = "\n".join(p.text for p in doc.paragraphs)
    present  = [k for k in list(_PROMPT_BUILDERS) + ["LAND_ADJUSTMENT_NARRATIVE"]
                if f"[[{k}]]" in all_text]

    if not present:
        return {}

    results = {}

    for key in present:
        command = _MODEL_COMMANDS.get(key, "draft")
        model = _get_model(config_path, command=command)
        print(f"  Generating {key} via {model} ...")
        try:
            if key == "LAND_ADJUSTMENT_NARRATIVE":
                comps = _read_land_adj(workbook_path)
                if not comps:
                    print(f"    ⚠  land tab empty — {key} skipped")
                    results[key] = "skipped (land tab empty)"
                    continue
                prompt, max_tok = _prompt_land_adjustment(comps, variables)
            else:
                prompt, max_tok = _PROMPT_BUILDERS[key](variables)

            narrative = _call_claude(prompt, model, max_tokens=max_tok)
            injected  = _inject_narrative_into_doc(doc, key, narrative)

            if injected:
                n_paras = len([c for c in narrative.split("\n\n") if c.strip()])
                print(f"    ✓  {key} — {n_paras} paragraph(s) injected")
                results[key] = "injected"
            else:
                print(f"    ⚠  {key} — placeholder not found in document after scan")
                results[key] = "skipped (placeholder not found)"

        except Exception as exc:
            print(f"    ✗  {key} — error: {exc}")
            results[key] = f"error: {exc}"

    doc.save(str(doc_path))
    return results


# ── Legacy entry point ────────────────────────────────────────────────────────

def generate_adjustment_narrative(workbook_path, variables=None, config_path=None):
    """
    Legacy function — generates only LAND_ADJUSTMENT_NARRATIVE from workbook.
    Kept for backward compatibility; inject_all_narratives() is preferred.
    """
    if variables is None:
        variables = {}
    workbook_path = Path(workbook_path)
    comps = _read_land_adj(workbook_path)
    if not comps:
        print("  Warning: land tab empty — [[LAND_ADJUSTMENT_NARRATIVE]] skipped.")
        return ""
    model = _get_model(config_path, command="adj-justify")
    prompt, max_tok = _prompt_land_adjustment(comps, variables)
    print(f"  Generating LAND_ADJUSTMENT_NARRATIVE via {model} ({len(comps)} comps) ...")
    narrative = _call_claude(prompt, model, max_tokens=max_tok)
    print(f"  ✓  {len(narrative)} chars")
    return narrative


# ── CLI test ──────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python narrative_generator.py <path/to/workbook.xlsx> [variables.json]")
        sys.exit(1)

    vars_path = sys.argv[2] if len(sys.argv) > 2 else None
    v = {}
    if vars_path and Path(vars_path).exists():
        with open(vars_path) as f:
            v = json.load(f)

    result = generate_adjustment_narrative(sys.argv[1], v)
    if result:
        print("\n── Generated Narrative ──")
        print(result)
