"""
fill_engine.py -- Axiom Commercial Appraisal
Document fill engine. Reads variables from JSON and/or the Excel
outputs tab, then replaces [[KEY]] placeholders in Word templates.
Templates are never modified -- output always goes to a new file.

Placeholder types
-----------------
  [[KEY]]         -- simple value substitution (this module handles these)
  [[X_BLOCK]]     -- multi-row block (ownership table, photos, etc.) -- left intact
  [[X_IMAGE]]     -- image insertion -- left intact
  [[X_NARRATIVE]] -- AI-generated prose -- left intact
  [[X_OVERVIEW]]  -- AI-generated prose -- left intact
  [[X_MAP]]       -- map image block -- left intact

Any key not found in variables is left as [[KEY]] in the output so
missing content is obvious rather than silently blank.

Conditional section removal
---------------------------
  When a flag variable (e.g. CA_DEVELOPED) is not "Yes", the
  corresponding major section is deleted from the output document.
  Sections are identified by their MainSectionHeading style paragraph.

  SECTION_REMOVAL_MAP: section_heading_text -> flag_variable_name
  If flag != "Yes" (case-insensitive), the section and everything
  up to (but not including) the next MainSectionHeading is removed.
"""

import json
import re
import shutil
from pathlib import Path

import openpyxl
from docx import Document
from docx.oxml.ns import qn

from presentation_variants import derive_presentation_variants


# -- Placeholder classification ------------------------------------------------

_BLOCK_SUFFIXES = (
    '_BLOCK', '_IMAGE', '_NARRATIVE', '_OVERVIEW', '_MAP',
    '_SECTION', '_TABLE',
)

def _is_block_placeholder(key):
    return any(key.endswith(s) for s in _BLOCK_SUFFIXES)


# -- Conditional section removal map ------------------------------------------

SECTION_REMOVAL_MAP = {
    'Cost Approach': 'CA_DEVELOPED',
}


# -- Variable loading ----------------------------------------------------------

def load_variables(json_path=None, workbook_path=None):
    variables = {}

    if json_path and Path(json_path).exists():
        with open(json_path, encoding='utf-8') as f:
            variables.update(json.load(f))

    if workbook_path and Path(workbook_path).exists():
        wb = openpyxl.load_workbook(workbook_path, data_only=True)
        if 'outputs' in wb.sheetnames:
            ws = wb['outputs']
            for row in ws.iter_rows(min_row=2, values_only=True):
                key   = str(row[1]).strip() if row[1] is not None else ''
                val_d = str(row[3]).strip() if row[3] is not None else ''
                val_c = str(row[2]).strip() if row[2] is not None else ''
                val   = val_d if val_d and val_d != 'None' else val_c
                if (key and re.fullmatch(r'[A-Z][A-Z0-9_]*', key)
                        and val and val != 'None'):
                    variables[key] = val

    return derive_presentation_variants(variables)


# -- Run-merge replacement -----------------------------------------------------

def _replace_in_paragraph(paragraph, variables, missing_keys):
    full_text = ''.join(run.text for run in paragraph.runs)
    if '[[' not in full_text:
        return

    new_text = full_text
    for key, value in variables.items():
        new_text = new_text.replace('[[' + key + ']]', str(value))

    remaining = re.findall(r'\[\[([A-Z0-9_]+)\]\]', new_text)
    for key in remaining:
        missing_keys.add(key)  # blocks and non-blocks; separated in fill_document

    if new_text == full_text:
        return

    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ''


def _replace_in_doc(doc, variables):
    missing_keys = set()

    def _walk(paragraphs):
        for para in paragraphs:
            _replace_in_paragraph(para, variables, missing_keys)

    _walk(doc.paragraphs)

    def _walk_table(table):
        for row in table.rows:
            for cell in row.cells:
                _walk(cell.paragraphs)
                for nested in cell.tables:
                    _walk_table(nested)

    for table in doc.tables:
        _walk_table(table)

    for section in doc.sections:
        for container in (section.header, section.footer):
            _walk(container.paragraphs)
            for table in container.tables:
                _walk_table(table)

    return missing_keys


# -- Blank-row removal ---------------------------------------------------------

def _remove_blank_table_rows(doc):
    for table in doc.tables:
        rows_to_delete = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) < 2:
                continue
            label = cells[0]
            values = cells[1:]
            if label and all(not v for v in values):
                rows_to_delete.append(row)

        for row in rows_to_delete:
            try:
                tbl = row._tr.getparent()
                tbl.remove(row._tr)
            except Exception:
                pass


# -- Conditional section removal -----------------------------------------------

def _get_heading_text(para_el):
    return ''.join(t.text or '' for t in para_el.iter(qn('w:t')))


def _is_main_section_heading(para_el):
    pPr = para_el.find(qn('w:pPr'))
    if pPr is None:
        return False
    pStyle = pPr.find(qn('w:pStyle'))
    return pStyle is not None and pStyle.get(qn('w:val')) == 'MainSectionHeading'


def _remove_conditional_sections(doc, variables):
    removed = []
    body = doc.element.body

    for section_text, flag_key in SECTION_REMOVAL_MAP.items():
        flag_val = variables.get(flag_key, 'No').strip().lower()
        if flag_val in ('yes', 'true', '1'):
            continue

        children = list(body)
        start_idx = None
        end_idx   = None

        for i, el in enumerate(children):
            if el.tag != qn('w:p'):
                continue
            if not _is_main_section_heading(el):
                continue
            text = _get_heading_text(el)
            if section_text in text:
                if start_idx is None:
                    start_idx = i
            elif start_idx is not None:
                end_idx = i
                break

        if start_idx is None:
            continue

        end = end_idx if end_idx is not None else len(children)
        to_remove = children[start_idx:end]

        for el in to_remove:
            pPr = el.find(qn('w:pPr')) if el.tag == qn('w:p') else None
            if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                continue
            try:
                body.remove(el)
            except Exception:
                pass

        removed.append(section_text)

    return removed


# -- Bookmark and TOC cleanup --------------------------------------------------

def _remove_orphaned_bookmark_ends(doc):
    """
    Remove w:bookmarkEnd elements with no matching w:bookmarkStart.
    Word shows a repair dialog when these exist. Caused by section removal
    taking a bookmarkStart while the matching bookmarkEnd sits outside that range.
    Also catches pre-existing orphans in the template.
    """
    body = doc.element.body
    start_ids = {el.get(qn('w:id'))
                 for el in body.iter(qn('w:bookmarkStart'))
                 if el.get(qn('w:id'))}
    for bk_end in list(body.iter(qn('w:bookmarkEnd'))):
        if bk_end.get(qn('w:id')) not in start_ids:
            parent = bk_end.getparent()
            if parent is not None:
                parent.remove(bk_end)


def _remove_orphaned_toc_entries(doc):
    """
    Remove TOC paragraphs whose w:hyperlink anchor points to a bookmark
    that no longer exists in the document.

    When a section is removed, its heading bookmark is deleted. Any TOC
    entry that hyperlinks to that bookmark becomes a dead entry and causes
    Word's 'unreadable content' repair dialog. This function removes those
    dead TOC entries entirely.

    Also removes any remaining PAGEREF field runs (begin/instrText/end) that
    reference missing bookmarks from any paragraph where they appear.
    """
    body = doc.element.body

    # Collect all bookmark names currently in the document
    existing_names = {
        el.get(qn('w:name'))
        for el in body.iter(qn('w:bookmarkStart'))
        if el.get(qn('w:name'))
    }

    # 1. Remove entire TOC paragraphs whose hyperlink anchor is dead
    for para in list(body.iter(qn('w:p'))):
        # Check direct hyperlink children (TOC entries are structured as
        # a paragraph containing a single hyperlink element)
        for hl in list(para):
            ns_hl = qn('w:hyperlink') if qn('w:hyperlink') else None
            if hl.tag != qn('w:hyperlink'):
                continue
            anchor = hl.get(qn('w:anchor'))
            if anchor and anchor not in existing_names:
                # The hyperlink targets a deleted bookmark -- remove this paragraph
                parent = para.getparent()
                if parent is not None:
                    try:
                        parent.remove(para)
                    except Exception:
                        pass
                break  # paragraph removed, move on

    # 2. Remove orphaned PAGEREF field groups from any remaining paragraphs
    #    A field group is: fldChar(begin) ... instrText ... fldChar(end)
    #    We identify it by instrText containing "PAGEREF <missing_name>"
    for para in list(body.iter(qn('w:p'))):
        _remove_pageref_fields_for_missing_bookmarks(para, existing_names)


def _remove_pageref_fields_for_missing_bookmarks(para_el, existing_names):
    """
    Within a single paragraph element, find and remove PAGEREF field groups
    that reference bookmark names not in existing_names.

    Field structure (all as w:r children of para or its hyperlink children):
      w:r/w:fldChar[@fldCharType='begin']
      w:r/w:instrText  (contains 'PAGEREF bookmarkname \\h')
      w:r/w:fldChar[@fldCharType='separate']
      w:r/w:t  (cached value)
      w:r/w:fldChar[@fldCharType='end']
    """
    # Operate on direct children of para AND on hyperlink children within para
    containers = [para_el] + list(para_el.iter(qn('w:hyperlink')))

    for container in containers:
        runs = list(container)
        i = 0
        while i < len(runs):
            run = runs[i]
            if run.tag != qn('w:r'):
                i += 1
                continue
            fc = run.find(qn('w:fldChar'))
            if fc is None or fc.get(qn('w:fldCharType')) != 'begin':
                i += 1
                continue
            # We're at a field begin -- look ahead for instrText
            target_name = None
            j = i + 1
            while j < len(runs):
                r2 = runs[j]
                instr = r2.find(qn('w:instrText'))
                if instr is not None and instr.text:
                    m = re.search(r'PAGEREF\s+(\S+)', instr.text)
                    if m:
                        name = m.group(1)
                        if name not in existing_names:
                            target_name = name
                    break
                fc2 = r2.find(qn('w:fldChar'))
                if fc2 is not None and fc2.get(qn('w:fldCharType')) == 'end':
                    break
                j += 1

            if target_name is not None:
                # Find the matching end and remove all runs from i to end (inclusive)
                k = i
                depth = 0
                while k < len(runs):
                    fc_k = runs[k].find(qn('w:fldChar'))
                    if fc_k is not None:
                        ftype = fc_k.get(qn('w:fldCharType'))
                        if ftype == 'begin':
                            depth += 1
                        elif ftype == 'end':
                            depth -= 1
                            if depth == 0:
                                break
                    k += 1
                # Remove runs[i:k+1] from container
                for r in runs[i:k+1]:
                    try:
                        container.remove(r)
                    except Exception:
                        pass
                # Refresh run list and don't advance i (index shifted)
                runs = list(container)
            else:
                i += 1


# -- Public API ----------------------------------------------------------------

def fill_document(template_path, output_path, variables, remove_blank_rows=True):
    """
    Fill a Word template with variables and write to output_path.
    Template is never modified.

    Returns dict: filled (int), missing (list), blocks (list), removed_sections (list).
    """
    template_path = Path(template_path)
    output_path   = Path(output_path)

    shutil.copy(template_path, output_path)
    doc = Document(str(output_path))

    # 1. Remove undeveloped approach sections
    removed_sections = _remove_conditional_sections(doc, variables)

    # 2. Substitute [[KEY]] placeholders
    missing_keys = _replace_in_doc(doc, variables)

    # 3. Remove blank table rows
    if remove_blank_rows:
        _remove_blank_table_rows(doc)

    # 4. Remove orphaned bookmarkEnd elements (no matching bookmarkStart).
    _remove_orphaned_bookmark_ends(doc)

    # 5. Remove dead TOC entries and orphaned PAGEREF fields left by
    #    section removal. Prevents Word's 'unreadable content' repair dialog.
    _remove_orphaned_toc_entries(doc)

    doc.save(str(output_path))

    blocks  = sorted(k for k in missing_keys if _is_block_placeholder(k))
    missing = sorted(k for k in missing_keys if not _is_block_placeholder(k))

    return {
        'filled':           len(variables),
        'missing':          missing,
        'blocks':           blocks,
        'removed_sections': removed_sections,
    }
