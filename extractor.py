"""
extractor.py — Axiom Commercial Appraisal
Extracts comp, lease comp, and assignment data from old appraisal reports
(.docx) and comp workbooks (.xlsx).

Strategy
--------
  Excel-first: column headers mapped via synonym dict → high-confidence rows.
  Word supplement: tables for any comps not in Excel; narrative regex for
  subject property, effective date, client, value conclusions, cap rate ranges.

Confidence levels
-----------------
  high   — extracted from a structured table (Excel or Word)
  medium — extracted from narrative text via regex
  low    — inferred or partially matched; flag for Derek to review
"""

import re
import json
from pathlib import Path
from datetime import datetime

from comparable_contract import comparable_identity
from financial_extractor import extract_financial_workbook
from pdf_financial_extractor import extract_financial_pdf
from observation_extractor import extract_market_observations
from artifact_extractor import extract_assignment_artifacts


# ─── Try importing document libraries ────────────────────────────────────────

try:
    from docx import Document as DocxDocument
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

try:
    import openpyxl
    _XLSX_OK = True
except ImportError:
    _XLSX_OK = False


# ─── Field synonym dictionaries ───────────────────────────────────────────────
# Each entry: canonical_field → list of header text variants (case-insensitive).
# Partial matches also work: if the header *contains* a variant it maps.

SALE_SYNONYMS = {
    # ── Property identification ──
    "address_street":       ["address", "street address", "property address",
                             "location", "street", "addr", "property location"],
    "address_city":         ["city"],
    "address_county":       ["county"],
    "address_state":        ["state"],
    "address_zip":          ["zip", "zip code", "postal code"],
    "submarket":            ["submarket", "market area", "area", "neighborhood",
                             "sub-market"],
    "property_type":        ["property type", "type", "prop type", "use type",
                             "land use", "property use"],
    "property_subtype":     ["subtype", "sub-type", "property subtype"],

    # ── Physical ──
    "gba_sf":               ["gba", "gross building area", "building area",
                             "total building area", "total sf", "building sf",
                             "bldg area", "building size", "gross area",
                             "total sqft", "bldg sf", "gross leasable area",
                             "total rentable area", "building", "area (sf)",
                             "gba size s/f", "building area sf", "building sf gba"],
    "nla_sf":               ["nla", "net leasable area", "net rentable area",
                             "rentable area", "rentable sf"],
    "site_area_sf":         ["site area", "lot area", "land area", "site",
                             "lot size", "land sf", "site (sf)", "site size",
                             "lot (sf)"],
    "year_built":           ["year built", "yr built", "built", "year constructed",
                             "year of construction", "yb", "constructed"],
    "stories":              ["stories", "floors", "# stories", "no. stories",
                             "number of stories", "story", "# of stories"],
    "construction_type":    ["construction", "construction type", "frame type",
                             "building type", "const type", "const."],
    "condition":            ["condition", "overall condition", "physical condition",
                             "quality/condition", "quality & condition"],
    "zoning":               ["zoning", "zone", "zoning class"],
    "flood_zone":           ["flood zone", "flood", "fema zone", "fema",
                             "flood area"],
    "parcel_id":            ["parcel", "parcel id", "apn", "tax id", "pin",
                             "parcel number", "tax parcel"],

    # ── Transaction ──
    "sale_price":           ["sale price", "sales price", "selling price",
                             "price", "sp", "closed price", "transaction price",
                             "purchase price", "consideration"],
    "sale_date":            ["sale date", "sales date", "date of sale",
                             "closing date", "close date", "date sold",
                             "sold date", "transfer date", "recording date",
                             "date closed"],
    "price_per_sf":         ["price/sf", "price per sf", "$/sf", "price per sqft",
                             "ppsf", "$/sqft", "price psf", "sale price/sf",
                             "sales price/sf", "sp/sf", "price per gba"],
    "cap_rate":             ["cap rate", "cap", "capitalization rate",
                             "going-in cap", "going in cap", "cr", "cap%",
                             "overall rate", "ro", "oar", "overall or cap rate",
                             "overall cap rate"],
    "noi":                  ["noi", "net operating income", "net income",
                             "stabilized noi"],
    "noi_per_sf":           ["noi/sf", "noi per sf", "$/sf noi", "noi psf"],
    "grantor":              ["grantor", "seller", "vendor", "from", "sold by",
                             "previous owner"],
    "grantee":              ["grantee", "buyer", "purchaser", "to", "bought by",
                             "new owner"],
    "deed_ref":             ["deed ref", "deed book", "instrument", "book/page",
                             "deed reference", "deed no", "deed number",
                             "instrument no", "db/pg", "deed book/page",
                             "instrument number", "recording"],
    "verification_source":  ["verification", "source", "data source",
                             "verified by", "verify", "verified via",
                             "confirmed by"],
}

LEASE_SYNONYMS = {
    # ── Property identification ──
    "address_street":       ["address", "street address", "property address",
                             "location", "street", "addr", "property location",
                             "building address"],
    "address_city":         ["city"],
    "address_county":       ["county"],
    "submarket":            ["submarket", "market area", "area", "neighborhood"],
    "property_type":        ["property type", "type", "use type", "land use"],

    # ── Tenant / transaction ──
    "tenant_name":          ["tenant", "lessee", "tenant name", "occupant",
                             "renter", "company"],
    "tenant_use":           ["use", "tenant use", "space use", "use type",
                             "occupancy type", "business type", "tenant type"],
    "lease_date":           ["lease date", "commencement", "commencement date",
                             "start date", "lease start", "effective date",
                             "date", "date of lease", "execution date"],
    "lease_expiration":     ["lease expiration", "expiration date", "expiration",
                             "lease expiry", "expiry date", "expiry",
                             "lease end", "end date", "termination date",
                             "lease termination", "lease expiration date"],
    "term_years":           ["term", "lease term", "term (years)", "years",
                             "initial term", "primary term"],
    "sf_leased":            ["sqft", "space sf", "leased sf", "sf leased",
                             "size s/f", "size sf", "size (sf)", "sf (leased)",
                             "demised premises", "gla", "nla",
                             "space size", "rentable sf", "leased area",
                             "premises", "space"],
    "base_rent_psf":        ["rent/sf", "rent psf", "base rent/sf",
                             "annual rent/sf", "$/sf", "rent per sf",
                             "asking rent", "base rent psf", "annual rent psf",
                             "contract rent", "contract rent/sf",
                             "rent per sqft", "net rent",
                             "rental rate psf", "rental rate", "rental rate/sf",
                             "market rent psf", "market rent/sf"],
    "base_rent_monthly":    ["monthly rent", "monthly", "rent/mo",
                             "rent per month", "monthly base rent"],
    "rent_structure":       ["structure", "rent type", "lease type", "nnn",
                             "gross", "expense structure", "lease structure",
                             "expense basis", "expense type", "reimbursement",
                             "expense reimbursement", "lease terms",
                             "lease basis", "expense basis", "terms"],
    "expense_stop_psf":     ["expense stop", "stop", "expense stop/sf",
                             "base year stop", "stop amount"],
    "ti_allowance_psf":     ["t.i.", "ti allowance", "tenant improvement",
                             "tia", "ti/sf", "ti psf", "improvement allowance",
                             "tenant improvement allowance", "fit-out",
                             "t.i. allowance", "allowance psf"],
    "free_rent_months":     ["free rent", "concession", "rent abatement",
                             "free months", "abatement", "free period"],
    "escalations":          ["escalation", "escalations", "increases", "bumps",
                             "steps", "rent increases", "annual increases",
                             "rent escalations", "rental increases"],
    "renewal_options":      ["options", "renewal", "renewal options",
                             "option terms", "renewal terms", "options to renew"],

    # ── Physical (shared with SALE_SYNONYMS) ──
    "gba_sf":               ["gba", "gross building area", "building area",
                             "total sf", "building sf"],
    "year_built":           ["year built", "yr built", "built", "yb"],
    "flood_zone":           ["flood zone", "fema zone", "flood"],
}


# ─── Utility ──────────────────────────────────────────────────────────────────

def _norm(text):
    """Lowercase, strip, collapse whitespace."""
    if text is None:
        return ""
    return re.sub(r"\s+", " ", str(text).lower().strip())


def _map_header(header, synonyms):
    """
    Return the canonical field name for a column header, or None if unrecognized.
    Tries exact match first, then substring.
    """
    h = _norm(header)
    if not h:
        return None
    for field, variants in synonyms.items():
        for v in variants:
            if h == v:
                return field
    # Second pass: substring match
    for field, variants in synonyms.items():
        for v in variants:
            if v in h or h in v:
                return field
    return None


def _clean_number(val):
    """
    Parse '$892,000' / '8.10%' / '9,780 SF' / '9780' → float.
    Returns None if unparseable.
    """
    if val is None:
        return None
    s = re.sub(r"[$,%]", "", str(val))
    s = re.sub(r"\s*(sf|sqft|sq\.?\s*ft\.?)\s*$", "", s, flags=re.IGNORECASE)
    s = s.replace(",", "").strip()
    # Handle ranges like "8.5 - 9.0" — take the midpoint
    range_match = re.match(r"^([\d.]+)\s*[-–to]+\s*([\d.]+)$", s)
    if range_match:
        try:
            return (float(range_match.group(1)) + float(range_match.group(2))) / 2
        except ValueError:
            pass
    try:
        return float(s)
    except ValueError:
        return None


def _clean_date(val):
    """Parse various date formats → ISO string YYYY-MM-DD, or original string."""
    if val is None:
        return None
    if hasattr(val, "strftime"):
        return val.strftime("%Y-%m-%d")
    s = str(val).strip()
    formats = [
        "%m/%d/%Y", "%m/%d/%y", "%m-%d-%Y",
        "%B %d, %Y", "%b %d, %Y", "%B %d %Y", "%b %d %Y",
        "%Y-%m-%d", "%d/%m/%Y",
        "%B %Y", "%b %Y",               # month/year only
    ]
    for fmt in formats:
        try:
            return datetime.strptime(s, fmt).strftime("%Y-%m-%d")
        except ValueError:
            continue
    return s  # return as-is; flag as medium confidence


_NUMBER_FIELDS = {
    "gba_sf", "nla_sf", "site_area_sf", "year_built", "stories",
    "sale_price", "price_per_sf", "cap_rate", "noi", "noi_per_sf",
    "sf_leased", "base_rent_psf", "base_rent_monthly", "expense_stop_psf",
    "ti_allowance_psf", "free_rent_months", "term_years",
}
_DATE_FIELDS = {"sale_date", "lease_date", "lease_expiration"}


def _coerce(field, raw):
    if field in _NUMBER_FIELDS:
        return _clean_number(raw)
    if field in _DATE_FIELDS:
        return _clean_date(raw)
    return str(raw).strip() if raw is not None else None


# ─── Excel extraction ─────────────────────────────────────────────────────────

def _is_lease_sheet_name(ws):
    """Heuristic on worksheet name."""
    name = _norm(ws.title)
    return any(k in name for k in ["lease", "rental", "rent comp", "income comp"])


# Patterns that appear in legitimate comp header rows
_HEADER_PATTERNS = {
    "sale date", "sale price", "price", "location", "address",
    "comp. no.", "comp no", "gba", "price/sf", "price psf",
    "price per sf", "oar", "cap rate", "noi", "year built",
    "rental rate", "rental rate psf", "rent psf", "rent/sf",
    "lease start", "lease date", "tenant", "size s/f", "sf leased",
    "lease terms", "grantor", "grantee", "deed",
}

def _sheet_header_row(ws):
    """
    Find the actual comp table header row in a worksheet.

    Strategy 1 (preferred): first row where ≥2 cells match known appraisal
    field patterns. Skips title/summary rows like "SUMMARY OF MARKET COMPARABLES".

    Strategy 2 (fallback): first row with ≥3 non-empty cells.
    """
    candidate = None
    for i, row in enumerate(ws.iter_rows(values_only=True), start=1):
        non_empty = [c for c in row if c is not None and str(c).strip()]
        if len(non_empty) < 2:
            continue
        known_hits = sum(
            1 for c in non_empty
            if any(pat in str(c).lower() for pat in _HEADER_PATTERNS)
        )
        if known_hits >= 2:
            return i, list(row)
        # Track first row with ≥3 non-empty as fallback
        if candidate is None and len(non_empty) >= 3:
            candidate = (i, list(row))

    return candidate if candidate else (None, [])


def _build_field_map(header, synonyms):
    """
    Map column indices → canonical field names.

    Rules:
    - First-wins-per-field: once a field is claimed by a column, later
      columns that would resolve to the same field are skipped (prevents
      the right-side adjustment grid from overwriting left-side data columns).
    - Stop at the second occurrence of a "COMP. NO." style header: that
      marks the start of the qualitative adjustment grid in Axiom Market
      Charts, where column names repeat and differ semantically.
    """
    _COMP_NO_HEADERS = {"comp. no.", "comp no", "comp #", "comp no.", "comparable no",
                        "comparable no.", "comp number"}
    field_map = {}
    field_claimed = set()
    comp_no_count = 0

    for col_idx, h in enumerate(header):
        if h is None:
            continue
        norm_h = _norm(str(h))
        if not norm_h:
            continue

        # Detect repeated COMP. NO. header → start of adjustment grid
        if norm_h in _COMP_NO_HEADERS:
            comp_no_count += 1
            if comp_no_count >= 2:
                break  # Everything past here is the adjustment/qualitative grid

        for field, variants in synonyms.items():
            if field in field_claimed:
                continue  # Already mapped — first-wins
            if any(v == norm_h or v in norm_h or norm_h in v for v in variants):
                field_map[col_idx] = field
                field_claimed.add(field)
                break

    return field_map


def _looks_like_comp_sheet(ws, header):
    """Return ('sale'|'lease'|None) based on sheet name and header content."""
    name = _norm(ws.title)
    header_text = " ".join(_norm(h) for h in header if h)

    lease_score = sum(1 for k in ["rental rate", "tenant", "lease", "lease date",
                                   "lease terms", "lessee", "size s/f", "sf leased"]
                      if k in header_text or k in name)
    # "rent" alone triggers on "rental" which appears in sale comp sheets too — use full phrases
    if "rent" in name and "rental" not in name:
        lease_score += 1
    sale_score  = sum(1 for k in ["sale price", "sale date", "gba", "cap rate",
                                   "oar", "grantor", "grantee", "deed",
                                   "price/sf", "price per sf", "price psf"]
                      if k in header_text)
    if lease_score >= 2:
        return "lease"
    if sale_score >= 2:
        return "sale"
    # Fewer hits — rely on name
    if any(k in name for k in ["sale", "comp", "comparab", "adjustment"]):
        return "sale"
    if any(k in name for k in ["lease", "rent", "income"]):
        return "lease"
    return None


def extract_from_xlsx(path):
    """
    Extract sale comps and lease comps from an Excel workbook.

    Returns
    -------
    dict with keys:
      'comps'       — list of {"data": {...}, "confidence": {...}, "sheet": str}
      'lease_comps' — same structure
      'warnings'    — list of warning strings (unmapped headers, etc.)
    """
    if not _XLSX_OK:
        return {"comps": [], "lease_comps": [], "warnings": ["openpyxl not installed"]}

    path = Path(path)
    try:
        wb = openpyxl.load_workbook(str(path), data_only=True)
    except Exception as e:
        return {"comps": [], "lease_comps": [], "warnings": [f"Could not open {path.name}: {e}"]}

    results  = {"comps": [], "lease_comps": [], "warnings": []}
    warnings = results["warnings"]

    for ws in wb.worksheets:
        # Skip obviously non-comp sheets
        skip_names = ["inputs", "intake", "output", "outputs", "toc", "cover",
                      "summary", "dilmore", "size", "adj", "deprec", "cost",
                      "income", "expense", "cash flow", "dcf",
                      "qualitative", "inflation"]
        if any(_norm(ws.title).startswith(s) for s in skip_names):
            continue

        header_row_idx, header = _sheet_header_row(ws)
        if not header:
            continue

        comp_type = _looks_like_comp_sheet(ws, header)
        if not comp_type:
            continue

        synonyms = LEASE_SYNONYMS if comp_type == "lease" else SALE_SYNONYMS

        # Map headers — first-wins-per-field, stops at second COMP. NO.
        field_map = _build_field_map(header, synonyms)

        # Report unmapped headers (informational only)
        mapped_cols = set(field_map.keys())
        unmapped = []
        for col_idx, h in enumerate(header):
            if h is None or not str(h).strip():
                continue
            if col_idx not in mapped_cols:
                unmapped.append(str(h).strip())

        if unmapped:
            warnings.append(
                f"  [{ws.title}] Unmapped headers: {', '.join(unmapped)}"
            )

        if len(field_map) < 2:
            warnings.append(f"  [{ws.title}] Too few recognizable columns — skipped")
            continue

        # Extract data rows
        for row in ws.iter_rows(min_row=header_row_idx + 1, values_only=True):
            if all(c is None or str(c).strip() == "" for c in row):
                continue

            data       = {}
            confidence = {}
            for col_idx, field in field_map.items():
                if col_idx >= len(row) or row[col_idx] is None:
                    continue
                raw = row[col_idx]
                if str(raw).strip() == "":
                    continue
                data[field]       = _coerce(field, raw)
                confidence[field] = "high"

            if not data:
                continue

            # Require minimum anchor fields — filters MEAN/SUBJECT/placeholder rows
            if comp_type == "sale":
                has_price = isinstance(data.get("sale_price"), (int, float)) and data["sale_price"] > 0
                has_addr  = bool(data.get("address_street", ""))
                if not (has_price and has_addr):
                    continue
            elif comp_type == "lease":
                has_rent = isinstance(data.get("base_rent_psf"), (int, float)) and data["base_rent_psf"] > 0
                has_addr = bool(data.get("address_street", ""))
                if not (has_rent and has_addr):
                    continue

            entry = {
                "data":       data,
                "confidence": confidence,
                "sheet":      ws.title,
                "source":     str(path),
                "type":       comp_type,
            }
            if comp_type == "lease":
                results["lease_comps"].append(entry)
            else:
                results["comps"].append(entry)

    wb.close()
    return results


# ─── Narrative patterns ───────────────────────────────────────────────────────

_PATTERNS = {
    "effective_date": [
        r"as of\s+([A-Za-z]+\.?\s+\d{1,2},?\s+\d{4})",
        r"effective date[:\s]+([A-Za-z]+\.?\s+\d{1,2},?\s+\d{4})",
        r"date of value[:\s]+([A-Za-z]+\.?\s+\d{1,2},?\s+\d{4})",
        r"value date[:\s]+([A-Za-z]+\.?\s+\d{1,2},?\s+\d{4})",
    ],
    "report_date": [
        r"date of report[:\s]+([A-Za-z]+\.?\s+\d{1,2},?\s+\d{4})",
        r"report date[:\s]+([A-Za-z]+\.?\s+\d{1,2},?\s+\d{4})",
        r"dated[:\s]+([A-Za-z]+\.?\s+\d{1,2},?\s+\d{4})",
    ],
    "client": [
        r"prepared for[:\s\n]+([^\n\.]{5,80})",
        r"addressed to[:\s\n]+([^\n\.]{5,80})",
        r"client[:\s]+([^\n\.]{5,60})",
    ],
    "reconciled_value": [
        r"market value[^\$\n]{0,40}\$\s*([\d,]+)",
        r"reconciled value[^\$\n]{0,40}\$\s*([\d,]+)",
        r"opinion of value[^\$\n]{0,40}\$\s*([\d,]+)",
        r"value conclusion[^\$\n]{0,40}\$\s*([\d,]+)",
        r"concluded[^\$\n]{0,20}\$\s*([\d,]+)",
        r"market value\s+(?:of|is)[^\$\n]{0,30}\$\s*([\d,]+)",
    ],
    "sca_value": [
        r"sales comparison approach[^\$\n]{0,60}\$\s*([\d,]+)",
        r"sales comparison[^\$\n]{0,40}indicate[sd]?[^\$\n]{0,20}\$\s*([\d,]+)",
    ],
    "ia_value": [
        r"income (?:capitalization )?approach[^\$\n]{0,60}\$\s*([\d,]+)",
    ],
    "ca_value": [
        r"cost approach[^\$\n]{0,60}\$\s*([\d,]+)",
    ],
    "cap_rate_range": [
        r"cap rates?\s+(?:range(?:d)?(?:\s+from)?|of|between)\s+([\d.]+)\s*%?\s*(?:to|[-–and])\s*([\d.]+)\s*%",
        r"([\d.]+)\s*%?\s*(?:to|[-–])\s*([\d.]+)\s*%\s+(?:cap|capitalization)",
    ],
    "address_street": [
        r"(?:subject property|property(?:\s+location)?)[:\s]+(\d+[^\n,]{5,60}(?:street|st|avenue|ave|blvd|boulevard|road|rd|drive|dr|lane|ln|way|pkwy|parkway|highway|hwy|circle|cir)[^\n,]{0,20})",
    ],
}


def _extract_narrative(text):
    """
    Apply regex patterns to report narrative text.
    Returns (data_dict, confidence_dict).
    """
    data       = {}
    confidence = {}
    lower      = text.lower()

    for field, patterns in _PATTERNS.items():
        for pattern in patterns:
            m = re.search(pattern, lower, re.IGNORECASE)
            if m:
                if field == "cap_rate_range":
                    try:
                        lo = float(m.group(1))
                        hi = float(m.group(2))
                        data["market_cap_rate_low"]  = lo
                        data["market_cap_rate_high"] = hi
                        confidence["market_cap_rate_low"]  = "medium"
                        confidence["market_cap_rate_high"] = "medium"
                    except (ValueError, IndexError):
                        pass
                elif field in ("reconciled_value", "sca_value", "ia_value", "ca_value"):
                    val = _clean_number(m.group(1))
                    if val and val > 50_000:   # sanity: appraisals are > $50k
                        data[field]       = val
                        confidence[field] = "medium"
                elif field in ("effective_date", "report_date"):
                    raw = m.group(1).strip().strip(".,;:")
                    data[field]       = _clean_date(raw)
                    confidence[field] = "medium"
                else:
                    data[field]       = m.group(1).strip().strip(".,;:")
                    confidence[field] = "medium"
                break   # first matching pattern wins

    return data, confidence


_INCOME_PATTERNS = {
    "period_year": [
        r"(?:period|fiscal|projection)\s+year[:\s]+(20\d{2})",
        r"\b(20\d{2})\s+(?:actual|pro\s*forma|projected|stabilized)\b",
    ],
    "period_type": [
        r"period\s+type[:\s]+(actual|pro\s*forma|projected|stabilized)",
        r"\b20\d{2}\s+(actual|pro\s*forma|projected|stabilized)\b",
    ],
    "pgi": [
        r"(?:potential gross income|pgi)[:\s\$]+([\d,]+(?:\.\d+)?)",
    ],
    "vacancy_pct": [
        r"(?:vacancy(?:\s+and\s+collection\s+loss)?|vacancy rate)[:\s]+([\d.]+)\s*%",
    ],
    "egi": [
        r"(?:effective gross income|egi)[:\s\$]+([\d,]+(?:\.\d+)?)",
    ],
    "total_expenses": [
        r"(?:total operating expenses|total expenses)[:\s\$]+([\d,]+(?:\.\d+)?)",
    ],
    "expense_ratio": [
        r"(?:expense ratio|operating expense ratio)[:\s]+([\d.]+)\s*%",
    ],
    "noi": [
        r"(?:net operating income|noi)[:\s\$]+([\d,]+(?:\.\d+)?)",
    ],
    "cap_rate_applied": [
        r"(?:capitalization rate|cap rate|overall rate)(?:\s+applied)?[:\s]+([\d.]+)\s*%",
    ],
}


def _extract_income_snapshot(text):
    """Extract a compact historical income summary from report text."""
    data = {}
    confidence = {}
    for field, patterns in _INCOME_PATTERNS.items():
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if not match:
                continue
            raw = match.group(1).strip()
            if field == "period_type":
                data[field] = raw
            elif field == "period_year":
                data[field] = int(raw)
            else:
                data[field] = _clean_number(raw)
            confidence[field] = "medium"
            break
    return data, confidence


# ─── Word document extraction ─────────────────────────────────────────────────

def _header_suggests_comp_table(table):
    """Return ('sale'|'lease'|None) for a Word table based on its first row."""
    if not table.rows:
        return None
    header_text = " ".join(_norm(c.text) for c in table.rows[0].cells)
    sale_hits  = sum(1 for k in ["sale price", "gba", "cap rate", "grantor",
                                  "grantee", "deed", "price/sf", "price per sf",
                                  "year built", "address"]
                     if k in header_text)
    lease_hits = sum(1 for k in ["rent", "tenant", "lease", "term",
                                  "ti", "lessee", "escalation", "option"]
                     if k in header_text)
    if lease_hits >= 2:
        return "lease"
    if sale_hits >= 2:
        return "sale"
    return None


def _extract_word_table(table, comp_type):
    """Extract records from a Word table. Returns list of {data, confidence}."""
    if len(table.rows) < 2:
        return []

    synonyms   = LEASE_SYNONYMS if comp_type == "lease" else SALE_SYNONYMS
    header_row = [c.text.strip() for c in table.rows[0].cells]
    field_map  = {}
    unmapped   = []

    for col_idx, h in enumerate(header_row):
        field = _map_header(h, synonyms)
        if field:
            field_map[col_idx] = field
        elif h:
            unmapped.append(h)

    if len(field_map) < 2:
        return []

    records = []
    for row in table.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if all(c == "" for c in cells):
            continue
        data       = {}
        confidence = {}
        for col_idx, field in field_map.items():
            if col_idx < len(cells) and cells[col_idx]:
                data[field]       = _coerce(field, cells[col_idx])
                confidence[field] = "high"
        if data:
            records.append({"data": data, "confidence": confidence,
                             "unmapped": unmapped})
    return records


def extract_from_docx(path):
    """
    Extract data from a Word appraisal report.

    Returns
    -------
    dict with keys:
      'comps'           — list of {data, confidence}
      'lease_comps'     — list of {data, confidence}
      'narrative'       — {data, confidence} from narrative text
      'income_data'     — {data, confidence} historical income summary
      'market_observations' — bounded, labeled narrative sections
      'warnings'        — list of warning strings
    """
    if not _DOCX_OK:
        return {"comps": [], "lease_comps": [], "narrative": {},
                "income_data": {},
                "market_observations": [],
                "warnings": ["python-docx not installed"]}

    path = Path(path)
    try:
        doc = DocxDocument(str(path))
    except Exception as e:
        return {"comps": [], "lease_comps": [], "narrative": {},
                "income_data": {},
                "market_observations": [],
                "warnings": [f"Could not open {path.name}: {e}"]}

    results = {
        "comps": [],
        "lease_comps": [],
        "narrative": {},
        "income_data": {},
        "market_observations": [],
        "artifacts": [],
        "warnings": [],
    }
    warnings = results["warnings"]

    # ── Narrative extraction ──────────────────────────────────────────────────
    full_text            = "\n".join(p.text for p in doc.paragraphs)
    narr_data, narr_conf = _extract_narrative(full_text)
    results["narrative"] = {"data": narr_data, "confidence": narr_conf}
    income_data, income_conf = _extract_income_snapshot(full_text)
    results["income_data"] = {
        "data": income_data,
        "confidence": income_conf,
        "source": str(path),
    } if income_data else {}
    results["market_observations"] = extract_market_observations(doc, path)

    # ── Reconciliation table → effective_date + value conclusions ─────────────
    recon = _extract_axiom_reconciliation(doc)
    if recon:
        for k, v in recon.items():
            if k not in narr_data or not narr_data[k]:
                narr_data[k] = v
                narr_conf[k] = "high"   # table-extracted, more reliable than regex

    # ── Axiom-format comp pages (two tables per comp) ─────────────────────────
    # Scan table list sequentially; pair left + right tables.
    tables      = doc.tables
    axiom_idxs  = set()   # track tables consumed by Axiom parser
    for i, table in enumerate(tables):
        if i in axiom_idxs:
            continue
        if _is_axiom_comp_table(table):
            left_raw  = _parse_axiom_left(table)
            right_raw = {}
            if i + 1 < len(tables) and _is_axiom_cont_table(tables[i + 1]):
                right_raw = _parse_axiom_right(tables[i + 1])
                axiom_idxs.add(i + 1)
            axiom_idxs.add(i)
            comp = _build_axiom_comp(left_raw, right_raw, path)
            if comp["data"]:
                results["comps"].append(comp)

    # ── Generic table extraction (fallback for non-Axiom reports) ─────────────
    for i, table in enumerate(tables):
        if i in axiom_idxs:
            continue
        comp_type = _header_suggests_comp_table(table)
        if not comp_type:
            continue
        records = _extract_word_table(table, comp_type)
        for rec in records:
            rec["source"] = str(path)
            if rec.get("unmapped"):
                warnings.append(
                    f"  [Word table] Unmapped headers: {', '.join(rec['unmapped'])}"
                )
            if comp_type == "lease":
                results["lease_comps"].append(rec)
            else:
                results["comps"].append(rec)

    return results


# ─── Axiom-format comp page parser ───────────────────────────────────────────
#
# Old Axiom / Falkner Firm reports use a two-table-per-comp layout:
#
#   LEFT table  (3 cols, ~31 rows): "Improved Sale No. X"
#     col 0 = label (or section heading), col 1 = label (repeated), col 2 = value
#     Sections: Property Identification / Sale Data / Land Data
#
#   RIGHT table (2 cols, ~24–29 rows): "Improved Sale No. X (Cont.)"
#     col 0 = label, col 1 = value
#     Sections: General Physical Data / Income Analysis / Indicators / Remarks
#
# Detection: first cell of table contains "Improved Sale No." (case-insensitive).
# Pairs are found by scanning the table list sequentially.

_AXIOM_LEFT_MAP = {
    # label text (lowercased) → canonical field or special key
    "record id":        "_record_id",       # not in schema → goes into notes
    "property type":    "property_type",
    "address":          "_address_raw",     # needs city/zip split
    "tax id":           "parcel_id",
    "msa":              "submarket",
    "grantor":          "grantor",
    "grantee":          "grantee",
    "sale date":        "sale_date",
    "deed book/page":   "deed_ref",
    "verification":     "verification_source",
    "sale price":       "sale_price",
    "adjusted price":   "_adjusted_price",  # use only if sale_price missing
    "land size":        "_land_size_raw",   # "X.XX Acres or YY,YYY SF"
    "flood info":       "_flood_raw",       # '"Zone X" 01117C... - date'
}

_AXIOM_RIGHT_MAP = {
    "building type":                "property_subtype",
    "sf":                           "gba_sf",
    "construction type":            "construction_type",
    "stories":                      "stories",
    "year built":                   "year_built",
    "condition":                    "condition",
    "net operating income":         "noi",
    "sale price/ sf":               "price_per_sf",
    "sale price/sf":                "price_per_sf",
    "overall or cap rate":          "_cap_rate_raw",  # "7.5%" → divide by 100
    "net operating income/sq. ft.": "noi_per_sf",
    "net operating income/sq ft":   "noi_per_sf",
    "remarks":                      "_remarks",
}


def _parse_land_size_raw(raw):
    """'6.390 Acres or 278,348 SF' → SF as float.  Falls back to acre→SF."""
    if not raw:
        return None
    m = re.search(r"or\s+([\d,]+)\s*sf", raw, re.IGNORECASE)
    if m:
        return _clean_number(m.group(1))
    m = re.search(r"([\d.]+)\s*acres?", raw, re.IGNORECASE)
    if m:
        acres = _clean_number(m.group(1))
        return round(acres * 43560) if acres else None
    return _clean_number(raw)


def _parse_axiom_address(raw):
    """
    Split 'Street, City, County, State ZIP' → dict of address components.
    Works on the Axiom/Falkner format:
      '4201-4209 University Blvd E, Tuscaloosa, Tuscaloosa County, Alabama 35404'
    """
    if not raw:
        return {}
    parts = [p.strip() for p in raw.split(",")]
    out = {}
    if parts:
        out["address_street"] = parts[0]
    if len(parts) >= 2:
        out["address_city"] = parts[1]
    if len(parts) >= 3:
        county_raw = parts[2]
        out["address_county"] = re.sub(r"\s*county\s*$", "", county_raw,
                                        flags=re.IGNORECASE).strip()
    if len(parts) >= 4:
        state_zip = parts[-1].strip()
        m = re.match(r"([A-Za-z ]+?)\s+(\d{5})", state_zip)
        if m:
            out["address_state"] = m.group(1).strip()
            out["address_zip"]   = m.group(2)
        else:
            out["address_state"] = state_zip
    return out


def _is_axiom_comp_table(table):
    """True if this is a left-side 'Improved Sale No. X' table (not Cont.)."""
    if not table.rows:
        return False
    txt = _norm(table.rows[0].cells[0].text)
    return "improved sale no" in txt and "cont" not in txt


def _is_axiom_cont_table(table):
    """True if this is a right-side 'Improved Sale No. X (Cont.)' table."""
    if not table.rows:
        return False
    txt = _norm(table.rows[0].cells[0].text)
    return "improved sale no" in txt and "cont" in txt


def _parse_axiom_left(table):
    """
    Parse a 3-col Axiom comp left table.
    Returns {raw_field → raw_value_string}.
    """
    raw = {}
    for row in table.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 2:
            continue
        label = _norm(cells[0])
        if not label:
            continue
        # Value: last non-empty cell that isn't a repeat of the label
        value = None
        for cell in reversed(cells[1:]):
            cell_s = cell.strip()
            if cell_s and _norm(cell_s) != label:
                value = cell_s
                break
        # Fallback: if all non-label cells repeat the label, still grab col 1
        if value is None and len(cells) > 1 and cells[1].strip():
            value = cells[1].strip()
        if value and label:
            raw[label] = value
    return raw


def _parse_axiom_right(table):
    """
    Parse a 2-col Axiom comp right (Cont.) table.
    Returns {raw_field → raw_value_string}.
    """
    raw = {}
    for row in table.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 2:
            continue
        label = _norm(cells[0])
        value = cells[1].strip() if cells[1].strip() else ""
        if label and value:
            raw[label] = value
    return raw


def _build_axiom_comp(left_raw, right_raw, source_path):
    """
    Merge left + right raw dicts into a structured {data, confidence, source} record.
    """
    data = {}
    conf = {}

    # ── Left table ────────────────────────────────────────────────────────────
    for label, value in left_raw.items():
        field = _AXIOM_LEFT_MAP.get(label)
        if not field:
            continue

        if field == "_address_raw":
            for k, v in _parse_axiom_address(value).items():
                data[k] = v
                conf[k] = "high"

        elif field == "_land_size_raw":
            sf = _parse_land_size_raw(value)
            if sf:
                data["site_area_sf"] = sf
                conf["site_area_sf"] = "high"

        elif field == "_adjusted_price":
            if "sale_price" not in data:
                v = _clean_number(value)
                if v:
                    data["sale_price"] = v
                    conf["sale_price"] = "high"

        elif field == "_record_id":
            data["notes"] = f"Record ID: {value}"
            conf["notes"] = "high"

        elif field == "_flood_raw":
            m = re.search(r'"?(Zone\s+[A-Z]+)"?', value, re.IGNORECASE)
            data["flood_zone"] = m.group(1) if m else value
            conf["flood_zone"] = "high"

        elif field == "sale_date":
            data["sale_date"] = _clean_date(value)
            conf["sale_date"] = "high"

        elif field == "sale_price":
            v = _clean_number(value)
            if v:
                data["sale_price"] = v
                conf["sale_price"] = "high"

        else:
            # grantor, grantee, deed_ref, verification_source, submarket,
            # parcel_id, property_type
            data[field] = value
            conf[field] = "high"

    # ── Right table ───────────────────────────────────────────────────────────
    for label, value in right_raw.items():
        field = _AXIOM_RIGHT_MAP.get(label)
        if not field:
            continue

        if field == "_remarks":
            existing = data.get("notes", "")
            data["notes"] = (existing + "\n\n" + value).strip() if existing else value
            conf["notes"] = "high"

        elif field == "_cap_rate_raw":
            v = _clean_number(value)          # "7.5%" → 7.5
            if v is not None:
                data["cap_rate"] = v / 100 if v > 1 else v   # store as 0.075
                conf["cap_rate"] = "high"

        elif field in ("gba_sf", "noi", "price_per_sf", "noi_per_sf"):
            v = _clean_number(value)
            if v is not None:
                data[field] = v
                conf[field] = "high"

        elif field in ("stories", "year_built"):
            v = _clean_number(value)
            if v is not None:
                data[field] = int(v)
                conf[field] = "high"

        else:
            # property_subtype, construction_type, condition
            data[field] = value
            conf[field] = "high"

    return {"data": data, "confidence": conf, "source": str(source_path),
            "confidence_source": "axiom_comp_page"}


def _extract_axiom_reconciliation(doc):
    """
    Find the reconciliation table and return effective_date + value conclusions.
    Looks for a table whose first row contains 'Reconciled' or 'Market Value'.
    Returns dict with keys: effective_date, reconciled_value (floats/strings).
    """
    result = {}
    for table in doc.tables:
        if not table.rows:
            continue
        header_txt = _norm(table.rows[0].cells[0].text)
        if "reconciled" not in header_txt and "market value" not in header_txt:
            continue
        # Scan rows for the data row: "As Is" Market Value | Estate | Date | $Value
        # Skip header rows — only accept rows where cells[3] looks like a dollar value
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 4:
                continue
            # Must be a value row: col 3 starts with $ or is a number
            if not cells[3] or not re.match(r"^\$?[\d,]+", cells[3].strip()):
                continue
            if "market value" in _norm(cells[0]) or "as is" in _norm(cells[0]):
                if cells[2]:
                    result["effective_date"] = _clean_date(cells[2])
                v = _clean_number(cells[3])
                if v and v > 50_000:
                    result["reconciled_value"] = v
                return result
    return result


# ─── Folder scanner ───────────────────────────────────────────────────────────
#
# Based on observed folder structure:
#
#  [file_no] - [property_type] - [address], [city]/   ← assignment root
#    REPORT [file_no].docx                             ← main report (narrative)
#    EXHIBITS [file_no].docx                           ← comp pages (tables)
#    Improved Sales.docx / Campground Sales Comps.docx ← standalone sale comps
#    MARKET CHART.xlsx / Capitalization Rate Comps.xlsx← cap rate / market data
#    MARKET CHART - RENTAL.xlsx / RentalChart_*.xlsx   ← lease comps
#    Smith Lake RV Park RR rev2.xlsx                   ← rent rolls (RR)
#    income chart.docx                                 ← income approach
#    email/  File Scans/  pic/  REPORT/                ← archive subfolders (skip)
#
# The folder NAME itself encodes: file_no, property_type, street, city.

# Subfolders to skip entirely — nothing to extract there
_SKIP_SUBFOLDERS = {"email", "file scans", "pic", "report", "research info",
                    "sketches", "as complete", "as is", "infomation provided",
                    "information provided"}

# Filenames to skip (engagement letters, contracts, deeds, surveys, tax cards, etc.)
_SKIP_FILENAME_FRAGMENTS = [
    "eng letter", "engagement letter", "executed contract", "deed",
    "tax card", "survey", "measurements", "benchmark", "parcel update",
    "quick stats", "newmark", "submarket report", "market report",
]

def _skip_file(path):
    """Return True if this file should be ignored."""
    name = path.name.lower()
    if name.startswith("~"):
        return True
    return any(frag in name for frag in _SKIP_FILENAME_FRAGMENTS)


def _classify_docx(path):
    """
    Return the role of a Word document based on its filename.
    Returns one of: 'report' | 'exhibits' | 'sale_comps' | 'income' | 'other'
    """
    name = path.stem.lower()
    if name.startswith("report"):
        return "report"
    if "exhibit" in name:
        return "exhibits"
    if any(k in name for k in ["sales", "sale comp", "improved sales",
                                "campground sales", "comp"]):
        return "sale_comps"
    if "income" in name or "income chart" in name:
        return "income"
    return "other"


def _classify_xlsx(path):
    """
    Return the role of an Excel file based on its filename.
    Returns one of: 'cap_rates' | 'rental_comps' | 'rent_roll' | 'expense' |
    'market' | 'other'
    """
    name = path.stem.lower()
    if any(k in name for k in ["cap rate", "capitalization rate",
                                "cap comp", "market chart"]) and "rental" not in name:
        return "cap_rates"
    if any(k in name for k in ["rental", "rent comp", "lease comp",
                                "market chart - rental", "rentchart", "rentalchart"]):
        return "rental_comps"
    if "rr" in name.split() or "rent roll" in name or name.endswith(" rr"):
        return "rent_roll"
    if any(k in name for k in [
        "expense",
        "operating history",
        "historical operations",
        "operating statement",
        "income statement",
        "profit and loss",
        "p&l",
    ]):
        return "expense"
    if "market" in name:
        return "market"
    return "other"


def _classify_pdf(path):
    """
    Return the role of a PDF based on its filename.
    Returns one of: 'rent_roll' | 'expense' | 'other'
    """
    name = path.stem.lower()
    if (
        "rent roll" in name
        or "tenant rent roll" in name
        or "rr" in name.split()
        or name.endswith(" rr")
    ):
        return "rent_roll"
    if any(k in name for k in [
        "expense",
        "operating history",
        "operating statement",
        "income statement",
        "profit and loss",
        "p&l",
        "finance statement",
    ]):
        return "expense"
    return "other"


def _parse_folder_name(folder_name):
    """
    Extract file_no, property_type, and city from a folder name like:
      '25C008 - retail - 14th street s, bessemer'
      '25c005 Crane Hill'
    Returns dict with keys file_no, property_type, city (all may be None).
    """
    result = {"file_no": None, "property_type": None, "city": None}

    # File number: first token matching pattern like 25C008, 25c005, 24C106
    m = re.match(r"^(\d{2}[Cc]\d+)", folder_name.strip())
    if m:
        result["file_no"] = m.group(1).upper()

    # Pattern: [file_no] - [property_type] - [address], [city]
    dash_parts = re.split(r"\s+-\s+", folder_name)
    if len(dash_parts) >= 3:
        result["property_type"] = dash_parts[1].strip().title()
        # City is often after last comma in the last part
        last = dash_parts[-1]
        if "," in last:
            result["city"] = last.split(",")[-1].strip().title()

    return result


def scan_assignment_folder(folder_path):
    """
    Scan one assignment folder (root level only — skip archive subfolders).

    Returns dict:
      folder, name, folder_meta (file_no/type/city from name),
      reports, exhibits, sale_comp_docs, income_docs, other_docs,
      cap_rate_xls, rental_comp_xls, rent_roll_xls, expense_xls, market_xls,
      other_xls
    """
    folder = Path(folder_path)
    result = {
        "folder":           str(folder),
        "name":             folder.name,
        "folder_meta":      _parse_folder_name(folder.name),
        "reports":          [],   # REPORT *.docx — narrative + value conclusions
        "exhibits":         [],   # EXHIBITS *.docx — comp pages (tables)
        "sale_comp_docs":   [],   # standalone sale comp Word docs
        "income_docs":      [],   # income chart Word docs
        "other_docs":       [],   # other Word docs
        "cap_rate_xls":     [],   # cap rate / market chart Excel files
        "rental_comp_xls":  [],   # rental / lease comp Excel files
        "rent_roll_xls":    [],   # rent roll Excel files
        "expense_xls":      [],   # operating expense/history Excel files
        "market_xls":       [],   # general market Excel files
        "other_xls":        [],   # other Excel files
        "rent_roll_pdfs":   [],   # native/scanned rent roll PDFs
        "expense_pdfs":     [],   # native/scanned operating statement PDFs
        "other_pdfs":       [],   # other PDF files
    }

    for item in folder.iterdir():
        # Skip archive subfolders
        if item.is_dir():
            if item.name.lower() in _SKIP_SUBFOLDERS:
                continue
            # Recurse one level into non-archive subfolders? No — keep root only.
            continue

        if _skip_file(item):
            continue

        ext = item.suffix.lower()

        if ext == ".docx":
            role = _classify_docx(item)
            result[{
                "report":     "reports",
                "exhibits":   "exhibits",
                "sale_comps": "sale_comp_docs",
                "income":     "income_docs",
                "other":      "other_docs",
            }[role]].append(str(item))

        elif ext == ".xlsx":
            role = _classify_xlsx(item)
            result[{
                "cap_rates":    "cap_rate_xls",
                "rental_comps": "rental_comp_xls",
                "rent_roll":    "rent_roll_xls",
                "expense":      "expense_xls",
                "market":       "market_xls",
                "other":        "other_xls",
            }[role]].append(str(item))

        elif ext == ".pdf":
            role = _classify_pdf(item)
            result[{
                "rent_roll": "rent_roll_pdfs",
                "expense":   "expense_pdfs",
                "other":     "other_pdfs",
            }[role]].append(str(item))

    return result


def scan_projects_root(root_path):
    """
    Walk root_path where each immediate subdirectory is one assignment.
    Returns list of scan dicts for folders that contain extractable files.
    """
    root        = Path(root_path)
    assignments = []
    for subfolder in sorted(root.iterdir()):
        if not subfolder.is_dir() or subfolder.name.startswith("."):
            continue
        scan = scan_assignment_folder(subfolder)
        has_data = any([
            scan["reports"], scan["exhibits"], scan["sale_comp_docs"],
            scan["income_docs"], scan["cap_rate_xls"],
            scan["rental_comp_xls"], scan["rent_roll_xls"],
            scan["expense_xls"], scan["rent_roll_pdfs"],
            scan["expense_pdfs"],
        ])
        if has_data:
            assignments.append(scan)
    return assignments


# ─── Combined extraction for one assignment folder ────────────────────────────

def extract_assignment(scan):
    """
    Run full extraction on one scanned assignment folder.
    Processes all relevant files; merges and deduplicates results.

    Parameters
    ----------
    scan : dict from scan_assignment_folder()

    Returns
    -------
    dict with keys:
      folder_name, folder_meta, comps, lease_comps, narrative,
      income_data, warnings, sources
    """
    result = {
        "folder_name":  scan["name"],
        "folder_meta":  scan["folder_meta"],
        "comps":        [],
        "lease_comps":  [],
        "narrative":    {},
        "income_data":  {},
        "rent_roll_entries": [],
        "expense_records": [],
        "market_observations": [],
        "warnings":     [],
        "sources":      [],
    }

    def _add_comps(extracted, source_path):
        """Merge extracted comps using stable transaction identity."""
        existing_keys = {
            comparable_identity("sale", comp["data"])
            for comp in result["comps"]
        }
        for comp in extracted:
            identity_key = comparable_identity("sale", comp["data"])
            if identity_key in existing_keys:
                continue
            comp["source"] = str(source_path)
            result["comps"].append(comp)
            existing_keys.add(identity_key)

    def _add_lease_comps(extracted, source_path):
        """Merge extracted lease comps using stable transaction identity."""
        existing_keys = {
            comparable_identity("lease", comp["data"])
            for comp in result["lease_comps"]
        }
        for lc in extracted:
            identity_key = comparable_identity("lease", lc["data"])
            if identity_key in existing_keys:
                continue
            lc["source"] = str(source_path)
            result["lease_comps"].append(lc)
            existing_keys.add(identity_key)

    # ── 1. Cap rate / market Excel files (highest confidence for IA data) ────
    for xls_path in scan["cap_rate_xls"] + scan["market_xls"]:
        result["sources"].append(xls_path)
        data = extract_from_xlsx(Path(xls_path))
        _add_comps(data["comps"], xls_path)
        _add_lease_comps(data["lease_comps"], xls_path)
        result["warnings"].extend(data["warnings"])

    # ── 2. Rental / lease comp Excel files ───────────────────────────────────
    for xls_path in scan["rental_comp_xls"]:
        result["sources"].append(xls_path)
        data = extract_from_xlsx(Path(xls_path))
        _add_lease_comps(data["lease_comps"], xls_path)
        # Rental Excel sometimes has sale comps too
        _add_comps(data["comps"], xls_path)
        result["warnings"].extend(data["warnings"])

    # ── 3. Exhibits Word docs (comp pages — rich table data) ─────────────────
    for doc_path in scan["exhibits"]:
        result["sources"].append(doc_path)
        data = extract_from_docx(Path(doc_path))
        _add_comps(data["comps"], doc_path)
        _add_lease_comps(data["lease_comps"], doc_path)
        result["warnings"].extend(data["warnings"])

    # ── 4. Standalone sale comp Word docs ────────────────────────────────────
    for doc_path in scan["sale_comp_docs"]:
        result["sources"].append(doc_path)
        data = extract_from_docx(Path(doc_path))
        _add_comps(data["comps"], doc_path)
        result["warnings"].extend(data["warnings"])

    # ── 5. Main report (narrative: effective date, client, value conclusions) ─
    # Use newest report if multiple (sort by file size desc as proxy for final version)
    reports = sorted(scan["reports"], key=lambda p: Path(p).stat().st_size, reverse=True)
    if reports:
        rpt_path = reports[0]
        result["sources"].append(rpt_path)
        data = extract_from_docx(Path(rpt_path))
        result["narrative"] = data["narrative"]
        result["assignment_source"] = str(rpt_path)
        if data.get("income_data") and not result["income_data"]:
            result["income_data"] = data["income_data"]
        result["market_observations"] = data.get("market_observations", [])
        # Reports also have comp tables — fill any gaps
        _add_comps(data["comps"], rpt_path)
        _add_lease_comps(data["lease_comps"], rpt_path)
        result["warnings"].extend(data["warnings"])

    # ── 6. Standalone income summaries ───────────────────────────────────────
    for doc_path in scan["income_docs"]:
        result["sources"].append(doc_path)
        data = extract_from_docx(Path(doc_path))
        if data.get("income_data"):
            result["income_data"] = data["income_data"]
            result["income_source"] = str(doc_path)
        result["warnings"].extend(data["warnings"])

    # ── 7. Rent rolls and operating-expense workbooks ────────────────────────
    for workbook_path in scan["rent_roll_xls"] + scan["expense_xls"]:
        result["sources"].append(workbook_path)
        data = extract_financial_workbook(Path(workbook_path))
        result["rent_roll_entries"].extend(data["rent_roll_entries"])
        result["expense_records"].extend(data["expense_records"])
        result["warnings"].extend(data["warnings"])

    # ── 8. External and Word-embedded source artifacts ───────────────────────
    for pdf_path in scan["rent_roll_pdfs"]:
        result["sources"].append(pdf_path)
        data = extract_financial_pdf(Path(pdf_path))
        result["rent_roll_entries"].extend(data["rent_roll_entries"])
        result["expense_records"].extend(data["expense_records"])
        result["warnings"].extend(data["warnings"])

    for pdf_path in scan["expense_pdfs"]:
        result["sources"].append(pdf_path)
        data = extract_financial_pdf(Path(pdf_path))
        result["rent_roll_entries"].extend(data["rent_roll_entries"])
        result["expense_records"].extend(data["expense_records"])
        result["warnings"].extend(data["warnings"])

    office_containers = (
        scan["reports"]
        + scan["exhibits"]
        + scan["sale_comp_docs"]
        + scan["income_docs"]
        + scan["other_docs"]
        + scan["cap_rate_xls"]
        + scan["rental_comp_xls"]
        + scan["rent_roll_xls"]
        + scan["expense_xls"]
        + scan["market_xls"]
        + scan["other_xls"]
        + scan["rent_roll_pdfs"]
        + scan["expense_pdfs"]
        + scan["other_pdfs"]
    )
    artifact_data = extract_assignment_artifacts(
        scan["folder"],
        office_containers,
    )
    result["artifacts"] = artifact_data["artifacts"]
    result["sources"].extend(artifact_data["sources"])
    result["warnings"].extend(artifact_data["warnings"])

    # ── Inject folder metadata into narrative if not already present ──────────
    meta = scan["folder_meta"]
    narr = result["narrative"].setdefault("data", {})
    conf = result["narrative"].setdefault("confidence", {})
    if meta.get("property_type") and not narr.get("property_type"):
        narr["property_type"] = meta["property_type"]
        conf["property_type"] = "high"
    if meta.get("city") and not narr.get("city"):
        narr["address_city"] = meta["city"]
        conf["address_city"] = "high"
    return result

def extract_from_assignment(folder_path):
    """
    Convenience wrapper: scan a folder then run full extraction.
    Accepts a string or Path; returns the same dict as extract_assignment().
    """
    scan = scan_assignment_folder(folder_path)
    return extract_assignment(scan)
