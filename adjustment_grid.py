"""
adjustment_grid.py — Axiom Commercial Appraisal Platform
==========================================================
Reads the Phase 6 adjustment-grid workbook tabs (sca_adjustment_grid,
land_adjustment_grid, sca_qualitative, land_qualitative) and injects each
as a Word table at its [[..._GRID_BLOCK]] marker.

Unlike comp_builder.py's fixed COMP_COLUMNS letter map, these sheets' column
sets are variable (they depend on adjustment_factors.json's per-property-
type preset), so the header row is read at runtime and the column map is
built from header text -> column index, following docs/ADJUSTMENT_GRID_
DESIGN.md's Pipeline step 6. A hand-added one-off category column, or a
preset with fewer columns, both just work without a code change.

Category adjustments and qualitative factor scores are manual-entry cells;
Time Adjustment, Adjusted Price, Net Adjustment/Overall, and Indicated
Value/Rating are Excel formulas in the sheet (matching size_adj/dilmore's
existing convention). This module opens the workbook with data_only=True
and reads already-computed results -- it performs no calculation itself,
so the workbook must have gone through a real Excel/LibreOffice
recalculation pass for these cells to have cached values.

Called automatically by axiom.py during the deliver stage, gated by the
same `inject_comps` doc_cfg flag as inject_comp_section/inject_media_blocks/
inject_ownership_history. Can also be run standalone:
    python adjustment_grid.py <report_path> <workbook_path>
"""

import sys
from pathlib import Path

from docx import Document
import openpyxl


# Each grid tab's header row and first data row. These match the tabs as
# built (Phase 6 steps 1, 2, 4). MAX_DATA_ROW gives headroom to keep
# reading further populated rows if the sheet is ever hand-expanded beyond
# its current 10-comp capacity, without a code change.
HEADER_ROW = 6
FIRST_DATA_ROW = 7
MAX_DATA_ROW = 40

# The one fixed expectation every grid sheet must satisfy: its header row's
# first column must literally read "Comp". Everything else is discovered
# from the header row text, not a hardcoded position -- if this anchor is
# missing, the sheet doesn't match what this injector expects, and it fails
# loudly (AdjustmentGridError) instead of silently reading garbage, per the
# design doc's drift-protection requirement.
ANCHOR_HEADER = "Comp"

GRIDS = {
    "SCA_ADJUSTMENT_GRID_BLOCK": "sca_adjustment_grid",
    "LAND_ADJUSTMENT_GRID_BLOCK": "land_adjustment_grid",
    "SCA_QUALITATIVE_GRID_BLOCK": "sca_qualitative",
    "LAND_QUALITATIVE_GRID_BLOCK": "land_qualitative",
}

# Excel's fixed set of formula-error tokens. openpyxl (data_only=True) hands
# these back as plain strings when a formula errored the last time the
# workbook was recalculated (e.g. Ratio (Ac/As) with a blank Comp GBA divides
# by zero). Left undetected, one of these would render verbatim in a
# delivered comp grid, indistinguishable from a legitimate text cell like a
# Flood Zone or Notes entry -- caught by the Fable adversarial review of this
# module (finding A3).
EXCEL_ERROR_TOKENS = {
    "#DIV/0!", "#N/A", "#NAME?", "#NULL!", "#NUM!", "#REF!", "#VALUE!",
    "#SPILL!", "#CALC!", "#GETTING_DATA",
}

# Common non-ISO date strings a comp's Sale Date might be typed in as plain
# text rather than a real Excel date (e.g. pasted from a source document).
_TEXT_DATE_FORMATS = ("%Y-%m-%d", "%m/%d/%Y", "%m/%d/%y", "%B %d, %Y")


class AdjustmentGridError(Exception):
    """Raised when a grid sheet's header row doesn't match expectations, or
    when it has more populated comp rows than this module read (see
    read_grid_rows's MAX_DATA_ROW truncation check)."""


def _format_value(header, value):
    """Render a computed cell value for display, guided by its header text
    (not a fixed column position, since positions vary by preset)."""
    if value is None or value == "":
        return ""
    if isinstance(value, bool):
        return str(value)
    if isinstance(value, str) and value.strip().upper() in EXCEL_ERROR_TOKENS:
        # Surface loudly rather than let a raw Excel error token slip into
        # a delivered report looking like an ordinary text cell.
        return f"[FORMULA ERROR -- {value.strip()}]"
    if isinstance(value, (int, float)):
        header_lower = header.lower()
        if "$" in header:
            # Checked before "%": a header could in principle contain both
            # (none currently do), and a dollar amount must never be scaled
            # by 100. Sign goes before the "$" (e.g. "-$2.32"), not after
            # it (Python's default f"${value:,.2f}" on a negative number
            # produces the confusing "$-2.32").
            sign = "-" if value < 0 else ""
            return f"{sign}${abs(value):,.2f}"
        if "%" in header or (
            header_lower.endswith("rate") and "$" not in header
        ):
            # Appraisal "rate" columns without an explicit unit (e.g.
            # "Monthly Mkt Rate") are, in every real preset, a fractional
            # rate meant to display as a percentage just like an explicit
            # "... Adj %" column -- e.g. 0.005 must read "0.5%", not be
            # silently rounded away to "0.01" by the plain-number branch
            # below (Fable finding A4).
            return f"{value * 100:.1f}%"
        if header_lower == "overall":
            return f"{value:.2f}"
        if isinstance(value, int):
            return f"{value:,}"
        return f"{value:,.2f}"
    if hasattr(value, "strftime"):
        return value.strftime("%m/%d/%Y")
    if isinstance(value, str) and "date" in header.lower():
        # A date typed as plain text (not a real Excel date cell) skips the
        # strftime branch above entirely and used to render however it was
        # typed (e.g. ISO "2025-03-15") instead of the "%m/%d/%Y" every real
        # date cell gets (Fable finding A4). Best-effort reparse; if it
        # doesn't match a known format, fall through and show it as typed
        # rather than fail the whole grid over one comp's date.
        import datetime

        text = value.strip()
        for fmt in _TEXT_DATE_FORMATS:
            try:
                return datetime.datetime.strptime(text, fmt).strftime("%m/%d/%Y")
            except ValueError:
                continue
    return str(value).strip()


def _read_header_map(ws):
    """Build {header_text: column_index} from HEADER_ROW, skipping blanks."""
    header_map = {}
    for cell in ws[HEADER_ROW]:
        text = cell.value
        if text is None or str(text).strip() == "":
            continue
        header_map[str(text).strip()] = cell.column

    if header_map.get(ANCHOR_HEADER) != 1:
        raise AdjustmentGridError(
            f"Sheet '{ws.title}' header row {HEADER_ROW} doesn't have "
            f"'{ANCHOR_HEADER}' in column A as expected -- refusing to "
            "guess at a different layout. Check the sheet hasn't been "
            "restructured, or update adjustment_grid.py's ANCHOR_HEADER "
            "expectation if this is an intentional template change."
        )
    return header_map


def read_grid_rows(workbook_path, sheet_name):
    """Read a grid sheet's header row and populated data rows.

    Returns (headers, rows): headers is an ordered list of column header
    strings (left to right, as they appear in the sheet); rows is a list of
    dicts {header: formatted_value_str}, one per populated comp row.

    A row counts as a real comp row only if its Comp/anchor cell literally
    starts with "Sale No." -- every grid tab writes that label into each of
    its fixed comp rows when the tab is built, so it's the one reliable
    signal that distinguishes a genuine comp row from a summary row below
    the comp rows (MEAN, LAND VALUE CONCLUSION, etc.), which can otherwise
    have non-blank values in the same header columns by coincidence of
    position, not because they're actually comp data. Scanning stops at the
    first row whose Comp cell doesn't match, rather than scanning all the
    way to MAX_DATA_ROW and collecting every row with any value -- the
    latter approach collected the MEAN and LAND VALUE CONCLUSION rows on
    land_adjustment_grid's first real test against realistic fixture data.
    """
    wb = openpyxl.load_workbook(str(workbook_path), data_only=True)
    if sheet_name not in wb.sheetnames:
        wb.close()
        return [], []

    ws = wb[sheet_name]
    header_map = _read_header_map(ws)
    headers = [h for h, _ in sorted(header_map.items(), key=lambda kv: kv[1])]
    anchor_col = header_map[ANCHOR_HEADER]

    rows = []
    for r in range(FIRST_DATA_ROW, MAX_DATA_ROW + 1):
        comp_label = ws.cell(row=r, column=anchor_col).value
        if not (isinstance(comp_label, str) and comp_label.startswith("Sale No.")):
            break

        row_values = {}
        any_non_label_value = False
        for header, col_idx in header_map.items():
            value = ws.cell(row=r, column=col_idx).value
            if header != ANCHOR_HEADER and value not in (None, ""):
                any_non_label_value = True
            row_values[header] = _format_value(header, value)
        if not any_non_label_value:
            continue
        rows.append(row_values)

    # If a comp row still matches the "Sale No." anchor immediately past
    # MAX_DATA_ROW, the sheet has been hand-expanded beyond this module's
    # read window and comps are being silently dropped from every delivered
    # report -- caught by the Fable adversarial review (finding A5). Fail
    # loudly rather than quietly deliver an incomplete comp grid.
    overflow_label = ws.cell(row=MAX_DATA_ROW + 1, column=anchor_col).value
    if isinstance(overflow_label, str) and overflow_label.startswith("Sale No."):
        wb.close()
        raise AdjustmentGridError(
            f"Sheet '{sheet_name}' has a comp row at row {MAX_DATA_ROW + 1}, "
            f"past this module's MAX_DATA_ROW ({MAX_DATA_ROW}) -- comps "
            "would be silently dropped from the delivered report. Raise "
            "adjustment_grid.py's MAX_DATA_ROW to cover the expanded sheet."
        )

    wb.close()
    return headers, rows


def _build_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        if header_cells[idx].paragraphs[0].runs:
            header_cells[idx].paragraphs[0].runs[0].bold = True

    for row in rows:
        cells = table.add_row().cells
        for idx, header in enumerate(headers):
            cells[idx].text = row.get(header, "")

    return table


def inject_adjustment_grid_block(doc_path, workbook_path, block_name, sheet_name):
    """Replace ``[[<block_name>]]`` with a table built from *sheet_name*.

    Returns the number of comp rows injected. Returns 0 (leaving the marker
    text in place) if the marker isn't found in the document, the sheet
    doesn't exist in the workbook, or the sheet has no populated rows --
    matching inject_ownership_history's convention of a silent no-op rather
    than an error for "nothing to inject yet". A header-row mismatch
    (AdjustmentGridError) is not caught here -- it propagates, per the
    design doc's loud-failure requirement for a sheet that doesn't match
    what this injector expects.
    """
    doc_path = Path(doc_path)
    doc = Document(str(doc_path))
    marker = f"[[{block_name}]]"

    marker_paragraph = None
    for paragraph in doc.paragraphs:
        if marker in paragraph.text:
            marker_paragraph = paragraph
            break

    if marker_paragraph is None:
        return 0

    headers, rows = read_grid_rows(workbook_path, sheet_name)
    if not rows:
        return 0

    table = _build_table(doc, headers, rows)
    marker_paragraph._p.addnext(table._tbl)
    parent = marker_paragraph._p.getparent()
    parent.remove(marker_paragraph._p)

    doc.save(str(doc_path))
    return len(rows)


def inject_all_adjustment_grids(doc_path, workbook_path):
    """Inject all four adjustment-grid blocks found in *doc_path*.

    Returns {block_name: rows_injected}. Each block is handled
    independently -- a missing/empty sheet for one grid doesn't stop the
    others from being injected.
    """
    results = {}
    for block_name, sheet_name in GRIDS.items():
        results[block_name] = inject_adjustment_grid_block(
            doc_path, workbook_path, block_name, sheet_name
        )
    return results


# ── Standalone entry point ────────────────────────────────────────────────────

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python adjustment_grid.py <report.docx> <workbook.xlsx>")
        sys.exit(1)

    report_path = Path(sys.argv[1])
    workbook_path = Path(sys.argv[2])

    if not report_path.exists():
        print(f"Error: report not found: {report_path}")
        sys.exit(1)
    if not workbook_path.exists():
        print(f"Error: workbook not found: {workbook_path}")
        sys.exit(1)

    grid_results = inject_all_adjustment_grids(report_path, workbook_path)
    for block, count in grid_results.items():
        print(f"{block}: {count} row(s) injected")
