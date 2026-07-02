"""Inject structured, non-media report components into Word documents."""

from pathlib import Path

from docx import Document


OWNERSHIP_HISTORY_BLOCK = "OWNERSHIP_HISTORY_TABLE"
OWNERSHIP_HISTORY_FIELDS = (
    ("Owner of Record", "OWNER_NAME"),
    ("Three-Year Transfer History", "PRIOR_SALE_DATE"),
    ("Prior Sale Price", "PRIOR_SALE_PRICE"),
)


def ownership_history_missing_fields(variables):
    """Return required ownership fields that have no usable value."""
    return [
        key
        for _, key in OWNERSHIP_HISTORY_FIELDS
        if not str(variables.get(key, "")).strip()
    ]


def inject_ownership_history(doc_path, variables):
    """Replace ``[[OWNERSHIP_HISTORY_TABLE]]`` with a formatted Word table."""
    doc_path = Path(doc_path)
    doc = Document(str(doc_path))
    marker = f"[[{OWNERSHIP_HISTORY_BLOCK}]]"
    injected = False

    for paragraph in doc.paragraphs:
        if marker not in paragraph.text:
            continue

        missing = ownership_history_missing_fields(variables)
        if missing:
            break

        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for label, key in OWNERSHIP_HISTORY_FIELDS:
            cells = table.add_row().cells
            cells[0].text = label
            cells[0].paragraphs[0].runs[0].bold = True
            cells[1].text = str(variables[key])

        paragraph._p.addnext(table._tbl)
        parent = paragraph._p.getparent()
        parent.remove(paragraph._p)
        injected = True
        break

    if injected:
        doc.save(str(doc_path))
    return injected
